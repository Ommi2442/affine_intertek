import platform
import subprocess
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import json
from utility.trf_report.trf_generation import *
from azure.storage.blob import ContentSettings
from fastapi import HTTPException


def convert_docx_to_pdf(docx_path: str, pdf_path: str):
    print('*************** 3.1 ******************************', docx_path)

    print('*************** 3.2 ******************************', pdf_path)


    system = platform.system().lower()

    print('*************** 3.3 ******************************', system)
    

    # ---------- WINDOWS (MS WORD via COM) ----------
    if system == "windows":
        import pythoncom
        from docx2pdf import convert
        pythoncom.CoInitialize()     # ✅ REQUIRED
        try:
            convert(docx_path, pdf_path)
        finally:
            pythoncom.CoUninitialize()  # ✅ REQUIRED
        return

    # ---------- LINUX (LIBREOFFICE) ----------
    if system == "linux":
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                os.path.dirname(pdf_path),
                docx_path,
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        return

    raise RuntimeError(f"Unsupported OS for DOCX conversion: {system}")



def fetch_final_json_record(container, project_id: str) -> dict:
    query = """
        SELECT * FROM c
        WHERE c.project_id = @pid
        AND c.file_type = ".json"
    """

    params = [{"name": "@pid", "value": project_id}]

    items = list(
        container.query_items(
            query=query,
            parameters=params,
            enable_cross_partition_query=True
        )
    )

    # ---------- HARD GUARANTEE ----------
    if not items:
        raise HTTPException(
            status_code=404,
            detail=f"No final JSON record found in Cosmos for project_id={project_id}"
        )

    record = items[0]

    # ---------- SCHEMA VALIDATION ----------
    if "blob_path" not in record or "blob_url" not in record:
        raise HTTPException(
            status_code=500,
            detail="Cosmos record is invalid: blob_path or blob_url missing"
        )

    return record




def replace_json_blob(
    blob_service,
    container_name: str,
    blob_path: str,
    json_data: dict
):
    blob_client = blob_service.get_blob_client(
        container=container_name,
        blob=blob_path
    )

    blob_client.upload_blob(
        json.dumps(json_data, indent=2, ensure_ascii=False),
        overwrite=True,
        content_settings=ContentSettings(
            content_type="application/json"
        )
    )



def replace_local_final_json(project_id: str, json_data: dict):
    BASE_DIR = Path(__file__).resolve().parent.parent
    DATA_DIR = BASE_DIR / "data" / project_id
    DATA_DIR.mkdir(parents=True, exist_ok=True)

    final_path = DATA_DIR / "final_output.json"

    with open(final_path, "w", encoding="utf-8") as f:
        json.dump(json_data, f, indent=2, ensure_ascii=False)

    return final_path




def update_docx_tables_from_json_arial(docx_path, json_path, output_path):
    """
    Update DOCX tables using JSON values.
    - Updates only cells that have ai_fillable = True OR task_type = verdict_dependency.
    - Applies Arial 10 font to updated cells.
    - Saves to output_path (you will call the same name again to avoid temporary versions).
    """

    doc = Document(docx_path)

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    for table_obj in data["Tables"]:
        table_index = table_obj["Table"]

        # If table index missing → skip
        if table_index >= len(doc.tables):
            print(f"⚠️ Skipping missing table index: {table_index}")
            continue

        doc_table = doc.tables[table_index]

        for item in table_obj["Items"]:

            ai_fillable = item.get("ai_fillable", False)
            task_type = item.get("task_type")

            # Only update allowed fields
            if not ai_fillable and task_type != "verdict_dependency":
                continue

            row = item.get("answer_row")
            col = item.get("answer_column")
            value = item.get("value")

            if value is None:
                continue

            try:
                cell = doc_table.cell(row, col)
                cell.text = ""  # clear

                p = cell.paragraphs[0]
                run = p.add_run(str(value))
                font = run.font
                font.name = "Arial"
                font.size = Pt(10)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

            except Exception as e:
                print(f"⚠️ Error updating table {table_index} cell ({row},{col}): {e}")

    doc.save(output_path)
    print(f"✅ DOCX first update saved → {output_path}")



def table2_json_rows_only(json_path):
    """
    Reads Table 2 from JSON and returns ONLY rows (no header).
    """
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Find Table 2
    table_2 = next(t for t in data["Tables"] if t["Table"] == 2)

    columns = []
    for item in table_2["Items"]:
        if isinstance(item.get("value"), list):
            columns.append(item["value"])

    # Transpose columns → rows
    df = pd.DataFrame(list(zip(*columns)))

    return df

from docx import Document
from docx.shared import Pt

def insert_rows_into_attachments_table(
    input_docx,
    output_docx,
    df,
    anchor_text
):
    if df is None or df.empty:
        print("⚠️ DataFrame is empty. Skipping insertion.")
        return

    doc = Document(input_docx)
    target_table = None

    # 1️⃣ Find the table containing the anchor text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if anchor_text.strip() in cell.text:
                    target_table = table
                    break
            if target_table:
                break
        if target_table:
            break

    if target_table is None:
        raise ValueError("Anchor table not found.")

    # 2️⃣ Remove all rows BELOW header row (keep title + header)
    while len(target_table.rows) > 2:
        target_table._tbl.remove(target_table.rows[2]._tr)

    # 3️⃣ Insert rows immediately under header
    for _, row_data in df.iterrows():
        cells = target_table.add_row().cells

        for i, value in enumerate(row_data):
            cell = cells[i]
            cell.text = ""  # clear default paragraph

            p = cell.paragraphs[0]
            run = p.add_run(str(value))

            # 🔤 FONT SETTINGS
            run.font.name = "Arial"
            run.font.size = Pt(10)

    doc.save(output_docx)
    print("✅ Rows inserted in Arial, font size 7, with no extra space.")

import json
import pandas as pd

def table3_json_to_df(json_path):
    """
    Creates a row-ready DataFrame for Table 3.
    No header. Columns are positional.
    """

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    table_3 = next(t for t in data["Tables"] if t["Table"] == 3)

    # If values exist later, collect them
    columns = []
    for item in table_3["Items"]:
        if "rendering_column" in item:
            values = item.get("value")
            if isinstance(values, list):
                columns.append(values)

    # If JSON has no rows yet → return empty DF with 3 columns
    if not columns:
        return pd.DataFrame(columns=[0, 1, 2])

    # Otherwise transpose into rows
    df = pd.DataFrame(list(zip(*columns)))
    return df

from docx import Document
from docx.shared import Pt

def insert_rows_into_table_by_anchor_no_gap(
    input_docx,
    output_docx,
    df,
    anchor_text
):
    if df is None or df.empty:
        print("⚠️ DataFrame empty. No rows inserted.")
        return

    doc = Document(input_docx)
    target_table = None

    # 1️⃣ Locate table by anchor text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if anchor_text.strip() in cell.text:
                    target_table = table
                    break
            if target_table:
                break
        if target_table:
            break

    if target_table is None:
        raise ValueError("Target table not found for anchor.")

    # 2️⃣ Remove all rows below header
    while len(target_table.rows) > 2:
        target_table._tbl.remove(target_table.rows[2]._tr)

    # 3️⃣ Insert rows directly under header (Arial, size 7)
    for _, row_data in df.iterrows():
        cells = target_table.add_row().cells

        for i, value in enumerate(row_data):
            cell = cells[i]
            cell.text = ""  # clear default content

            p = cell.paragraphs[0]
            run = p.add_run(str(value))

            # 🔤 FONT SETTINGS
            run.font.name = "Arial"
            run.font.size = Pt(10)

    doc.save(output_docx)
    print("✅ Table updated (Arial, size 7) with no gaps.")

def run_table_insert_pipeline(
    json_path,
    input_docx,
    output_docx
):
    """
    Orchestrates Table 2 and Table 3 insertion
    using already-defined helper functions.
    """

    # ----------------------------
    # TABLE 2
    # ----------------------------
    df_table_2 = table2_json_rows_only(json_path)

    insert_rows_into_attachments_table(
        input_docx=input_docx,
        output_docx=output_docx,
        df=df_table_2,
        anchor_text="List of Attachments (including a total number of pages in each attachment)"
    )

    # ----------------------------
    # TABLE 3
    # ----------------------------
    df_table_3 = table3_json_to_df(json_path)

    insert_rows_into_table_by_anchor_no_gap(
        input_docx=output_docx,  # IMPORTANT: chain output
        output_docx=output_docx,
        df=df_table_3,
        anchor_text="Documents referenced by this report (available on request):"
    )

    print("✅ Table 2 and Table 3 pipeline completed.")


def update_docx_from_existing_json(
    input_docx_path: str,
    input_json_path: str,
    output_docx_path: str,
    projectId: str,
):

    print("\n===============================================")
    print("   JSON → DOCX POPULATION (NO LLM)")
    print("===============================================")

    # ---------------------------------------------------------
    # STEP 1 — LOAD JSON
    # ---------------------------------------------------------
    with open(input_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # ---------------------------------------------------------
    # STEP 2 — APPLY JSON POST-PROCESSING RULES
    # (safe + deterministic)
    # ---------------------------------------------------------

    print("\n[STEP] Applying JSON-only post processing...")

    # Prefix rules
    targets = [
        "General product information and other remarks:\nDescription of unit:\n",
        "Description of model differences:\n",
        "Description of special features:\n(HV circuits, high pressure systems etc.)\n"
    ]

    # data = apply_prefixes_to_items(data, targets)
    # data = prefix_summary_of_compliance(data)

    # Attach blob URLs
    # data = attach_blob_urls_to_text_support(data, blob_urls)
    # data = attach_blob_urls_to_image_support(data, blob_urls)

    # ---------------------------------------------------------
    # STEP 3 — SAVE FINAL JSON (OPTIONAL BUT SAFE)
    # ---------------------------------------------------------
    # export_results_to_json(data, input_json_path)

    # ---------------------------------------------------------
    # STEP 4 — UPDATE DOCX TABLE VALUES (Arial 10)
    # ---------------------------------------------------------
    print("\n[STEP] Filling DOCX tables from JSON...")

    update_docx_tables_from_json_arial(
        docx_path=input_docx_path,
        json_path=input_json_path,
        output_path=output_docx_path
    )

    # ---------------------------------------------------------
    # STEP 5 — APPLY CHECKBOX VALUES (INDEX BASED)
    # ---------------------------------------------------------
    print("\n[STEP] Applying checkbox ticks...")
    apply_checkboxes_from_json(
        docx_path=output_docx_path,
        output_path=output_docx_path,
        json_data=data
    )

    insert_legacy_checkbox_with_text(
        docx_path=output_docx_path,     # update in-place
        output_path=output_docx_path,   # overwrite same file
        sentence='The product fulfils the requirements of IEC 61010-1:2010, IEC 61010-1:2010/AMD1:2016',
        table_index=5,
        row=12,
        col=0,
        new_lines=5
    )


    # ---------------------------------------------------------
    # STEP 6 — INSERT MARKING PLATE IMAGES (FROM JSON)
    # ---------------------------------------------------------
    print("\n[STEP] Inserting marking images...")

    BASE_DIR = Path(__file__).resolve().parent.parent

    project_data_dir = BASE_DIR / "data" / projectId

    image_paths = download_marking_images_from_json_new(data, project_data_dir)
    insert_marking_images_from_json_into_docx_new(
        docx_path=output_docx_path,
        output_path=output_docx_path,
        image_paths=image_paths,
        table_index=6,
        row=0,
        col=0,
        width_inches=2
    )

    run_table_insert_pipeline(
        input_json_path,
        output_docx_path,
        output_docx_path
    )

    print("\n✅ JSON → DOCX POPULATION COMPLETE")
    return output_docx_path


 
def delete_blobs_inside_folder(folder_path: str, container_client):
    # Normalize folder path to act like a prefix
    if not folder_path.endswith("/"):
        folder_path += "/"

    blobs = container_client.list_blobs(name_starts_with=folder_path)

    deleted_count = 0
    for blob in blobs:
        # Safety check: skip the "folder marker" if it exists
        if blob.name != folder_path:
            container_client.delete_blob(blob.name)
            deleted_count += 1

    return deleted_count > 0




def delete_local_project_outputs(project_id: str, reportType: str) -> bool:
    """
    Deletes selected files and folders inside:
    <BASE_DIR>/data/<project_id>/
    """
    BASE_DIR = Path(__file__).resolve().parent.parent
    project_dir = BASE_DIR / "data" / project_id

    if not project_dir.exists():
        return False

    if not project_dir.is_dir():
        raise RuntimeError(f"Expected directory but found file: {project_dir}")

    deleted = False

    if reportType == "TRF":
        target_items = [
            "final_output.json",
            "final_output.docx",
            "pta_final_6_3_1_part1_output.json",
            "pta_final_6_3_1_part2_output.json",
            "pta_final_6_3_1_part3_output.json",
            "pta_final_6_3_1_part4_output.json",
            "pta_final_6_3_1_part5_output.json",
            "src_files"
        ]

    elif reportType == "CDR":
        target_items = [
            f"iec_output_cdr_{project_id}.json",
            f"iec_output_sheet_{project_id}.xlsx",
        ]

    else:
        target_items = [
            f"letter_body_iec_output_{project_id}.json",
            f"letter_header_iec_output_{project_id}.json",
            f"letter_iec_output_{project_id}.docx",
            "letter_page_images",
            "trf_src_files",
            "letter_src_files",
            "non_conforming_images",
        ]

    for name in target_items:
        path = project_dir / name

        if path.exists():
            if path.is_file():
                path.unlink()
                deleted = True

            elif path.is_dir():
                shutil.rmtree(path)
                deleted = True
    print('Deleted')

    return deleted

