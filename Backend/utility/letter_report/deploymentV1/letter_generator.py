# ============================================================
# IMPORTS
# ============================================================

import os
import json
from azure.cosmos import CosmosClient
from langchain_openai import AzureChatOpenAI
from openai import AzureOpenAI
import re
import tempfile
import shutil
import requests
from urllib.parse import urlparse, unquote
from email import policy
from email.parser import BytesParser
import extract_msg
import uuid
from langchain_core.documents import Document
import io
import openpyxl
import xlrd
#from utils import *
# from config import *
# from ocr_image_processor import load_and_process_images
# from trf_essential import *
# from trf_utils import *
# from core import *
from utility.letter_report.deploymentV1.config import *
from utility.letter_report.deploymentV1.ocr_image_processor import load_and_process_images
from utility.letter_report.deploymentV1.trf_essential import *
from utility.letter_report.deploymentV1.trf_utils import *
from utility.letter_report.deploymentV1.core import *
from utility.letter_report.deploymentV1.prompts import get_iec61010_non_conformance_prompt

from azure.storage.blob import BlobClient
from azure.core.exceptions import ResourceNotFoundError, AzureError

import pandas as pd
import math
import copy
import time
from azure.cosmos import CosmosClient, PartitionKey, exceptions
import json, os
from azure.cosmos import CosmosClient, ConsistencyLevel
from typing import List, Dict, Any, Tuple
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
from langchain_azure_ai.vectorstores import AzureCosmosDBNoSqlVectorSearch
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from azure.cosmos import CosmosClient
from langchain_openai import AzureOpenAIEmbeddings
from operator import itemgetter
from langchain_core.runnables import (
    RunnableParallel, RunnableLambda, RunnableMap
)
from langchain_core.output_parsers import StrOutputParser
from langchain_openai import AzureChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from tenacity import retry, retry_if_exception_type, wait_exponential, stop_never, RetryCallState
from openai import RateLimitError  # Make sure this import exists
from types import SimpleNamespace
from concurrent.futures import ThreadPoolExecutor, as_completed
from azure.core.exceptions import HttpResponseError
import time
from langchain_community.callbacks import get_openai_callback

import requests
from openai import AzureOpenAI
from azure.storage.blob import BlobServiceClient

from azure.storage.blob import ContainerClient

import os
import pdfplumber
from fuzzywuzzy import fuzz

#START FROM HERE
import fitz

import platform


# def contains_prepared_for_table(pdf_path):
#     """
#     Check if PDF contains the specific 'Prepared For:' table with expected structure.
#     Looks for key indicators like name, company, address, phone, email pattern.
#     """
#     text = ""
#     with fitz.open(pdf_path) as doc:
#         for page in doc:
#             t = page.get_text()
#             if t:
#                 text += t + "\n"
    
#     # Look for "Prepared For:" header
#     if not re.search(r"Prepared\s*For\s*:", text, re.IGNORECASE):
#         return False
    
#     # Extract the section after "Prepared For:"
#     match = re.search(r"Prepared\s*For\s*:(.{0,500})", text, re.IGNORECASE | re.DOTALL)
#     if not match:
#         return False
    
#     section = match.group(1)
    
#     # Check for multiple indicators that this is the specific table structure
#     indicators = [
#         r"Gener8\s*LLC",  # Company name
#         r"\(\d{3}\)\s*\d{3}-\d{4}",  # Phone number pattern like (650) 940-9898
#         r"\w+@\w+\.\w+",  # Email pattern
#         r"San\s*Jose|CA\s*95134|USA",  # Address components
#         r"Consultant"  # Role/title
#     ]
    
#     # Require at least 3 of these indicators to be present in the section
#     matches = sum(1 for pattern in indicators if re.search(pattern, section, re.IGNORECASE))
    
#     return matches >= 3


def contains_prepared_for_table(pdf_path):
    """
    Check if PDF contains the specific 'Prepared For:' table with expected structure.
    Looks for key indicators like name, company, address, phone, email pattern.
    """
    """
    Detect Intertek Quote / Project Proposal PDF based on fixed template structure.
    Client-specific values are ignored.
    """

    full_text = ""

    with fitz.open(pdf_path) as doc:
        for page in doc:
            text = page.get_text()
            if text:
                full_text += text + "\n"

    # Normalize text
    text = full_text.lower()

    # Required structural keywords (format-level)
    required_patterns = [
        r"project\s+proposal",
        r"quote\s+no",
        r"project\s+name",
        r"compiled\s+by",
        r"date",
        r"prepared\s+for\s*:",
        r"prepared\s+by\s*:",
        r"intertek"
    ]

    matches = 0
    for pattern in required_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            matches += 1

    # Require at least 5 strong structural indicators
    return matches >= 5


def find_prepared_for_pdf(folder_path: str):
    """
    Find the PDF that contains the specific 'Prepared For:' table.
    Returns the full PDF path or None if not found.
    """
    for file in os.listdir(folder_path):
        if file.lower().endswith(".pdf"):
            pdf_path = os.path.join(folder_path, file)
            if contains_prepared_for_table(pdf_path):
                return pdf_path  # Return full path instead of just filename
    
    return None


import re
from PyPDF2 import PdfReader
def extract_prepared_for_generic(pdf_path: str) -> dict:
    """
    Extracts ONLY the Scope of Work from an Intertek proposal PDF.

    Returns:
        {
            "Scope of Work": "<scope text>" 
        }
        or
        {
            "Scope of Work": None
        }
    """

    text = "\n".join(page.extract_text() or "" for page in PdfReader(pdf_path).pages)

    match = re.search(
        r"Scope\s*of\s*Work:\s*\n([^\n]+)",
        text,
        re.IGNORECASE | re.MULTILINE
    )

    scope_of_work = match.group(1).strip() if match else None

    return {
        "Scope of Work": scope_of_work
    }

def process_quote_from_folder(folder_path: str):
    """
    Find the quote PDF in the folder and extract the 'Prepared For' information from it.
    
    Args:
        folder_path: Path to the folder containing PDFs
        
    Returns:
        Extracted fields from the quote PDF, or None if no quote PDF found
    """
    # Find the quote PDF
    quote_pdf_path = find_prepared_for_pdf(folder_path)
    
    if quote_pdf_path is None:
        print(f"No quote PDF found in folder: {folder_path}")
        return None
    
    print(f"Found quote PDF: {quote_pdf_path}")
    
    # Extract fields from the quote PDF
    fields = extract_prepared_for_generic(quote_pdf_path)
    
    return fields

# find_prepared_for_pdf
# extract_prepared_for_generic
# process_quote_from_folder
# update_scope_of_work_in_json
import json

def update_scope_of_work_in_json(input_json_path, output_json_path, extracted_data):
    """
    Reads input JSON file
    Updates value of:
      - key == "KEY1"
      - key == "<ETL Listing/CB/Other Evaluation>"
    using Scope of Work from extracted_data
    Writes updated JSON to output file
    """

    scope_of_work = extracted_data.get("Scope of Work", "")

    # Load input JSON
    with open(input_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Traverse pages and items
    for page in data.get("pages", []):
        for item in page.get("items", []):
            if item.get("key") in ["KEY1", "<ETL Listing/CB/Other Evaluation> of the «ProductType» «ProductCovModels»"]:
                item["value"] = scope_of_work

    # Write output JSON
    with open(output_json_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

    print(f"✅ Updated JSON saved to: {output_json_path}")

# replace_keys_with_values_no_format_change_all
# this will chnage filders from json. Run after UI recivees new data
from docx import Document

from docx import Document

def replace_keys_with_values_no_format_change_all(
    input_docx: str,
    output_docx: str,
    data: dict
):
    """
    Replaces placeholders with values WITHOUT changing formatting.
    Skips specific keys completely.
    """

    doc = Document(input_docx)

    # Keys to ignore (case-insensitive)
    SKIP_KEYS = {
        "non-conformance table",
        "critical components table"
    }

    # Build replacement map
    replacements = {}
    for page in data.get("pages", []):
        for item in page.get("items", []):
            key = str(item.get("key", "")).strip()
            value = str(item.get("value", "")).strip()

            if not key:
                continue

            # Skip excluded keys
            if key.lower() in SKIP_KEYS:
                continue

            replacements[key] = value
            replacements[f"<{key}>"] = value  # optional bracket form

    # --- RUN JOIN FALLBACK REPLACEMENT ---
    def replace_in_para(para):
        full_text = "".join(run.text for run in para.runs)

        replaced = False
        for old, new in replacements.items():
            if old in full_text:
                full_text = full_text.replace(old, new)
                replaced = True

        if not replaced:
            return

        # Push replaced text back while keeping formatting
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = full_text

    # Replace in paragraphs
    for para in doc.paragraphs:
        replace_in_para(para)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

    doc.save(output_docx)
 

from docx import Document



def delete_cosmos_container(
    endpoint: str,
    key: str,
    database_name: str,
    container_name: str
):
    """
    Deletes a Cosmos DB container.
    """

    client = CosmosClient(endpoint, credential=key)

    try:
        database = client.get_database_client(database_name)
        database.delete_container(container_name)
        
        print("\n===============================================")
        print(f" DELETING THE VECTOR COSMOS DB Container '{container_name}'")
        print("===============================================")

        print(f"Container '{container_name}' deleted successfully.")


    except ResourceNotFoundError:
        print(f"Container '{container_name}' or database '{database_name}' not found.")

    except Exception as e:
        raise RuntimeError(f"Failed to delete container: {e}")



def replace_keys_with_values_no_format_change_v2(
    input_docx,
    output_docx,
    data
):
    """
    Replaces placeholders with values WITHOUT changing formatting.
    Works even if Word splits placeholder across multiple runs.
    Only replaces keys where ai_fillable == True.
    """

    doc = Document(input_docx)

    # -------------------------------------------------
    # Build replacement map (ai_fillable only)
    # -------------------------------------------------
    replacements = {}

    for page in data.get("pages", []):
        for item in page.get("items", []):

            if item.get("ai_fillable") is not True:
                continue

            key = str(item.get("key", "")).strip()
            value = str(item.get("value", "")).strip()

            if not key:
                continue

            replacements[key] = value
            replacements[f"<{key}>"] = value  # optional bracket form

    # -------------------------------------------------
    # Run-join fallback replacement
    # -------------------------------------------------
    def replace_in_para(para):
        if not para.runs:
            return

        full_text = "".join(run.text for run in para.runs)

        for old, new in replacements.items():
            if old in full_text:
                full_text = full_text.replace(old, new)

        # Clear runs but keep formatting
        for run in para.runs:
            run.text = ""

        para.runs[0].text = full_text

    # -------------------------------------------------
    # Replace in paragraphs
    # -------------------------------------------------
    for para in doc.paragraphs:
        replace_in_para(para)

    # -------------------------------------------------
    # Replace in tables
    # -------------------------------------------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

    doc.save(output_docx)

# replace_keys_with_values_no_format_change_v2
# get_first_doc_or_docx

def get_first_doc_or_docx(folder_path):
    """
    Returns the full path of the first .doc or .docx file found in a folder.
    Returns None if no such file exists.
    """
    if not os.path.isdir(folder_path):
        raise FileNotFoundError(f"Folder not found: {folder_path}")

    for filename in sorted(os.listdir(folder_path)):
        if filename.lower().endswith((".doc", ".docx")):
            return os.path.join(folder_path, filename)

    return None

# extract_table1a

# format_critical_components_df
# insert_dataframe_below_anchor
# read_docx_full
# extract_iec61010_non_conformances_full_doc

# def insert_dataframe_below_anchor(
#     input_docx,
#     output_docx,
#     df,
#     anchor_text
# ):
#     doc = Document(input_docx)

#     anchor_paragraph = None
#     for para in doc.paragraphs:
#         if anchor_text in para.text:
#             anchor_paragraph = para
#             break

#     if anchor_paragraph is None:
#         raise ValueError("Anchor text not found in document.")

#     # -------------------------------------------------
#     # INSERT BLANK PARAGRAPH AFTER ANCHOR (XML SAFE)
#     # -------------------------------------------------
#     blank_p = OxmlElement("w:p")
#     anchor_paragraph._p.addnext(blank_p)

#     # -------------------------------------------------
#     # CREATE TABLE
#     # -------------------------------------------------
#     rows, cols = df.shape
#     table = doc.add_table(rows=rows + 1, cols=cols)
#     table.style = "Table Grid"

#     # -----------------------------
#     # Header row (Amber background)
#     # -----------------------------
#     for col_idx, col_name in enumerate(df.columns):
#         cell = table.rows[0].cells[col_idx]
#         cell.text = str(col_name)

#         tc_pr = cell._tc.get_or_add_tcPr()
#         shd = OxmlElement("w:shd")
#         shd.set(qn("w:fill"), "FFC000")  # Amber
#         tc_pr.append(shd)

#     # -----------------------------
#     # Data rows
#     # -----------------------------
#     for row_idx in range(rows):
#         for col_idx in range(cols):
#             table.rows[row_idx + 1].cells[col_idx].text = str(df.iat[row_idx, col_idx])

#     # -------------------------------------------------
#     # INSERT TABLE AFTER BLANK PARAGRAPH
#     # -------------------------------------------------
#     blank_p.addnext(table._element)

#     doc.save(output_docx)



from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def insert_dataframe_below_anchor(
    input_docx,
    output_docx,
    df,
    anchor_text
):
    doc = Document(input_docx)

    # -------------------------------------------------
    # FIND ANCHOR PARAGRAPH
    # -------------------------------------------------
    anchor_paragraph = None
    for para in doc.paragraphs:
        if anchor_text in para.text:
            anchor_paragraph = para
            break

    if anchor_paragraph is None:
        raise ValueError("Anchor text not found in document.")

    # -------------------------------------------------
    # INSERT BLANK PARAGRAPH AFTER ANCHOR (XML SAFE)
    # -------------------------------------------------
    blank_p = OxmlElement("w:p")
    anchor_paragraph._p.addnext(blank_p)

    # -------------------------------------------------
    # HANDLE NONE OR EMPTY DATAFRAME SAFELY
    # -------------------------------------------------
    if df is None or df.empty:
        # Optional: insert a message instead of a table
        msg_para = OxmlElement("w:p")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "No non-conformances were identified."
        run.append(text)
        msg_para.append(run)

        blank_p.addnext(msg_para)
        doc.save(output_docx)
        return

    # -------------------------------------------------
    # CREATE TABLE
    # -------------------------------------------------
    rows, cols = df.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.style = "Table Grid"

    # -----------------------------
    # Header row (Amber background)
    # -----------------------------
    for col_idx, col_name in enumerate(df.columns):
        cell = table.rows[0].cells[col_idx]
        cell.text = str(col_name)

        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "FFC000")  # Amber
        tc_pr.append(shd)

    # -----------------------------
    # Data rows
    # -----------------------------
    for row_idx in range(rows):
        for col_idx in range(cols):
            table.rows[row_idx + 1].cells[col_idx].text = str(
                df.iat[row_idx, col_idx]
            )

    # -------------------------------------------------
    # INSERT TABLE AFTER BLANK PARAGRAPH
    # -------------------------------------------------
    blank_p.addnext(table._element)

    doc.save(output_docx)




from docx import Document



from docx import Document
from pathlib import Path

def read_docx_full(path):
    """
    Reads a DOCX file using a relative or absolute path
    and returns full text content (paragraphs + tables).
    """

    # Resolve relative path safely
    path = Path(path).expanduser().resolve()

    doc = Document(path)
    content = []

    # Read normal paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            content.append(p.text.strip())

    # Read ALL tables (CRITICAL for IEC TRFs)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                content.append(" | ".join(row_text))

    return "\n".join(content)




def run_iec61010_non_conformance_extraction(
    docx_path,
    deployment_name,
    chunk_size=800,
    chunk_overlap=100,
    fuzzy_threshold=40
):
    """
    One-call function.
    Returns final DataFrame with REAL PDF page numbers.

    REQUIREMENTS:
    pip install docx2pdf pdfplumber rapidfuzz langchain-text-splitters python-docx
    """



def convert_docx_to_pdf(docx_path, pdf_path):
    system = platform.system().lower()
    if system == "windows":
        import pythoncom
        from docx2pdf import convert
        pythoncom.CoInitialize()   # 🔑 REQUIRED per-thread
        try:
            convert(docx_path, pdf_path)
        finally:
            pythoncom.CoUninitialize()

    # ---------------- imports ----------------
    from pathlib import Path
    from docx import Document
    from docx2pdf import convert
    import pdfplumber
    from rapidfuzz import fuzz
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    # ---------------- 1. DOCX → PDF ----------------
    pdf_path = Path(docx_path).with_suffix(".pdf")
    convert_docx_to_pdf(docx_path, pdf_path)

    # ---------------- 2. Read PDF pages ----------------
    pdf_pages = {}
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            pdf_pages[i] = (page.extract_text() or "").replace("\n", " ")

    # ---------------- 3. Read DOCX blocks ----------------
    doc = Document(docx_path)
    blocks = []

    para_idx = 0
    table_idx = 0

    for p in doc.paragraphs:
        if p.text.strip():
            para_idx += 1
            blocks.append({
                "text": p.text.strip(),
                "source_ref": f"Paragraph {para_idx}"
            })

    for table in doc.tables:
        table_idx += 1
        for r_idx, row in enumerate(table.rows, start=1):
            row_text = " | ".join(
                c.text.strip() for c in row.cells if c.text.strip()
            )
            if row_text:
                blocks.append({
                    "text": row_text,
                    "source_ref": f"Table {table_idx}, Row {r_idx}"
                })

    # # ---------------- 4. Map blocks → REAL PDF pages ----------------
    # for block in blocks:
    #     best_page = None
    #     best_score = 0

    #     for page_num, page_text in pdf_pages.items():
    #         score = fuzz.partial_ratio(
    #             block["text"][:300],   # anchor slice (REQUIRED)
    #             page_text
    #         )
    #         if score > best_score:
    #             best_score = score
    #             best_page = page_num

    #     block["pdf_page"] = best_page if best_score >= fuzzy_threshold else None
    # ---------------- 4. Map blocks → REAL PDF pages ----------------
    for block in blocks:
        best_page = None
        best_score = 0
    
        for page_num, page_text in pdf_pages.items():
            score = fuzz.partial_ratio(
                block["text"][:300],   # anchor slice (REQUIRED)
                page_text
            )
            if score > best_score:
                best_score = score
                best_page = page_num
    
        # 👇 ADD THIS DEBUG LINE HERE
        if best_score < fuzzy_threshold:
            print(
                "UNMAPPED:",
                block["source_ref"],
                "|",
                block["text"][:120]
            )
    
        block["pdf_page"] = best_page if best_score >= fuzzy_threshold else None


    # ---------------- 5. Split using RecursiveCharacterTextSplitter ----------------
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " | ", " ", ""]
    )

    document_chunks = []
    ref_id = 1

    for block in blocks:
        for chunk in splitter.split_text(block["text"]):
            document_chunks.append({
                "ref_id": ref_id,
                "text": chunk,
                "pdf_page": block["pdf_page"],
                "source_ref": block["source_ref"]
            })
            ref_id += 1

    # ---------------- 6. CALL YOUR EXISTING FUNCTION ----------------
    df = extract_iec61010_non_conformances_full_doc(
        document_chunks=document_chunks,
        deployment_name=deployment_name
    )

    return df

IEC_61010_NON_CONFORMANCE_PROMPT = get_iec61010_non_conformance_prompt()




# def extract_iec61010_non_conformances_full_doc(
#     document_text,
#     deployment_name
# ):
#     system_prompt = """
# You are an IEC 61010-1 CB Scheme compliance expert.

# TASK:
# - Review the FULL Test Report Form (TRF)
# - Compare each clause with IEC 61010-1 requirements
# - Identify ONLY NON-CONFORMANCES

# RULES:
# - Verdict = F → Non-conformance
# - Missing mandatory markings or documentation → Non-conformance
# - TBD = testing not yet performed → DO NOT list
# - N/A with justification → Ignore
# - Do NOT invent issues

# OUTPUT (STRICT JSON ARRAY):
# [
#   {
#     "clause": "5.1.3",
#     "requirement": "Equipment shall be marked with rated voltage, current, and frequency",
#     "finding": "No electrical ratings marked on the equipment"
#   }
# ]

# Return [] if no non-conformances exist.
# """
#     client = AzureOpenAI(
#         api_key=AOAI_KEY,
#         api_version=API_VERSION,
#         azure_endpoint=AOAI_ENDPOINT
#     )

#     response = client.chat.completions.create(
#         model=deployment_name,   # MUST be Azure deployment name
#         temperature=0,
#         messages=[
#             {"role": "system", "content": system_prompt},
#             {
#                 "role": "user",
#                 "content": f"""
# FULL TRF DOCUMENT:
# ------------------
# {document_text}
# ------------------
# """
#             }
#         ]
#     )

#     output = response.choices[0].message.content.strip()
#     findings = json.loads(output)

#     return pd.DataFrame(
#         findings,
#         columns=["clause", "requirement", "finding"]
#     ).rename(columns={
#         "clause": "Clause",
#         "requirement": "Requirement of the Clause",
#         "finding": "Remark and Findings"
#     })




def extract_iec61010_non_conformances_full_doc(
    document_chunks,
    deployment_name
):
    import json
    import re
    import pandas as pd

    def extract_json_array(text):
        """
        Safely extract the first JSON array from LLM output.
        Returns [] if nothing valid is found.
        """
        if not text:
            return []

        match = re.search(r"\[\s*{.*?}\s*\]", text, re.DOTALL)
        if not match:
            return []

        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            return []

    def normalize(text):
        """Normalize text for stable comparison."""
        return " ".join(text.split()).lower()

    # ---------- Build LLM input ----------
    flat_text = "\n".join(
        f"[REF:{c['ref_id']} | Page {c['pdf_page']}] {c['text']}"
        for c in document_chunks
    )

    system_prompt = IEC_61010_NON_CONFORMANCE_PROMPT

    client = AzureOpenAI(
        api_key=AOAI_KEY,
        api_version=API_VERSION,
        azure_endpoint=AOAI_ENDPOINT
    )

    response = client.chat.completions.create(
        model=deployment_name,
        temperature=0,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": flat_text}
        ]
    )

    raw_output = response.choices[0].message.content
    findings = extract_json_array(raw_output)

    # ---------- Return empty but typed DF ----------
    if not findings:
        return pd.DataFrame(
            columns=[
                "Clause",
                "Requirement of the Clause",
                "Remark and Findings",
                "Page Reference"
            ]
        )

    # ---------- Map ref_id → page ----------
    ref_map = {c["ref_id"]: c for c in document_chunks}

    for f in findings:
        ref = ref_map.get(f.get("ref_id"))
        f["Page Reference"] = (
            f"Page {ref['pdf_page']}"
            if ref and ref.get("pdf_page")
            else "Unknown"
        )

    return pd.DataFrame(findings).rename(columns={
        "clause": "Clause",
        "requirement": "Requirement of the Clause",
        "finding": "Remark and Findings"
    })




# download_files_from_blob
# extract_images_from_docx
# extract_images_from_excel
# identify_iec61010_non_conforming_images_batch
import os
from urllib.parse import urlparse, unquote
from azure.storage.blob import BlobServiceClient
from azure.core.exceptions import ResourceNotFoundError


def download_files_from_blob(
    blob_file_list,
    download_dir,
    connection_string,
    container_name
):
    """
    Downloads files from Azure Blob Storage.
    Supports:
      - Full HTTPS blob URLs (with or without SAS)
      - Plain blob paths
      - Nested folder structures
    """

    # ----------------------------------------
    # Validate & normalize container name
    # ----------------------------------------
    if not isinstance(container_name, str):
        raise ValueError(f"container_name must be string. Got: {type(container_name)}")

    container_name = container_name.strip()

    if container_name.startswith("http"):
        parsed = urlparse(container_name)
        container_name = parsed.path.strip("/").split("/")[0]

    if not container_name:
        raise ValueError("container_name is empty after normalization")

    print(f"📦 Using container: {container_name}")

    # ----------------------------------------
    # Prepare clients & directory
    # ----------------------------------------
    os.makedirs(download_dir, exist_ok=True)

    blob_service_client = BlobServiceClient.from_connection_string(connection_string)
    container_client = blob_service_client.get_container_client(container_name)

    downloaded_files = []

    # ----------------------------------------
    # Download loop
    # ----------------------------------------
    for blob_item in blob_file_list:

        if not isinstance(blob_item, str) or not blob_item.strip():
            raise ValueError(f"Invalid blob entry: {blob_item}")

        blob_item = blob_item.strip()

        # ------------------------------------
        # Normalize blob name (HTTPS or plain)
        # ------------------------------------
        if blob_item.startswith("http"):
            parsed = urlparse(blob_item)
            path = parsed.path.lstrip("/")  # container/path/to/blob

            prefix = f"{container_name}/"
            if not path.startswith(prefix):
                raise ValueError(
                    f"Blob URL does not belong to container '{container_name}': {blob_item}"
                )

            blob_name = unquote(path[len(prefix):])
        else:
            blob_name = unquote(blob_item.lstrip("/"))

        if not blob_name:
            raise ValueError(f"Resolved empty blob name from: {blob_item}")

        print(f"⬇️ Downloading blob: {blob_name}")

        # ------------------------------------
        # Local path (preserve folders)
        # ------------------------------------
        local_path = os.path.join(download_dir, blob_name)
        os.makedirs(os.path.dirname(local_path), exist_ok=True)

        blob_client = container_client.get_blob_client(blob_name)

        # ------------------------------------
        # Download (no extra exists() call)
        # ------------------------------------
        try:
            stream = blob_client.download_blob()
            with open(local_path, "wb") as f:
                f.write(stream.readall())
        except ResourceNotFoundError:
            raise FileNotFoundError(
                f"Blob NOT FOUND in container '{container_name}': {blob_name}"
            )

        downloaded_files.append(local_path)

    return downloaded_files




import os
from docx import Document

def extract_images_from_docx(docx_path, output_dir=None):
    """
    Extracts all images from a .docx file and saves them to output_dir
    (defaults to current working directory).
    """
    image_paths= []
    if output_dir is None:
        output_dir = os.getcwd()

    os.makedirs(output_dir, exist_ok=True)

    doc = Document(docx_path)
    image_count = 0

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
            image = rel.target_part.blob
            image_ext = rel.target_part.content_type.split("/")[-1]

            image_filename = f"extracted_image_{image_count}.{image_ext}"
            image_path = os.path.join(output_dir, image_filename)

            with open(image_path, "wb") as f:
                f.write(image)
            image_paths.append(image_path)

    return image_paths,image_count



import os
import subprocess
import platform
from openpyxl import load_workbook

# def extract_images_from_excel(excel_path, output_dir=None):
#     """
#     Extract images from Excel files.
#     - .xlsx / .xlsm: direct
#     - .xls: converted using LibreOffice
#     Works on Linux and Windows.
#     """

#     if output_dir is None:
#         output_dir = os.getcwd()

#     os.makedirs(output_dir, exist_ok=True)
#     excel_path = os.path.abspath(excel_path)
#     lower = excel_path.lower()

#     # ----------------------------------
#     # Locate LibreOffice executable
#     # ----------------------------------
#     if platform.system() == "Windows":
#         soffice = r"C:\Program Files\LibreOffice\program\soffice.exe"
#         if not os.path.exists(soffice):
#             soffice = r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
#         if not os.path.exists(soffice):
#             raise RuntimeError("LibreOffice is not installed or soffice.exe not found")
#     else:
#         soffice = "soffice"

#     # ----------------------------------
#     # Convert .xls -> .xlsx if needed
#     # ----------------------------------
#     if lower.endswith(".xls") and not lower.endswith(".xlsx"):
#         subprocess.run(
#             [
#                 soffice,
#                 "--headless",
#                 "--convert-to",
#                 "xlsx",
#                 "--outdir",
#                 output_dir,
#                 excel_path,
#             ],
#             check=True,
#         )

#         base = os.path.splitext(os.path.basename(excel_path))[0]
#         excel_path = os.path.join(output_dir, base + ".xlsx")

#         if not os.path.exists(excel_path):
#             raise RuntimeError("LibreOffice conversion failed")

#     # ----------------------------------
#     # Extract images from .xlsx
#     # ----------------------------------
#     if not excel_path.lower().endswith((".xlsx", ".xlsm")):
#         raise ValueError("Unsupported Excel format")

#     wb = load_workbook(excel_path)
#     extracted_paths = []
#     counter = 1

#     for sheet in wb.worksheets:
#         for img in getattr(sheet, "_images", []):
#             data = img._data()
#             ext = img.format.lower() if img.format else "png"

#             img_path = os.path.join(
#                 output_dir,
#                 f"excel_image_{counter}.{ext}"
#             )

#             with open(img_path, "wb") as f:
#                 f.write(data)

#             extracted_paths.append(img_path)
#             counter += 1

#     return extracted_paths




def extract_images_from_excel(excel_path, output_dir=None):
    """
    Extract images from Excel files.
    - .xlsx / .xlsm: direct
    - .xls: converted using OS-aware logic
      - Windows: MS Excel via COM
      - Linux: LibreOffice
    """

    import os
    import platform
    import subprocess
    from openpyxl import load_workbook

    if output_dir is None:
        output_dir = os.getcwd()

    os.makedirs(output_dir, exist_ok=True)
    excel_path = os.path.abspath(excel_path)
    lower = excel_path.lower()
    system = platform.system().lower()

    # ----------------------------------
    # Convert .xls -> .xlsx if needed
    # ----------------------------------
    if lower.endswith(".xls") and not lower.endswith(".xlsx"):

        base = os.path.splitext(os.path.basename(excel_path))[0]
        converted_path = os.path.join(output_dir, base + ".xlsx")

        # ---------- WINDOWS (MS EXCEL via COM) ----------
        if system == "windows":
            import pythoncom
            import win32com.client

            pythoncom.CoInitialize()  # ✅ REQUIRED
            try:
                excel = win32com.client.Dispatch("Excel.Application")
                excel.Visible = False
                excel.DisplayAlerts = False

                wb = excel.Workbooks.Open(excel_path)
                wb.SaveAs(converted_path, FileFormat=51)  # 51 = xlsx
                wb.Close()
                excel.Quit()
            finally:
                pythoncom.CoUninitialize()  # ✅ REQUIRED

        # ---------- LINUX (LIBREOFFICE) ----------
        elif system == "linux":
            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "xlsx",
                    "--outdir",
                    output_dir,
                    excel_path,
                ],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )

        else:
            raise RuntimeError(f"Unsupported OS for Excel conversion: {system}")

        if not os.path.exists(converted_path):
            raise RuntimeError("Excel conversion failed")

        excel_path = converted_path

    # ----------------------------------
    # Extract images from .xlsx / .xlsm
    # ----------------------------------
    if not excel_path.lower().endswith((".xlsx", ".xlsm")):
        raise ValueError("Unsupported Excel format")

    wb = load_workbook(excel_path)
    extracted_paths = []
    counter = 1

    for sheet in wb.worksheets:
        for img in getattr(sheet, "_images", []):
            data = img._data()
            ext = img.format.lower() if img.format else "png"

            img_path = os.path.join(
                output_dir,
                f"excel_image_{counter}.{ext}"
            )

            with open(img_path, "wb") as f:
                f.write(data)

            extracted_paths.append(img_path)
            counter += 1

    return extracted_paths

 



import base64
import json
import os
from openai import AzureOpenAI

def identify_iec61010_non_conforming_images_batch(
    image_paths,
    clauses_start=4,
    clauses_end=17
):
    """
    Batch IEC 61010 image compliance checker using Azure OpenAI SDK.
    Returns per-image verdict with non-conformance details.
    """

    client = AzureOpenAI(
        api_key=AOAI_KEY,
        api_version=API_VERSION,
        azure_endpoint=AOAI_ENDPOINT
    )

    deployment = CHAT_DEPLOY

    results = []

    for idx, image_path in enumerate(image_paths):

        # ----------------------------
        # Encode image
        # ----------------------------
        with open(image_path, "rb") as f:
            image_base64 = base64.b64encode(f.read()).decode("utf-8")

        system_prompt = f"""
You are an IEC 61010 compliance expert.

Evaluate ONLY the provided image.
Check IEC 61010 clauses {clauses_start} to {clauses_end}.
Identify NON-CONFORMANCES only.

Rules:
- Do NOT assume compliance
- Do NOT infer hidden measurements
- If visual evidence is insufficient, return INSUFFICIENT_EVIDENCE
- Be conservative and audit-safe

Return STRICT JSON ONLY:

{{
  "image_verdict": "NON_CONFORMING | CONFORMING | INSUFFICIENT_EVIDENCE",
  "non_conforming_clauses": [
    {{
      "clause": "IEC 61010-X",
      "issue": "Clear technical reason",
      "risk": "Safety / Fire / Electric Shock / Mechanical / Thermal"
    }}
  ],
  "confidence_score": 0.0 to 1.0,
  "notes": "Optional"
}}
"""

        try:
            response = client.chat.completions.create(
                model=deployment,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{image_base64}"
                                }
                            }
                        ]
                    }
                ],
                temperature=0.1,
                max_tokens=800
            )

            raw_text = response.choices[0].message.content

            try:
                parsed = json.loads(raw_text)
            except Exception:
                parsed = {
                    "image_verdict": "ERROR",
                    "raw_response": raw_text
                }

        except Exception as e:
            parsed = {
                "image_verdict": "ERROR",
                "error": str(e)
            }

        results.append({
            "image_index": idx,
            "image_path": image_path,
            "result": parsed
        })

    return results




from docx import Document
from docx.shared import Inches

def _normalize(text):
    return " ".join(text.replace("\u00a0", " ").split()).upper()

def insert_photos_before_section_6_table(
    docx_path,
    output_path,
    image_paths,
    image_width_inches=3
):
    doc = Document(docx_path)

    for table in doc.tables:
        table_text = " ".join(
            _normalize(cell.text)
            for row in table.rows
            for cell in row.cells
        )

        if "SECTION 6" in table_text or "PROJECT STATUS" in table_text:

            tbl_element = table._tbl
            parent = tbl_element.getparent()
            tbl_index = parent.index(tbl_element)

            insert_index = tbl_index

            for img_path in image_paths:
                # Paragraph with image
                img_para = doc.add_paragraph()
                img_run = img_para.add_run()
                img_run.add_picture(img_path, width=Inches(image_width_inches))
                parent.insert(insert_index, img_para._p)
                insert_index += 1

                # Blank paragraph for spacing (NEW LINE)
                spacer_para = doc.add_paragraph("")
                parent.insert(insert_index, spacer_para._p)
                insert_index += 1

            doc.save(output_path)
            return

    raise ValueError("SECTION 6 table anchor not found")



import os

def inject_dataframes_into_letter_json(data, df_critical, df_nonpass):
    """
    Injects dataframe dictionaries into JSON placeholders:
      - 'Non-conformance Table' -> df_nonpass
      - 'Critical components table' -> df_critical
    """
 
    # Convert DataFrames to dict format
    critical_dict = df_critical.to_dict(orient="records") if df_critical is not None else []
    nonpass_dict = df_nonpass.to_dict(orient="records") if df_nonpass is not None else []
 
    for page in data.get("pages", []):
        for item in page.get("items", []):
 
            # Inject Non-Conformance Table
            if item.get("key") == "Non-conformance Table" and item.get("dataframe_table") is True:
                item["value"] = nonpass_dict
 
            # Inject Critical Components Table
            if item.get("key") == "Critical components table" and item.get("dataframe_table") is True:
                item["value"] = critical_dict
 
    return data

def orchestrate_iec61010_image_compliance_pipeline(
    blob_file_list,
    local_download_dir,
    source_docx_name,
    output_docx_path,
    connection_string,
    container_name
):
    """
    End-to-end orchestration:
    1. Download files from blob
    2. Extract images from DOCX and XLS/XLSX into src_files
    3. Identify IEC 61010 non-conforming images
    4. Insert non-conforming images into DOCX before Section 6 table
    """

    # # ---------------------------------------------------------
    # # 1. Download files from blob
    # # ---------------------------------------------------------
    # print("⬇️ Downloading files from blob...")
    # downloaded_files = download_files_from_blob(
    #     blob_file_list=blob_file_list,
    #     download_dir=local_download_dir,
    #     connection_string=connection_string,
    #     container_name=container_name
    # )

    # ---------------------------------------------------------
    # Build downloaded_files list from local_download_dir
    # ---------------------------------------------------------
    downloaded_files = [
        os.path.join(local_download_dir, f)
        for f in os.listdir(local_download_dir)
        if os.path.isfile(os.path.join(local_download_dir, f))
    ]

    print("📂 Using local files:")
    for f in downloaded_files:
        print(" -", f)

    print(downloaded_files)

    # ---------------------------------------------------------
    # 2. Extract images into src_files directory
    # ---------------------------------------------------------
    print("🖼️ Extracting images from documents into src_files...")
    extracted_images = []

    for file_path in downloaded_files:
        file_lower = file_path.lower()

        if file_lower.endswith(".docx"):
            images, count = extract_images_from_docx(file_path, output_dir=local_download_dir)
            extracted_images.extend(images)

        elif file_lower.endswith((".xls", ".xlsx")):
            images = extract_images_from_excel(file_path, output_dir=local_download_dir)
            extracted_images.extend(images)

    if not extracted_images:
        print("⚠️ No images extracted from documents.")
        return

    print("\n📂 Extracted Images:")
    for img in extracted_images:
        print(" -", os.path.abspath(img))

    # ---------------------------------------------------------
    # 3. Identify IEC 61010 non-conforming images
    # ---------------------------------------------------------
    print("\n🔍 Checking IEC 61010 compliance...")
    results = identify_iec61010_non_conforming_images_batch(extracted_images)

    non_conforming_images = []

    print("\n🚨 Non-Conforming Images:")
    for r in results:
        if r["result"]["image_verdict"] == "NON_CONFORMING":
            print("❌ NON-CONFORMING:", r["image_path"])
            non_conforming_images.append(r["image_path"])

    if not non_conforming_images:
        print("✅ No non-conforming images found.")
        return

    # ---------------------------------------------------------
    # 4. Insert non-conforming images into letter.docx
    # ---------------------------------------------------------
    print("\n📄 Inserting non-conforming images into letter...")

    try:
        insert_photos_before_section_6_table(
            docx_path=output_docx_path,
            output_path=output_docx_path,
            image_paths=non_conforming_images
        )

        print("✅ Non-conforming images inserted into letter successfully.")

        return non_conforming_images

    except Exception as e:
        print("⚠️ Failed to insert images into letter.")
        print(f"Reason: {e}")

import uuid
from pathlib import Path
from azure.storage.blob import BlobServiceClient, ContentSettings

def upload_images_to_blob(
    non_conforming_images,
    connection_string,
    container_name,
    base_path,
    final_letter_images_path
):
    """
    Upload NON-CONFORMING images to:
    Documents/<project_Id>/letter_images/
    Returns list of blob URLs.
    Fully OS-agnostic.
    """

    blob_service_client = BlobServiceClient.from_connection_string(
        connection_string
    )
    container_client = blob_service_client.get_container_client(
        container_name
    )

    uploaded_blob_urls = []

    # Normalize blob base path (Azure requires forward slashes)
    blob_base_path = Path(final_letter_images_path).as_posix()

    for image in non_conforming_images:
        image_path = Path(image).resolve()

        if not image_path.exists():
            print(f"⚠️ Skipping missing image: {image_path}")
            continue

        unique_name = f"{uuid.uuid4().hex}_{image_path.name}"
        blob_path = f"{blob_base_path}/{unique_name}"

        blob_client = container_client.get_blob_client(blob_path)

        suffix = image_path.suffix.lower()
        content_type = (
            "image/png" if suffix == ".png"
            else "image/jpeg"
        )

        with image_path.open("rb") as f:
            blob_client.upload_blob(
                f,
                overwrite=True,
                content_settings=ContentSettings(
                    content_type=content_type
                )
            )

        uploaded_blob_urls.append(blob_client.url)

    return uploaded_blob_urls



# insert_photos_before_section_6_table
# orchestrate_iec61010_image_compliance_pipeline
# replace_header_keys_with_values_header

#citation with blob urls



import os
from urllib.parse import quote

def map_pdf_path_to_blob_url(
    pdf_path: str,
    account_name: str = "saaffine",
    container_name: str = "nasa-ebooks-pdfs-all"
) -> str | None:
    """
    Convert local PDF path to Azure Blob URL.
    Returns None if input is None.
    """

    if not pdf_path:
        return None

    filename = os.path.basename(pdf_path)
    blob_name = quote(filename)

    return (
        f"https://{account_name}.blob.core.windows.net/"
        f"{container_name}/{blob_name}"
    )

import json
import os
from copy import deepcopy

# def add_text_support_with_blob_url_and_filename(
#     letter_json_path: str,
#     folder_path: str,
#     page_number: int = 3
# ):
#     """
#     Reads existing letter JSON and adds deterministic text_support
#     using:
#       - pdf_path from find_prepared_for_pdf(folder_path)
#       - filename extracted from pdf_path
#       - blob URL from map_pdf_path_to_blob_url(pdf_path)
#     """

#     TARGET_KEYS = {
#         "KEY1",
#         "<ETL Listing/CB/Other Evaluation> of the «ProductType» «ProductCovModels»"
#     }

#     # 1️⃣ Find prepared-for PDF
#     pdf_path = find_prepared_for_pdf(folder_path)

#     if not pdf_path:
#         raise ValueError("No prepared-for PDF found.")

#     # 2️⃣ Extract filename from path
#     filename = os.path.basename(pdf_path)

#     # 3️⃣ Get blob URL
#     blob_url = map_pdf_path_to_blob_url(pdf_path)

#     with open(letter_json_path, "r", encoding="utf-8") as f:
#         letter = json.load(f)

#     letter = deepcopy(letter)  # safety

#     for page in letter.get("pages", []):
#         for item in page.get("items", []):

#             key = item.get("key")
#             value = item.get("value")

#             # Only act on target keys with real values
#             if key not in TARGET_KEYS:
#                 continue

#             if not value or not isinstance(value, str):
#                 continue

#             # Attach deterministic text support
#             item["text_support"] = [
#                 {
#                     "filename": filename,
#                     "page": page_number,
#                     "similarity_score": 1,
#                     "preview_text": "",
#                     "url": blob_url
#                 }
#             ]

#             # Set confidence
#             item["confidence"] = 100

#     return letter

def add_text_support_with_blob_url_and_filename(
    letter_json_path: str,
    folder_path: str,
    page_number: int = 3
):
    """
    Reads existing letter JSON and adds deterministic text_support
    using:
      - pdf_path from find_prepared_for_pdf(folder_path)
      - filename extracted from pdf_path
      - blob URL from map_pdf_path_to_blob_url(pdf_path)
    """

    TARGET_KEYS = {
        "KEY1",
        "<ETL Listing/CB/Other Evaluation> of the «ProductType» «ProductCovModels»"
    }

    # 1️⃣ Find prepared-for PDF
    pdf_path = find_prepared_for_pdf(folder_path)

    if not pdf_path:
        raise ValueError("No prepared-for PDF found.")

    # 2️⃣ Extract filename from path
    filename = os.path.basename(pdf_path)

    # 3️⃣ Get blob URL
    #blob_url = map_pdf_path_to_blob_url(pdf_path)

    with open(letter_json_path, "r", encoding="utf-8") as f:
        letter = json.load(f)

    letter = deepcopy(letter)  # safety

    for page in letter.get("pages", []):
        for item in page.get("items", []):

            key = item.get("key")
            value = item.get("value")

            # Only act on target keys with real values
            if key not in TARGET_KEYS:
                continue

            if not value or not isinstance(value, str):
                continue

            # Attach deterministic text support
            item["text_support"] = [
                {
                    "filename": filename,
                    "page": page_number,
                    "similarity_score": 1,
                    "preview_text": ""
                    #"url": blob_url
                }
            ]

            # Set confidence
            item["confidence"] = 100

    return letter

 

 
from docx import Document

def replace_header_keys_with_values_header(input_docx, output_docx, data):
    """
    Replaces specific header placeholders from page 3 onwards:
    - «AppCOMPANYNAME»
    - Intertek Report: No:
    """

    TARGET_KEYS = {"«AppCOMPANYNAME»", "Intertek Report: No:"}

    doc = Document(input_docx)

    # Build replacement map ONLY from page_number >= 3
    replacements = {}
    for page in data.get("pages", []):
        if page.get("page_number", 0) < 3:
            continue

        for item in page.get("items", []):
            placeholder = str(item.get("key", "")).strip()
            value = str(item.get("value", "")).strip()

            if placeholder in TARGET_KEYS and value:
                replacements[placeholder] = value

    # If nothing to replace, exit early
    if not replacements:
        doc.save(output_docx)
        return output_docx

    # Replace helper (UNCHANGED logic)
    def replace_in_header(header):
        for para in header.paragraphs:
            for run in para.runs:
                for old, new in replacements.items():
                    if old in run.text:
                        run.text = run.text.replace(old, new)

        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for old, new in replacements.items():
                                if old in run.text:
                                    run.text = run.text.replace(old, new)

    # Replace in ALL section headers (important for linked headers)
    for section in doc.sections:
        replace_in_header(section.header)

    doc.save(output_docx)
    return output_docx

import pandas as pd
import numpy as np

def clean_dataframe_for_json(df, use_null=True):
    """
    Replace NaN values so JSON does not contain NaN.
    use_null=True  -> NaN becomes None (JSON null)
    use_null=False -> NaN becomes empty string ""
    """
    if df is None:
        return None

    df = df.copy()

    if use_null:
        df = df.replace({np.nan: None})
    else:
        df = df.fillna("")

    return df



 
#entry function

def generate_letter_pipeline(
    blob_urls,
    container_name,
    src_files_dir,
    letter_json_path,
    letter_header_json_path,
    letter_template_docx,
    output_letter_docx,
    letter_json_path_output,
    letter_header_json_path_output,
    project_Id,
    text_container,
    image_container
    ):
    """
    Single entry function for full letter generation pipeline.
    """

    print("\n==============================")
    print("🚀 LETTER GENERATION PIPELINE")
    print("==============================\n")

    # -------------------------------------------------------
    # VECTOR DATABASE CONNECTION
    # -------------------------------------------------------

    embeddings = build_embeddings(AOAI_ENDPOINT, AOAI_KEY, API_VERSION, EMBED_DEPLOY)

    vs = build_vectorstore(
        embeddings,
        COSMOS_URL,
        COSMOS_KEY,
        DB_NAME,
        text_container
    )

    vs2 = build_vectorstore2(
        embeddings,
        COSMOS_URL,
        COSMOS_KEY,
        DB_NAME_IMG,
        image_container
    )

    text_retriever_agent = vs.as_retriever(search_kwargs={"k": 5})

    llm = AzureChatOpenAI(
        azure_endpoint=AOAI_ENDPOINT,
        api_key=AOAI_KEY,
        openai_api_version=API_VERSION,
        azure_deployment=CHAT_DEPLOY,
        temperature=0.1,
    ).with_config({"response_format": "json_object"})



    rag_image = build_rag_image_pipeline_v5(
        text_retriever_agent,
        llm,
        build_vision_message_v5,
        attach_supporting_refs_grey,
        vs
    )

    # -------------------------------------------------------
    # STEP 1 — Fill Letter JSON
    # -------------------------------------------------------

    with open(letter_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    with open(letter_header_json_path, "r", encoding="utf-8") as f:
        data_header = json.load(f)

    # -------------------------------------------------------
    # STEP 2 — Build Tasks
    # -------------------------------------------------------

    tasks, item_refs = build_tasks_with_custom_prompt_letter(data)
    image_retriever_agent = vs2.as_retriever(search_kwargs={"k": 5})
    tasks = update_tasks_with_top5_images(tasks, image_retriever_agent)

    print(tasks[0])
    print("###############################")
    print(tasks[1])

    tasks_header, item_refs_header = build_tasks_with_custom_prompt_letter(data_header)
    tasks_header = update_tasks_with_top5_images(tasks_header, image_retriever_agent)

    # -------------------------------------------------------
    # STEP 3 — Run RAG
    # -------------------------------------------------------

    results = process_tasks_with_batches_parallel_grey(
        tasks,
        item_refs,
        rag_image,
        vs,
        batch_size=150,
        cooldown_sec=15,
        max_workers=6,
        use_llm_inGrey=False,
        stats=True
    )

    results_header = process_tasks_with_batches_parallel_grey(
        tasks_header,
        item_refs_header,
        rag_image,
        vs,
        batch_size=150,
        cooldown_sec=15,
        max_workers=6,
        use_llm_inGrey=False,
        stats=True
    )
    
    # -------------------------------------------------------
    # STEP 4 — Update Letter DOCX
    # -------------------------------------------------------

    replace_keys_with_values_no_format_change_v2(
        input_docx=letter_template_docx,
        output_docx=output_letter_docx,
        data=data
    )

    replace_header_keys_with_values_header(
        input_docx=output_letter_docx,
        output_docx=output_letter_docx,
        data=data_header
    )


    # downloaded_files = download_files_from_blob(
    #     blob_file_list=blob_urls,
    #     download_dir=src_files_dir,
    #     connection_string=AZURE_CONN_STRING,
    # container_name=BLOB_CONTAINER_NAME
    # )

    # -------------------------------------------------------
    # STEP 5 — Extract Critical Components
    # -------------------------------------------------------
    doc_path = get_first_doc_or_docx(src_files_dir)
    docx_trf = ensure_docx(doc_path)

    print(doc_path)
    print("############", docx_trf)

    try:
        df_1a = None

        # Step 1: Try extracting from DOCX only if available
        if docx_trf:
            try:
                df_1a = extract_table1a(docx_trf)
                df_1a = format_critical_components_df(df_1a)
            except IndexError:
                df_1a = None  # Table not found → fallback

        # Step 2: Fallback if extraction failed OR dataframe empty
        if df_1a is None :
            print("⚠️ Table 1.A not found or empty. Loading from CDR components.")
            df_1a = load_cdr_components_df(src_files_dir)
            #print(df_1a)       
            # df_1a = format_critical_components_df(df_1a)
            # Step 3: Try formatting (NON-BLOCKING)
            try:
                df_1a = format_critical_components_df(df_1a)
            except Exception as e:
                print("⚠️ Formatting failed. Inserting raw table.")
                print(f"Reason: {e}")

        # Clean NaN before using
        df_1a = clean_dataframe_for_json(df_1a, use_null=True)

        print("####################### Critical Components Table ######################", df_1a)

        # Step 3: Insert table
        
        insert_dataframe_below_anchor(
            input_docx=output_letter_docx,
            output_docx=output_letter_docx,
            df=df_1a,
            anchor_text="Details for the following critical components or materials have not been provided as required:"
            
        )

        print("✅ Critical components table inserted successfully.")

    except Exception as e:
        print("⚠️ Failed to insert critical components table. Skipping this step.")
        print(f"Reason: {e}")



    # -------------------------------------------------------
    # STEP 6 — Extract IEC 61010 Non-Conformances
    # -------------------------------------------------------
    #print(doc_path)
    #print("@@@@@@@@@@@@",docx_trf)
    text = read_docx_full(docx_trf)


    print("############  DOCUMENT READ    #######################################")

    #print(text)

    # df_9_nonpass = extract_iec61010_non_conformances_full_doc(
    #     document_text=text,
    #     deployment_name=CHAT_DEPLOY
    # )
    df_9_nonpass=run_iec61010_non_conformance_extraction(
        docx_trf,
        deployment_name=CHAT_DEPLOY,
        chunk_size=800,
        chunk_overlap=100,
        fuzzy_threshold=50
    )
    print('######'*5)
    #print(df_9_nonpass)
    # ✅ Clean NaN
    df_9_nonpass = clean_dataframe_for_json(df_9_nonpass, use_null=True)

    data = inject_dataframes_into_letter_json(
        data=data,
        df_critical=df_1a,
        df_nonpass=df_9_nonpass
    )

    insert_dataframe_below_anchor(
        input_docx=output_letter_docx,
        output_docx=output_letter_docx,
        df=df_9_nonpass,
        anchor_text="The shared documents were evaluated during the intrinsic safety analysis and constructional evaluation and following non-conformances were observed:"
    )

    # -------------------------------------------------------
    # STEP 7 — Image Compliance Pipeline (uses SAME blob_urls)
    # -------------------------------------------------------

    # non_conforming_images=orchestrate_iec61010_image_compliance_pipeline(
    #     blob_file_list=blob_urls,   # SAME LIST
    #     local_download_dir=src_files_dir,
    #     source_docx_name=output_letter_docx,
    #     output_docx_path=output_letter_docx,
    #     connection_string=AZURE_CONN_STRING,
    #     container_name=BLOB_CONTAINER_NAME
    # )
    # print("Conforming Images : #####################")

    # print(non_conforming_images)

    # from pathlib import Path


    # base_path = Path("Documents") / project_Id
    # final_letter_images_path = base_path / "letter_images"

    # blob_urls_non_conform = upload_images_to_blob(
    # non_conforming_images=non_conforming_images,
    # connection_string=AZURE_CONN_STRING,
    # container_name=BLOB_CONTAINER_NAME,
    # base_path=base_path,
    # final_letter_images_path=final_letter_images_path
    # )

    # # print("blob urls: ###################################")
    # # print(blob_urls_non_conform)

    # update_non_conforming_urls_from_blob(data, blob_urls_non_conform)

    try:
        # Run IEC61010 image compliance pipeline
        non_conforming_images = orchestrate_iec61010_image_compliance_pipeline(
            blob_file_list=blob_urls,   # SAME LIST
            local_download_dir=src_files_dir,
            source_docx_name=output_letter_docx,
            output_docx_path=output_letter_docx,
            connection_string=AZURE_CONN_STRING,
            container_name=BLOB_CONTAINER_NAME
        )

        # If pipeline returns empty or None
        if not non_conforming_images:
            raise ValueError("No non-conforming images found.")

        base_path = Path("Documents") / project_Id
        final_letter_images_path = base_path / "letter_images"

        # Upload images to blob
        blob_urls_non_conform = upload_images_to_blob(
            non_conforming_images=non_conforming_images,
            connection_string=AZURE_CONN_STRING,
            container_name=BLOB_CONTAINER_NAME,
            base_path=base_path,
            final_letter_images_path=final_letter_images_path
        )

        # Update URLs in final letter JSON
        update_non_conforming_urls_from_blob(data, blob_urls_non_conform)

        print("✔ Non-conforming images processed and uploaded successfully.")

    except Exception as e:
        print("⚠ Image compliance pipeline skipped.")
        print("Reason:", str(e))

        # Optional: ensure downstream code doesn't fail
        blob_urls_non_conform = []





    # -------------------------------------------------------
    # STEP 8 — Save Outputs
    # -------------------------------------------------------
    
    with open(letter_json_path_output, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

    


    with open(letter_header_json_path_output, "w", encoding="utf-8") as f:
        json.dump(data_header, f, indent=2, ensure_ascii=False)
    
 
    

    print("\n==============================")
    print("✅ LETTER GENERATION COMPLETE")
    print("==============================\n")

    return {
        "letter_json": "letter_output.json",
        "header_json": "letter_header_output.json",
        "letter_docx": output_letter_docx
    }

def letter_gen(blob_urls,    
               container_name, src_files_dir,src_files_trf, letter_json_path,
                letter_header_json_path,letter_template_docx,
                output_letter_docx,
                letter_json_path_output,
                letter_header_json_path_output,
                project_Id,blob_urls_trf,
                text_container,
                image_container):
        
        generate_letter_pipeline(
            blob_urls,
            container_name,
            src_files_dir,
            letter_json_path,
            letter_header_json_path,
            letter_template_docx,
            output_letter_docx,
            letter_json_path_output,
            letter_header_json_path_output,
            project_Id=project_Id,
            text_container=text_container,
            image_container=image_container)

        # extracted_data =  process_quote_from_folder(src_files_trf)
        # update_scope_of_work_in_json(letter_json_path_output,letter_json_path_output,extracted_data=extracted_data)
        # ------------------------------------------------
    # STEP 2 — Extract Scope of Work Data (SAFE)
    # ------------------------------------------------
        extracted_data = None
        try:
            extracted_data = process_quote_from_folder(src_files_trf)

            if not extracted_data:
                print("⚠️ No extracted quote data found. Skipping scope of work update.")
            else:
                update_scope_of_work_in_json(
                    letter_json_path_output,
                    letter_json_path_output,
                    extracted_data=extracted_data
                )
                print("✅ Scope of work updated successfully.")

        except Exception as e:
            print("⚠️ Failed to update scope of work. Continuing without it.")
            print(f"Reason: {e}")

        # updated_letter = add_text_support_with_blob_url_and_filename(
        #         letter_json_path=letter_json_path_output,
        #         folder_path=src_files_trf )

        # with open(letter_json_path_output, "w", encoding="utf-8") as f:
        #     json.dump(updated_letter, f, indent=2,ensure_ascii=False)

        try:
            updated_letter = add_text_support_with_blob_url_and_filename(
                letter_json_path=letter_json_path_output,
                folder_path=src_files_trf
            )

            if not updated_letter:
                print("⚠️ Updated letter content is empty. Skipping JSON overwrite.")
            else:
                with open(letter_json_path_output, "w", encoding="utf-8") as f:
                    json.dump(updated_letter, f, indent=2, ensure_ascii=False)

                print("✅ Letter JSON updated with blob support text.")

        except Exception as e:
            print("⚠️ Failed while adding blob support text.")
            print(f"Reason: {e}")

 

        with open(letter_json_path_output, "r", encoding="utf-8") as f:
            data_final = json.load(f)

        data_final=attach_blob_urls_to_image_support_letter(data_final, blob_urls_trf)
        data_final=attach_blob_urls_to_text_support_letter(data_final, blob_urls_trf)

        with open(letter_json_path_output, "w", encoding="utf-8") as f:
            json.dump(data_final, f, indent=2, ensure_ascii=False)



        with open(letter_header_json_path_output, "r", encoding="utf-8") as f:
            data_header_final = json.load(f)
        data_header_final=attach_blob_urls_to_image_support_letter(data_header_final, blob_urls_trf)
        data_header_final=attach_blob_urls_to_text_support_letter(data_header_final, blob_urls_trf)

        with open(letter_header_json_path_output, "w", encoding="utf-8") as f:
            json.dump(data_header_final, f, indent=2, ensure_ascii=False)
        

        replace_keys_with_values_no_format_change_all(
            input_docx=output_letter_docx,
            output_docx=output_letter_docx,
            data=data_final
        )

        delete_cosmos_container(
        endpoint=COSMOS_URL,
        key=COSMOS_KEY,
        database_name=DB_NAME,
        container_name=text_container
        )

        delete_cosmos_container(
        endpoint=COSMOS_URL,
        key=COSMOS_KEY,
        database_name=DB_NAME_IMG,
        container_name=image_container
        )
