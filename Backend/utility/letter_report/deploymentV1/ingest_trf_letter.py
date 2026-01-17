# ingestion_tool.py
# --------------------------------------------------------
# SECTION 2.1 — Imports, Configuration Loading, Constants
# --------------------------------------------------------
import re
import tempfile
import shutil
import requests
from urllib.parse import urlparse, unquote
from email import policy
from email.parser import BytesParser
import extract_msg
import uuid
import io
import openpyxl
import xlrd

from azure.storage.blob import BlobClient
from azure.core.exceptions import ResourceNotFoundError, AzureError
from utility.letter_report.deploymentV1.trf_essential import *
from utility.letter_report.deploymentV1.trf_utils import *
import pandas as pd
import math
import copy
import time
from azure.cosmos import CosmosClient, PartitionKey, exceptions
import json, os
from azure.cosmos import CosmosClient, ConsistencyLevel
from typing import List, Dict, Any, Tuple
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
from langchain_azure_ai.vectorstores import AzureCosmosDBNoSqlVectorSearch
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from azure.cosmos import CosmosClient
from langchain_openai import AzureOpenAIEmbeddings
from operator import itemgetter
from langchain_core.runnables import (
    RunnableParallel, RunnableLambda, RunnableMap
)
from langchain_core.output_parsers import StrOutputParser
from langchain_openai import AzureChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from tenacity import retry, retry_if_exception_type, wait_exponential, stop_never, RetryCallState
from openai import RateLimitError  # Make sure this import exists
from types import SimpleNamespace
from concurrent.futures import ThreadPoolExecutor, as_completed
from azure.core.exceptions import HttpResponseError
import time
from langchain_community.callbacks import get_openai_callback
import os
import re
import json
import time
import shutil
import tempfile
from types import SimpleNamespace
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlparse, unquote

# External dependencies
from azure.cosmos import CosmosClient, PartitionKey, exceptions
from azure.core.exceptions import AzureError, ResourceNotFoundError, HttpResponseError
from azure.storage.blob import BlobClient
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document

from openai import AzureOpenAI

from dotenv import load_dotenv
load_dotenv()

AZURE_CONN_STRING = os.getenv("LT_AZURE_CONN_STRING")
DB_NAME_IMG = os.getenv("LT_DB_NAME_IMG")
CONT_NAME_IMG = os.getenv("LT_CONT_NAME_IMG")
CHUNK_SIZE = int(os.getenv("LT_CHUNK_SIZE"))
CHUNK_OVERLAP = int(os.getenv("LT_CHUNK_OVERLAP"))
TOP_K = int(os.getenv("LT_TOP_K"))
EMBED_DIM = int(os.getenv("LT_EMBED_DIM"))
VECTOR_PATH = os.getenv("LT_VECTOR_PATH")

BLOB_CONTAINER_NAME = os.getenv("LT_BLOB_CONTAINER_NAME")
CONN_STR = os.getenv("LT_conn_str")

IMAGE_EXTS = os.getenv("LT_IMAGE_EXTS")

AOAI_ENDPOINT = os.getenv("LT_AOAI_ENDPOINT")
AOAI_KEY = os.getenv("LT_AOAI_KEY")
API_VERSION = os.getenv("LT_API_VERSION")
EMBED_DEPLOY = os.getenv("LT_EMBED_DEPLOY")
CHAT_DEPLOY = os.getenv("LT_CHAT_DEPLOY")
# COSMOS_URL = os.getenv("LT_COSMOS_URL")
# COSMOS_KEY = os.getenv("LT_COSMOS_KEY")
COSMOS_DB = os.getenv("LT_COSMOS_DB")
COSMOS_CONT = os.getenv("LT_COSMOS_CONT")
DB_NAME = os.getenv("LT_DB_NAME")
CONT_NAME = os.getenv("LT_CONT_NAME")
MAX_THREADS = int(os.getenv("LT_MAX_THREADS"))
MAX_RETRIES = int(os.getenv("LT_MAX_RETRIES"))
INITIAL_BACKOFF = int(os.getenv("LT_INITIAL_BACKOFF"))
# COSMOS_DB_TEXT  = os.getenv("COSMOS_DB_TEXT")
COSMOS_CONT_TEXT = COSMOS_CONT
COSMOS_DB_IMAGE  = DB_NAME_IMG
COSMOS_CONT_IMAGE = CONT_NAME_IMG
BLOB_CONT_NAME= os.getenv("BLOB_CONT_NAME")
INITIAL_BACKOFF = int(os.getenv("LT_INITIAL_BACKOFF"))
ENABLE_CAD_SCHEMATICS = os.getenv("ENABLE_CAD_SCHEMATICS")
DOWNLOAD_DIR = os.getenv("DOWNLOAD_DIR")
FLATTENED_DIR = os.getenv("FLATTENED_DIR")
IMAGES_ROOT =os.getenv("IMAGES_ROOT")
TRF_DOWNLOAD_DIR=os.getenv("TRF_DOWNLOAD_DIR")

COSMOS_URL  = "https://rag-intertek.documents.azure.com:443/"
COSMOS_KEY  = "AbhkomWJLtf8TR7odpABPqx1OrjlmCcpTXlKr9Vvp3RulZmFGollxQflIp3LLUAFt4XcMh70RbRxACDbuxyZLg=="
COSMOS_DB_TEXT  = "ragdatabase_new_itk"
COSMOS_CONT_TEXT = "vectorstorecontainer_new_itk"
# ----------------------------------------------------------------------------------------
# Azure OpenAI Client (shared for whole pipeline)
# ----------------------------------------------------------------------------------------
aoai_client = AzureOpenAI(
    api_key=AOAI_KEY,
    api_version=API_VERSION,
    azure_endpoint=AOAI_ENDPOINT,
)


# ----------------------------------------------------------------------------------------
# Embedding Builder — same as notebook
# ----------------------------------------------------------------------------------------
from langchain_openai import AzureOpenAIEmbeddings

def build_embeddings():
    return AzureOpenAIEmbeddings(
        azure_endpoint=AOAI_ENDPOINT,
        api_key=AOAI_KEY,
        openai_api_version=API_VERSION,
        azure_deployment=EMBED_DEPLOY,
    )


# ----------------------------------------------------------------------------------------
# Vector Store Builder (for TEXT)
# EXACT logic from notebook (not altered)
# ----------------------------------------------------------------------------------------


def build_vectorstore_text():
    cosmos_client = CosmosClient(
        url=COSMOS_URL,
        credential=COSMOS_KEY
    )

    return AzureCosmosDBNoSqlVectorSearch(
        cosmos_client=cosmos_client,
        embedding=build_embeddings(),
        database_name=COSMOS_DB_TEXT,
        container_name=COSMOS_CONT_TEXT,

        vector_embedding_policy={
            "vectorEmbeddings": [{
                "path": "/vector",
                "dataType": "float32",
                "dimensions": EMBED_DIM,
                "distanceFunction": "cosine"
            }]
        },
        indexing_policy={
            "includedPaths": [{"path": "/*"}],
            "excludedPaths": [{"path": "/\"_etag\"/?"}, {"path": "/vector/*"}],
            "vectorIndexes": [{"path": "/vector", "type": "quantizedFlat"}],
        },
        cosmos_container_properties={"partition_key": "/id"},
        cosmos_database_properties={},
        vector_search_fields={
            "text_field": "text",
            "embedding_field": "vector",
            "metadata_field": "metadata"
        }
    )


# ----------------------------------------------------------------------------------------
# Vector Store Builder (for IMAGES)
# EXACT logic from notebook (not altered)
# ----------------------------------------------------------------------------------------
def build_vectorstore_image():
    cosmos_client = CosmosClient(
        url=COSMOS_URL,
        credential=COSMOS_KEY
    )

    return AzureCosmosDBNoSqlVectorSearch(
        cosmos_client=cosmos_client,
        embedding=build_embeddings(),
        database_name=COSMOS_DB_IMAGE,
        container_name=COSMOS_CONT_IMAGE,

        vector_embedding_policy={
            "vectorEmbeddings": [{
                "path": "/vector",
                "dataType": "float32",
                "dimensions": EMBED_DIM,
                "distanceFunction": "cosine"
            }]
        },
        indexing_policy={
            "includedPaths": [{"path": "/*"}],
            "excludedPaths": [{"path": "/\"_etag\"/?"}, {"path": "/vector/*"}],
            "vectorIndexes": [{"path": "/vector", "type": "quantizedFlat"}],
        },
        cosmos_container_properties={"partition_key": "/id"},
        cosmos_database_properties={},
        vector_search_fields={
            "text_field": "text",
            "embedding_field": "vector",
            "metadata_field": "metadata"
        }
    )


# ingestion_tool.py (continued)
# --------------------------------------------------------
# SECTION 2.2 — PDF Text Loading, CAD/Schematic Image Extraction,
#                Chunking Logic, Image Upload to Azure Blob Storage
# --------------------------------------------------------

import fitz  # PyMuPDF


# -------------------------------------------------------------------------
# Utility: Sanitize blob names (same as in notebook)
# -------------------------------------------------------------------------
def sanitize_blob_name(name: str) -> str:
    name = name.replace(" ", "_")
    name = re.sub(r"[^A-Za-z0-9_\-./]", "_", name)
    return name


# -------------------------------------------------------------------------
# STRICT CAD/Schematic Page Detector (EXACT logic from notebook)
# -------------------------------------------------------------------------
def extract_relevant_pdf_page_images(pdf_path, dpi=200):
    """
    STRICT selection of pages containing diagrams / CAD / schematics.
    EXACT code copied from your notebook. No modifications.
    """

    import fitz
    pdf = fitz.open(pdf_path)
    base = os.path.basename(pdf_path)

    image_page_metadata = []

    for i, page in enumerate(pdf):
        page_num = i + 1

        text = page.get_text().strip()
        images = page.get_images(full=True)
        drawings = page.get_drawings()
        blocks = page.get_text("blocks")

        text_len = len(text)
        num_blocks = len(blocks)
        vector_ops = len(drawings)

        raster_area = 0
        for img in images:
            try:
                w = img[2]
                h = img[3]
                raster_area += (w * h)
            except:
                continue

        # --- EXACT strict rules from notebook ---
        should_extract = False

        if raster_area > 500000:
            should_extract = True
        elif vector_ops > 150:
            should_extract = True
        elif text_len < 30:
            should_extract = True
        elif num_blocks > 30 and text_len < 150:
            should_extract = True
        elif (
            any(k in text.lower() for k in ["label", "regulatory"])
            and vector_ops > 20
            and text_len < 600
        ):
            should_extract = True
        elif any(k in base.lower() for k in ["schematic", "cad", "drawing", "layout", "wiring"]) and text_len < 80:
            should_extract = True
        elif len(pdf) <= 3 and (vector_ops > 20 or text_len < 600):
            should_extract = True

        if not should_extract:
            continue

        # ---- Extract the page image EXACTLY like notebook ----
        try:
            pix = page.get_pixmap(dpi=dpi)
            img_path = f"{pdf_path}_page_{page_num}.png"
            pix.save(img_path)

            image_page_metadata.append({
                "pdf_file": base,
                "page": page_num,
                "local_image_path": img_path,
                "text_length": text_len,
                "raster_area": raster_area,
                "vector_ops": vector_ops,
                "blocks": num_blocks,
            })

        except Exception as e:
            print(f"[WARN] Image extraction failed for {base} page {page_num}: {e}")

    return image_page_metadata



 


# -------------------------------------------------------------------------
# PDF Loader + Text Chunking (EXACT logic from notebook)
# -------------------------------------------------------------------------
def load_and_split_pdfs_text(
    pdf_paths,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    extracted_texts=None,
    cad_schematics=True
):
    """
    EXACT implementation from your notebook.
    Returns:
        chunks → list of text Document objects
        image_page_metadata → list of schematic image metadata
    """

    docs = []
    image_page_metadata = []

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " "],
        keep_separator=False,
    )

    # ----- STEP 1: PDF TEXT EXTRACTION -----
    for path in pdf_paths:
        if not str(path).lower().endswith(".pdf"):
            continue

        loader = PyPDFLoader(str(path))
        raw_docs = loader.load()
        base = os.path.basename(str(path))

        for d in raw_docs:
            page = int(d.metadata.get("page", 1))
            d.metadata["source_file"] = base
            d.metadata["page"] = page
            d.metadata["citation"] = f"{base}#page={page}"

        docs.extend(raw_docs)

        # ----- STEP 2: CAD/Schematic Image Extraction -----
        if cad_schematics:
            try:
                extracted = extract_relevant_pdf_page_images(path)
                image_page_metadata.extend(extracted)
            except Exception as e:
                print(f"[WARN] selective image extraction failed for {path}: {e}")

    # ----- STEP 3: External extracted text -----
    # -------------------------------------------------------------
    # STEP 2 — Process externally extracted text files (unchanged)
    # -------------------------------------------------------------
    if extracted_texts:
        for item in extracted_texts:
            if not isinstance(item, dict):
                continue

            if "filename" in item and "text" in item:
                filename = item["filename"]
                text = item["text"]
            elif len(item) == 1:
                filename, text = next(iter(item.items()))
            elif "text" in item:
                filename = item.get("filename") or "unknown"
                text = item["text"]
            else:
                filename = item.get("filename") or "unknown"
                text = None
                for k, v in item.items():
                    if isinstance(v, str) and v.strip():
                        filename = k
                        text = v
                        break
                if text is None:
                    text = " ".join(str(v) for v in item.values())

            metadata = {
                "source_file": os.path.basename(str(filename)),
                "page": 1,
                "citation": os.path.basename(str(filename))
            }

            # ORIGINAL WORKING VERSION — KEEP SimpleNamespace
            docs.append(
                SimpleNamespace(
                    page_content=text or "",
                    metadata=metadata
                )
            )



    # ----- STEP 4: Chunking -----
    chunks = splitter.split_documents(docs)

    return chunks, image_page_metadata


# -------------------------------------------------------------------------
# Upload extracted PDF page images → Azure Blob Storage
# (EXACT logic from notebook)
# -------------------------------------------------------------------------


from concurrent.futures import ThreadPoolExecutor, as_completed
from azure.storage.blob import BlobClient
import os

def upload_pdf_images_and_append_urls(
    image_page_metadata,
    image_urls,
    conn_str,
    container,
    max_workers=8
):
    """
    Takes CAD/schematic page images extracted above,
    uploads each PNG to blob storage in parallel,
    appends URLs in same format as notebook.
    """

    def upload_single(item):
        local_path = item["local_image_path"]
        pdf_file = item["pdf_file"]
        page = item.get("page") or item.get("page_num")

        safe_pdf_name = sanitize_blob_name(pdf_file)
        safe_image_filename = sanitize_blob_name(os.path.basename(local_path))

        blob_name = f"{safe_pdf_name}/page_{page}.png"

        try:
            blob = BlobClient.from_connection_string(
                conn_str,
                container_name=container,
                blob_name=blob_name,
            )

            with open(local_path, "rb") as f:
                blob.upload_blob(f, overwrite=True)

            blob_url = blob.url

            return {
                "url": blob_url,
                "image_file": safe_image_filename,
                "pdf_file": pdf_file,
                "page": page
            }

        except Exception as e:
            print(f"[ERROR] Upload failed for {local_path}: {e}")
            return None

    # -------------------------------------------------
    # PARALLEL EXECUTION
    # -------------------------------------------------
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = [
            executor.submit(upload_single, item)
            for item in image_page_metadata
        ]

        for future in as_completed(futures):
            result = future.result()
            if result:
                image_urls.append(result)

    return image_urls


# ingestion_tool.py (continued)
# --------------------------------------------------------
# SECTION 3.1 — Image URL Extraction + Agent-Based OCR Pipeline
# --------------------------------------------------------

import requests


# --------------------------------------------------------
# Extract image name from blob URL — EXACT as notebook
# --------------------------------------------------------
def extract_clean_image_name(blob_url: str):
    parsed = urlparse(blob_url)
    path = parsed.path
    parts = path.split("/")

    pdf_file = None
    for part in parts:
        if part.lower().endswith(".pdf"):
            pdf_file = part
            break

    image_filename = os.path.basename(path)

    if pdf_file:
        return f"{pdf_file}/{image_filename}"
    else:
        return image_filename


# --------------------------------------------------------
# Wrapper: turn mixed list into plain list of URLs
# --------------------------------------------------------
def extract_urls(mixed_list):
    urls = []

    for item in mixed_list:
        if isinstance(item, str):
            urls.append(item)
        elif isinstance(item, dict) and "url" in item:
            urls.append(item["url"])

    return urls


# --------------------------------------------------------
# The ORIGINAL process_single_image replaced by an AGENT CALL
# --------------------------------------------------------
MAX_THREADS = 10
MAX_RETRIES = 5
INITIAL_BACKOFF = 3


def _direct_llm_fallback(url, index, total, vision_deploy_name):
    """
    Only used if agent not triggered correctly.
    Matches original notebook LLM behavior exactly.
    """

    resp = requests.get(url, timeout=20)
    resp.raise_for_status()

    completion = aoai_client.chat.completions.create(
        model=vision_deploy_name,
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": url},
                    {
                        "type": "text",
                        "text": (
                            "Perform OCR on this image and also provide a detailed "
                            "description. Combine both into one response."
                        ),
                    },
                ],
            }
        ],
        max_tokens=2048,
    )

    extracted_text = completion.choices[0].message.content.strip()
    image_name = extract_clean_image_name(url)

    return Document(
        page_content=extracted_text,
        metadata={
            "image_name": image_name,
            "blob_url": url,
            "source_type": "image",
        },
    )


# --------------------------------------------------------
# Agent Function Schema (EXACT as notebook)
# --------------------------------------------------------
image_desc_schema = {
    "name": "image_description_agent",
    "description": "OCR + description on an Azure Blob image URL.",
    "parameters": {
        "type": "object",
        "properties": {
            "url": {"type": "string"},
            "vision_deploy_name": {"type": "string"},
            "index": {"type": "integer"},
            "total": {"type": "integer"},
        },
        "required": ["url", "vision_deploy_name"],
    }
}


# --------------------------------------------------------
# run_function_agent_llm — EXACT as notebook
# --------------------------------------------------------
def run_function_agent_llm(client, user_prompt: str, function_schema: dict, python_callback: callable):
    response = client.chat.completions.create(
        model=CHAT_DEPLOY,
        messages=[{"role": "user", "content": user_prompt}],
        functions=[function_schema],
    )

    msg = response.choices[0].message

    if msg.function_call:
        fn_args = json.loads(msg.function_call.arguments)
        return python_callback(**fn_args)

    return msg.content


# --------------------------------------------------------
# Callback executed when LLM triggers function_call
# --------------------------------------------------------
def image_desc_callback(url, vision_deploy_name, index=1, total=1):
    """
    Internally calls _direct_llm_fallback so that return format matches EXACT original logic.
    """
    try:
        return _direct_llm_fallback(url, index, total, vision_deploy_name)
    except Exception as e:
        print(f"[ERROR] Fallback LLM failed for {url}: {e}")
        return None


# --------------------------------------------------------
# Main Agent Wrapper — same structure as notebook
# --------------------------------------------------------
def image_desc_agent(blob_url, vision_deploy_name="gpt-4.1", index=1, total=1):

    payload = {
        "url": blob_url,
        "vision_deploy_name": vision_deploy_name,
        "index": index,
        "total": total,
    }

    return run_function_agent_llm(
        aoai_client,
        user_prompt="Use the image_description_agent tool on: " + json.dumps(payload),
        function_schema=image_desc_schema,
        python_callback=image_desc_callback,
    )


# --------------------------------------------------------
# NEW process_single_image → replaced with AGENT call
# --------------------------------------------------------
def process_single_image(url, index, total, vision_deploy_name):
    """
    Calls the agent and ALWAYS returns a Document object.
    This guarantees downstream ingestion will not break.
    """

    print(f"[INFO] Processing image {index}/{total} → {url}")
    backoff = INITIAL_BACKOFF

    for attempt in range(1, MAX_RETRIES + 1):

        try:
            # Validate URL
            resp = requests.get(url, timeout=20)
            resp.raise_for_status()

            # Call the agent
            result = image_desc_agent(
                blob_url=url,
                vision_deploy_name=vision_deploy_name,
                index=index,
                total=total
            )

            # -------------------------------------------------------------------
            # NORMALIZE OUTPUT → ALWAYS A Document (required for add_ids_to_chunks)
            # -------------------------------------------------------------------
            cleaned_name = extract_clean_image_name(url)

            # CASE 1: Agent correctly returned Document
            if isinstance(result, Document):
                return result

            # CASE 2: Agent returned dict (tool responses sometimes do this)
            if isinstance(result, dict):
                text = result.get("text") or result.get("content") or str(result)
                return Document(
                    page_content=text,
                    metadata={"image_name": cleaned_name, "blob_url": url, "source_type": "image"}
                )

            # CASE 3: Agent returned string
            if isinstance(result, str):
                return Document(
                    page_content=result,
                    metadata={"image_name": cleaned_name, "blob_url": url, "source_type": "image"}
                )

            # CASE 4: Anything else (fallback)
            return Document(
                page_content=str(result),
                metadata={"image_name": cleaned_name, "blob_url": url, "source_type": "image"}
            )

        except Exception as e:
            print(f"[WARN] Attempt {attempt}/{MAX_RETRIES} failed for image {index} → {url}: {e}")

            if attempt == MAX_RETRIES:
                print(f"[ERROR] Giving up after {MAX_RETRIES} attempts → {url}")
                return Document(
                    page_content="OCR extraction failed.",
                    metadata={"image_name": extract_clean_image_name(url), "blob_url": url, "source_type": "image", "error": str(e)}
                )

            print(f"[INFO] Cooling down {backoff}s before retry…")
            time.sleep(backoff)
            backoff *= 2



# --------------------------------------------------------
# Multi-threaded loader (unchanged except for agent call inside)
# --------------------------------------------------------
def load_and_process_images(image_urls, vision_deploy_name=CHAT_DEPLOY):
    docs = []
    total = len(image_urls)

    print(f"[START] Processing {total} images with up to {MAX_THREADS} threads.\n")

    with ThreadPoolExecutor(max_workers=MAX_THREADS) as executor:

        futures = {
            executor.submit(process_single_image, url, idx + 1, total, vision_deploy_name): idx
            for idx, url in enumerate(image_urls)
        }

        for future in as_completed(futures):
            idx = futures[future]
            url = image_urls[idx]

            try:
                result = future.result()
                if result is not None:
                    docs.append(result)
                else:
                    print(f"[ERROR] Image {idx+1}/{total} → returned None")

            except Exception as e:
                print(f"[FATAL] Unexpected error for {url}: {e}")

    print(f"\n[COMPLETE] Finished processing {total} images.\n")
    return docs


# ingestion_tool.py (continued)
# --------------------------------------------------------
# SECTION 4 — Full Orchestration: run_full_ingestion()
# --------------------------------------------------------

def clear_cosmos_container(database_name, container_name):
    """
    Clears all items from a Cosmos DB container.
    EXACT logic from notebook.
    """
    client = CosmosClient(COSMOS_URL, credential=COSMOS_KEY)

    container = client.get_database_client(database_name).get_container_client(container_name)

    print(f"[INFO] Deleting all items in {database_name}.{container_name} ...")

    try:
        items = container.read_all_items()
        for item in items:
            try:
                container.delete_item(item=item, partition_key=item["id"])
            except Exception as e:
                print(f"[WARN] Failed deleting item {item.get('id')}: {e}")

        print("[SUCCESS] All documents deleted successfully!\n")

    except Exception as e:
        print(f"[ERROR] Could not list/delete items: {e}")

#!/usr/bin/env python
# coding: utf-8


from pathlib import Path
from typing import List, Dict, Any
from PyPDF2 import PdfReader
from PyPDF2.generic import IndirectObject
import fitz  # PyMuPDF
import base64
from dotenv import load_dotenv
import os
from pathlib import Path
import json
from openai import OpenAI
import os
import importlib
# importlib.reload(configs)
from openai import AzureOpenAI


# # 🔍 Optional sanity check (don’t print the key!)
# print("Endpoint:", AZURE_OPENAI_ENDPOINT)
# print("API version:", AZURE_OPENAI_API_VERSION)

# # ✅ Azure OpenAI client
# client = AzureOpenAI(
#     api_key=AZURE_OPENAI_API_KEY,
#     api_version=AZURE_OPENAI_API_VERSION,
#     azure_endpoint=AZURE_OPENAI_ENDPOINT,
# )

def _resolve_indirect(obj):
    if isinstance(obj, IndirectObject):
        return obj.get_object()
    return obj


def _has_freetext_annotations(reader: PdfReader, max_pages: int = 3) -> bool:
    """Return True if any of the first `max_pages` pages contain /FreeText annotations.

    Many "editable"-looking PDFs (Fill & Sign / typewriter tools) store user-entered
    text as /FreeText annotations under each page's /Annots array. These are NOT
    AcroForm/XFA fields, so they won't appear under /Root -> /AcroForm.
    """

    try:
        pages = reader.pages
        n = min(len(pages), max_pages)

        for i in range(n):
            page = pages[i]
            annots = page.get("/Annots")
            annots = _resolve_indirect(annots) or []

            for a in annots:
                aobj = _resolve_indirect(a)
                subtype = aobj.get("/Subtype")
                # subtype is usually a NameObject like '/FreeText'
                if subtype and str(subtype) == "/FreeText":
                    return True
        return False
    except Exception:
        return False
def is_editable_form_pdf(path: Path) -> bool:
    """
    Return True if the PDF appears to contain *editable* content:
      - AcroForm fields (/Root -> /AcroForm -> /Fields)
      - XFA forms (/Root -> /AcroForm -> /XFA)
      - FreeText annotations (/Annots -> /Subtype /FreeText)

    Note: FreeText annotations are NOT form fields; they are an annotation layer.
    Ensures the file handle is closed (important on Windows) by using a 'with' block.
    """
    try:
        # Open the file explicitly so it gets closed after we're done
        with open(path, "rb") as f:
            reader = PdfReader(f)

            root = _resolve_indirect(reader.trailer.get("/Root"))
            if not root:
                return False

            acroform = root.get("/AcroForm")
            # If there's no AcroForm, it still may be "editable" via FreeText annotations.
            if not acroform:
                return _has_freetext_annotations(reader)

            acroform = _resolve_indirect(acroform)

            fields = acroform.get("/Fields")
            if isinstance(fields, IndirectObject):
                fields = fields.get_object()

            # If we have any form fields, it's editable
            if fields and len(fields) > 0:
                return True

            # Check for XFA forms as well
            xfa = acroform.get("/XFA")
            if isinstance(xfa, IndirectObject):
                xfa = xfa.get_object()

            if xfa:
                return True

            # No fields/XFA found. Fall back to FreeText annotations.
            return _has_freetext_annotations(reader)
    except Exception:
        return False

def flatten_and_get_images(input_path: str, output_path: str, dpi: int = 200):
    """
    Flattens the PDF AND returns a list of page images (pixmaps).
    Uses context managers so files are always closed, even on errors.
    """
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    pixmaps = []

    # Ensure both documents are closed reliably
    with fitz.open(input_path) as src_doc, fitz.open() as out_doc:
        for page in src_doc:
            pix = page.get_pixmap(matrix=mat)
            pixmaps.append(pix)

            new_page = out_doc.new_page(width=pix.width, height=pix.height)
            new_page.insert_image(new_page.rect, pixmap=pix)

        out_doc.save(output_path)

    return pixmaps

# ============================
# 3) Convert pixmaps to PNG paths
# ============================

def save_pixmaps_to_images(pixmaps: List[fitz.Pixmap], out_dir: Path, stem: str):
    out_dir.mkdir(parents=True, exist_ok=True)
    image_paths = []

    for i, pix in enumerate(pixmaps, start=1):
        img_path = out_dir / f"{stem}_page{i}.png"
        pix.save(str(img_path))
        image_paths.append(img_path)

    return image_paths


prompt_lm = """
You are reading a scanned "Client Information Sheet".

Extract ALL fields into a STRICT JSON object with this exact shape:

{
  "Applicant": {
    "Legal Entity Name": string or null,
    "DBA": string or null,
    "Street Address": string or null,
    "City, State, Postal Code, Country": string or null,
    "Phone Number": string or null,
    "Email": string or null,
    "Contacts": [
      {
        "Name": string or null,
        "Role": string or null,
        "Phone Number": string or null,
        "Email": string or null
      }
    ]
  },
  "Bill-To": {
    "Legal Entity Name": string or null,
    "Street Address": string or null,
    "City, State, Postal Code, Country": string or null,
    "Accounts Payable Contact": string or null,
    "Phone Number": string or null,
    "Email": string or null
  },
  "Manufacturer": {
    "Legal Entity Name": string or null,
    "Street Address": string or null,
    "City, State, Postal Code, Country": string or null,
    "Contacts": [
      {
        "Name": string or null,
        "Role": string or null,
        "Phone Number": string or null,
        "Email": string or null
      }
    ],
    "Estimated Production Date": string or null,
    "Labeling Method": string or null
  },
  "Completed By": string or null,
  "Dates": {
    "Form Completion": string or null
  },
  "Signatures": string or null
}

Rules:
- DO NOT put fields like "Legal Entity Name" at the root.
- All applicant fields must be inside "Applicant".
- All bill-to fields must be inside "Bill-To".
- All manufacturer/production facility fields must be inside "Manufacturer".
- Use null if something is missing or unreadable.
- Return ONLY JSON, no extra text.
"""

def image_to_data_uri(path: Path) -> str:
    b = path.read_bytes()
    b64 = base64.b64encode(b).decode("utf-8")
    return f"data:image/png;base64,{b64}"

def extract_page_with_llm(img_path: Path) -> str:
    """
    Sends a single PNG page to GPT-4.1 (vision-enabled)
    and extracts structured JSON from the form.
    """
    data_uri = image_to_data_uri(img_path)

    prompt = prompt_lm

    response = aoai_client.chat.completions.create(
        model="gpt-4.1",
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": data_uri}},
                ],
            }
        ]
    )

    return response.choices[0].message.content

from pathlib import Path
import shutil
import time
from typing import Dict, Any

# def process_pdfs(
#     src_root: str = "src_files",
#     flattened_dir: str = "flattened_pdfs",
#     images_root: str = "page_images",
#     dpi: int = 200,
#     archive_root: str | None = None,
# ) -> Dict[str, Any]:
#     src_root = Path(src_root)
#     flattened_dir = Path(flattened_dir)
#     images_root = Path(images_root)

#     flattened_dir.mkdir(exist_ok=True)
#     images_root.mkdir(exist_ok=True)

#     # ✅ NEW: archive folder NEXT TO src_root (same parent)
#     if archive_root is None:
#         archive_root = src_root.parent / "archived_editable"
#     else:
#         archive_root = Path(archive_root)

#     archive_root.mkdir(parents=True, exist_ok=True)

#     results: Dict[str, Any] = {}

#     print(f"Scanning {src_root.resolve()} ...")

#     for pdf_path in src_root.glob("*.pdf"):
#         if not pdf_path.is_file():
#             continue

#         print(f"\nChecking {pdf_path.name}...")

#         if not is_editable_form_pdf(pdf_path):
#             print(" → Not editable. Skipping.")
#             continue

#         print(" → Editable PDF detected.")

#         # 1) Flatten PDF
#         flat_pdf_path = flattened_dir / f"{pdf_path.stem}_flat.pdf"
#         pixmaps = flatten_and_get_images(str(pdf_path), str(flat_pdf_path), dpi=dpi)
#         print(" ✔ Flattened PDF created.")

#         # 2) Archive original (now outside src_root)
#         archive_path = archive_root / pdf_path.name
#         print(f"   → Archiving original editable to {archive_path}")

#         for attempt in range(3):
#             try:
#                 shutil.move(str(pdf_path), str(archive_path))
#                 break
#             except PermissionError as e:
#                 print(f"   ⚠️ PermissionError on move (attempt {attempt+1}/3): {e}")
#                 if attempt == 2:
#                     print("   ❌ Skipping archive for this file. It may be open in another program.")
#                 else:
#                     time.sleep(0.5)

#         # 3) Save images + LLM extraction as you already do...
#         img_dir = images_root / pdf_path.stem
#         img_paths = save_pixmaps_to_images(pixmaps, img_dir, pdf_path.stem)

#         llm_outputs = []
#         for img_path in img_paths:
#             print(f"   → Sending {img_path.name} to GPT-4.1...")
#             output = extract_page_with_llm(img_path)
#             llm_outputs.append(output)

#         results[pdf_path.name] = {
#             "original": archive_path,
#             "flattened": flat_pdf_path,
#             "images": img_paths,
#             "extracted": llm_outputs,
#         }

#     return results

def process_pdfs(
    src_root: str = "src_files_trf",
    images_root: str = "page_images",
    dpi: int = 200,
) -> Dict[str, Any]:

    src_root = Path(src_root)
    images_root = Path(images_root)

    images_root.mkdir(exist_ok=True)

    results: Dict[str, Any] = {}

    print(f"Scanning {src_root.resolve()} ...")

    for pdf_path in src_root.glob("*.pdf"):
        if not pdf_path.is_file():
            continue

        print(f"\nChecking {pdf_path.name}...")

        if not is_editable_form_pdf(pdf_path):
            print(" → Not editable. Skipping.")
            continue

        print(" → Editable PDF detected.")

        # Extract images directly from original PDF
        pixmaps = []

        with fitz.open(pdf_path) as doc:
            zoom = dpi / 72.0
            mat = fitz.Matrix(zoom, zoom)

            for page in doc:
                pix = page.get_pixmap(matrix=mat)
                pixmaps.append(pix)

        # Save images
        img_dir = images_root / pdf_path.stem
        img_paths = save_pixmaps_to_images(pixmaps, img_dir, pdf_path.stem)

        # OCR / LLM extraction
        llm_outputs = []
        for img_path in img_paths:
            print(f"   → Sending {img_path.name} to GPT-4.1...")
            output = extract_page_with_llm(img_path)
            llm_outputs.append(output)

        # results[pdf_path.name] = {
        #     "original": pdf_path,
        #     "images": img_paths,
        #     "extracted": llm_outputs,
        # }
        results[pdf_path.name] = {
            "original": pdf_path,
            "images": img_paths,
            "extracted": llm_outputs,
            "is_editable": True
        }


    return results


def extract_cis(src_root, images_root):
    results = process_pdfs(
        src_root=src_root,
        images_root=images_root,
        dpi=200
    )

    all_cis = []
    editable_pdfs = []

    for pdfs in results.keys():
        dic_cis = {}
        dic_cis['filename'] = pdfs
        dic_cis['text'] = results[pdfs]['extracted'][0]
        all_cis.append(dic_cis)

        editable_pdfs.append(pdfs)

    return all_cis, editable_pdfs

# with open("src_files\\all_cis_info.txt", "w", encoding="utf-8") as f:
#     json.dump(all_cis, f, indent=4, default=str)

def copy_extracted_images_to_src(page_images_root: str, src_root: str):
    """
    Copies all extracted page images into src_files directory
    so they get ingested like normal uploaded images.
    """

    page_images_root = Path(page_images_root)
    src_root = Path(src_root)

    src_root.mkdir(exist_ok=True)

    copied = 0

    for subdir in page_images_root.iterdir():
        if not subdir.is_dir():
            continue

        for img in subdir.glob("*.png"):
            dest = src_root / img.name

            # avoid overwrite
            if dest.exists():
                continue

            shutil.copy2(img, dest)
            copied += 1

    print(f"✅ Copied {copied} extracted page images into {src_root}")

def append_cis_images_to_image_metadata(images_root: str, image_page_metadata: list):
    """
    Adds CIS extracted page images into CAD schematic image pipeline
    so they get uploaded + OCR + embedded exactly the same way.
    """

    images_root = Path(images_root)

    for subdir in images_root.iterdir():
        if not subdir.is_dir():
            continue

        pdf_name = subdir.name + ".pdf"   # fake source name for metadata consistency

        for img in subdir.glob("*.png"):
            image_page_metadata.append({
                "pdf_file": pdf_name,
                "page": None,
                "local_image_path": str(img),
                "reason": "editable_pdf_page"
            })

 


# ------------------------------------------------------------
# Master Orchestration Function
# ------------------------------------------------------------
def run_full_ingestion(blob_urls):
    """
    Master ingestion function to be called externally.

    Steps:
    1. Process blob URLs → download files, convert doc/docx → pdf, extract other texts.
    2. Extract PDF text + CAD/Schematic page images.
    3. Create text vector store container + ingest text.
    4. Upload schematic page images to Blob Storage.
    5. Process all images through AGENT-based OCR.
    6. Create image vector store container + ingest image embeddings.
    """

    print("\n==============================")
    print("   STEP 1 — PROCESS BLOB URLs ")
    print("==============================\n")

    extracted_texts, image_urls_raw, downloaded_pdf_paths, converted_pdf_paths = process_blob_urls_2(
        blob_urls,
        AZURE_CONN_STRING,
        container=BLOB_CONT_NAME,   # SAME as notebook
        download_dir=DOWNLOAD_DIR, #DOWNLOAD_DIR from environment
        keep_files=True,
        verbose=True
    )

    # pdf_paths = downloaded_pdf_paths + converted_pdf_paths
    pdf_paths = downloaded_pdf_paths + converted_pdf_paths
    


    # cis_info = extract_cis(src_root=DOWNLOAD_DIR,
    #     flattened_dir=FLATTENED_DIR,
    #     images_root=IMAGES_ROOT)
    cis_info, editable_pdfs = extract_cis(
        src_root=DOWNLOAD_DIR,
        images_root=IMAGES_ROOT
    )

    # ✅ Remove editable PDFs from PDF ingestion list (they are handled via OCR images)
    pdf_paths = [
        p for p in pdf_paths
        if os.path.basename(p) not in editable_pdfs
    ]

    
    copy_extracted_images_to_src(
    page_images_root=IMAGES_ROOT,
    src_root=DOWNLOAD_DIR)
    
    extracted_texts += cis_info

    print("###################### CIS INFO ##################################")
    print(cis_info)
    print("###################################################################")



    print(f"[INFO] Extracted text files: {len(extracted_texts)}")
    print(f"[INFO] Initial image URLs: {len(image_urls_raw)}")
    print(f"[INFO] Total PDFs: {len(pdf_paths)}\n")


    print("\n======================================")
    print("   STEP 2 — LOAD + SPLIT PDF TEXT      ")
    print("======================================\n")

    chunks, image_page_metadata = load_and_split_pdfs_text(
        pdf_paths,
        CHUNK_SIZE,
        CHUNK_OVERLAP,
        extracted_texts=extracted_texts,
        cad_schematics=ENABLE_CAD_SCHEMATICS,
    )

    # ✅ Add CIS extracted images into CAD schematic pipeline
    append_cis_images_to_image_metadata(
        images_root=IMAGES_ROOT,
        image_page_metadata=image_page_metadata
    )


    print(f"[INFO] Total text chunks produced: {len(chunks)}")
    print(f"[INFO] CAD/Schematic pages extracted: {len(image_page_metadata)}\n")


    print("\n======================================")
    print("   STEP 3 — CREATE TEXT VECTOR STORE   ")
    print("======================================\n")

    # Clear DB (EXACT notebook logic)
    clear_cosmos_container(COSMOS_DB_TEXT, COSMOS_CONT_TEXT)

    vectorstore_text = build_vectorstore_text()

    # Assign UUIDs to each chunk
    chunks_uuid = add_ids_to_chunks(chunks)

    # Ingest in parallel
    ingest_to_cosmos_parallel(vectorstore_text, chunks_uuid, batch_size=10, max_workers=10)

    print("\n[SUCCESS] Text ingestion completed.\n")


    print("\n=====================================================")
    print("   STEP 4 — UPLOAD CAD/SCHEMATIC PAGE IMAGES         ")
    print("=====================================================\n")

    # image_urls = upload_pdf_images_and_append_urls(
    #     image_page_metadata=image_page_metadata,
    #     image_urls=image_urls_raw,
    #     conn_str=AZURE_CONN_STRING,
    #     container=BLOB_CONT_NAME,   # same notebook container
    # )

    # Deduplicate local image paths
    seen = set()
    unique_metadata = []

    for item in image_page_metadata:
        key = item["local_image_path"]
        if key not in seen:
            seen.add(key)
            unique_metadata.append(item)

    image_urls = upload_pdf_images_and_append_urls(
        image_page_metadata=unique_metadata,
        image_urls=image_urls_raw,
        conn_str=AZURE_CONN_STRING,
        container=BLOB_CONT_NAME,
    )


    # Turn into flat list
    img_links = extract_urls(image_urls)

    print(f"[INFO] Total images for OCR (blob URLs): {len(img_links)}\n")


    print("\n==============================================")
    print("   STEP 5 — IMAGE OCR USING AGENT PIPELINE     ")
    print("==============================================\n")

    docs_image = load_and_process_images(img_links, vision_deploy_name=CHAT_DEPLOY)

    print(f"[SUCCESS] Finished OCR for {len(docs_image)} images.\n")


    print("\n==============================================")
    print("   STEP 6 — CREATE IMAGE VECTOR STORE          ")
    print("==============================================\n")

    # Clear DB (EXACT notebook logic)
    clear_cosmos_container(COSMOS_DB_IMAGE, COSMOS_CONT_IMAGE)

    vectorstore_image = build_vectorstore_image()

    # Assign UUIDs to image docs
    docs_image_uuid = add_ids_to_chunks(docs_image)

    # Ingest
    ingest_to_cosmos_parallel(vectorstore_image, docs_image_uuid, batch_size=10, max_workers=10)

    print("\n[SUCCESS] Image ingestion completed.\n")


    print("\n==============================================")
    print("         FULL INGESTION PIPELINE DONE         ")
    print("==============================================\n")

    return {
        "text_chunks_ingested": len(chunks_uuid),
        "images_ingested": len(docs_image_uuid),
        "pdf_pages_extracted": len(image_page_metadata),
    }


# if __name__ == "__main__":
#     # --------------------------------------
#     # Example: Pass your blob URLs here
#     # --------------------------------------
#     blob_urls = [  #source file's upload url
#         'https://saaffine.blob.core.windows.net/nasa-ebooks-pdfs-all/Accepted_-_Gener8_LLC_-_Qu-01390131-0.msg',
#         'https://saaffine.blob.core.windows.net/nasa-ebooks-pdfs-all/CFE-4LB011-E%20(1)%20(1).pdf',
#         'https://saaffine.blob.core.windows.net/nasa-ebooks-pdfs-all/CFE_block_diagram.png',
#         'https://saaffine.blob.core.windows.net/nasa-ebooks-pdfs-all/Client_Information_Sheet_-_FUS_CIS_1_.pdf',
#         'https://saaffine.blob.core.windows.net/nasa-ebooks-pdfs-all/PastedGraphic-1.png',
#         'https://saaffine.blob.core.windows.net/nasa-ebooks-pdfs-all/celFE_isol.docx',
#         'https://saaffine.blob.core.windows.net/nasa-ebooks-pdfs-all/RE__External_Re__Intertek_Order_Qu-01390131-0_processed_-_Gener8_LLC_Project_G105581614.eml',
#         'https://saaffine.blob.core.windows.net/nasa-ebooks-pdfs-all/Risk%20Assessment%20CFE_28nov.xlsx',
#         'https://saaffine.blob.core.windows.net/nasa-ebooks-pdfs-all/CellFE_Infinity_MTx_Operating_Manual-jsg-11-16-2023_Final_1_.docx',
#         'https://saaffine.blob.core.windows.net/nasa-ebooks-pdfs-all/CellFE_Infinity_MTx_Operating_Manual.docx',
#         'https://saaffine.blob.core.windows.net/nasa-ebooks-pdfs-all/Cell_Gener8_Agent_Agreement_2018_1_.pdf',
#         'https://saaffine.blob.core.windows.net/nasa-ebooks-pdfs-all/Gener_8_PO_P70909_INTERTEK_TESTING_SERVICES__EMC_Safety_Testings_Proj_13403_-.pdf',
#         'https://saaffine.blob.core.windows.net/nasa-ebooks-pdfs-all/Gener8_SAF_Electrical_Risk_Assessment_Form_2022-11-30.docx',
#         'https://saaffine.blob.core.windows.net/nasa-ebooks-pdfs-all/Qu-01390131-0.pdf'

#     ]

#     if not blob_urls:
#         print("[WARNING] No blob URLs provided. Add URLs inside the blob_urls list.")
#     else:
#         print("\n🚀 Starting ingestion pipeline...\n")
#         result = run_full_ingestion(blob_urls)
#         print("\n📌 Final ingestion summary:")
#         print(result)
