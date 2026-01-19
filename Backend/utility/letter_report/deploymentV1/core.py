import re
import tempfile
import shutil
import requests
from urllib.parse import urlparse, unquote
from email import policy
from email.parser import BytesParser
import extract_msg
import uuid
from langchain_core.documents import Document
import io
import openpyxl
import xlrd

from utility.letter_report.deploymentV1.config import AZURE_CONN_STRING, DB_NAME_IMG, CONT_NAME_IMG, CHUNK_SIZE, CHUNK_OVERLAP, TOP_K, EMBED_DIM, VECTOR_PATH, BLOB_CONTAINER_NAME, conn_str, IMAGE_EXTS, AOAI_ENDPOINT, AOAI_KEY, API_VERSION, EMBED_DEPLOY, CHAT_DEPLOY, COSMOS_URL, COSMOS_KEY, COSMOS_DB, COSMOS_CONT, DB_NAME, CONT_NAME, MAX_THREADS, MAX_RETRIES, INITIAL_BACKOFF

import os
import pdfplumber
from fuzzywuzzy import fuzz
# from utils import *
from azure.storage.blob import BlobClient
from azure.core.exceptions import ResourceNotFoundError, AzureError
# from templates import *
from utility.letter_report.deploymentV1.trf_essential import *
from utility.letter_report.deploymentV1.trf_utils import *
import pandas as pd
import math
import copy
import time
from azure.cosmos import CosmosClient, PartitionKey, exceptions
import json, os
from azure.cosmos import CosmosClient, ConsistencyLevel
from typing import List, Dict, Any, Tuple
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
from langchain_azure_ai.vectorstores import AzureCosmosDBNoSqlVectorSearch
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from azure.cosmos import CosmosClient
from langchain_openai import AzureOpenAIEmbeddings
from operator import itemgetter
from langchain_core.runnables import (
    RunnableParallel, RunnableLambda, RunnableMap
)
from langchain_core.output_parsers import StrOutputParser
from langchain_openai import AzureChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from tenacity import retry, retry_if_exception_type, wait_exponential, stop_never, RetryCallState
from openai import RateLimitError  # Make sure this import exists
from types import SimpleNamespace
from concurrent.futures import ThreadPoolExecutor, as_completed
from azure.core.exceptions import HttpResponseError
import time
from langchain_community.callbacks import get_openai_callback
pd.set_option('display.max_colwidth', None)  # Don't truncate cell text
pd.set_option('display.max_rows', None)      # Show all rows (optional)
pd.set_option('display.max_columns', None)

import os
import re
import pandas as pd
from typing import Optional, Tuple

def build_vectorstore(embeddings, COSMOS_URL, COSMOS_KEY,DB_NAME, CONT_NAME ):
    
    
    cosmos_client = CosmosClient(
        url=COSMOS_URL,
        credential=COSMOS_KEY,
        consistency_level=ConsistencyLevel.Eventual
    )

    # keep your existing policy helpers if your constructor requires them
    return AzureCosmosDBNoSqlVectorSearch(
        cosmos_client=cosmos_client,
        embedding=embeddings,
        database_name=DB_NAME,
        container_name=CONT_NAME,

        # if your version requires explicit policies, keep these as you already had:
        vector_embedding_policy={"vectorEmbeddings":[{"path":"/vector","dataType":"float32","dimensions":1536,"distanceFunction":"cosine"}]},
        indexing_policy={"includedPaths":[{"path":"/*"}],
                         "excludedPaths":[{"path":"/\"_etag\"/?"},{"path":"/vector/*"}],
                         "vectorIndexes":[{"path":"/vector","type":"quantizedFlat"}]},
        cosmos_container_properties={"partition_key":"/id"},
        cosmos_database_properties={}, # _db_props()

        # IMPORTANT: pass a dict, not a list
        vector_search_fields={
            "text_field": "text",
            "embedding_field": "vector",
            "metadata_field": "metadata"
        }
    )


def extract_relevant_pdf_page_images(pdf_path, dpi=200):
    """
    STRICT version:
    Extracts ONLY highly-likely diagram/CAD/schematic pages.
    """

    import fitz
    import os

    pdf = fitz.open(pdf_path)
    base = os.path.basename(pdf_path)

    image_page_metadata = []

    for i, page in enumerate(pdf):
        page_num = i + 1

        text = page.get_text().strip()
        images = page.get_images(full=True)
        drawings = page.get_drawings()
        blocks = page.get_text("blocks")

        text_len = len(text)
        num_blocks = len(blocks)
        vector_ops = len(drawings)

        # COMPUTE TOTAL RASTER AREA (sum of w*h for each image)

        raster_area = 0
        for img in images:
            try:
                w = img[2]   # width
                h = img[3]   # height
                raster_area += (w * h)
            except:
                continue

    

        # ---------------- STRICT RULES ----------------
        should_extract = False

        # 1. Large raster content (ignore tiny icons)
        if raster_area > 500000:   # around 220×220 pixels
            should_extract = True

        # 2. Heavy vector operations (CAD / big schematics)
        elif vector_ops > 150:
            should_extract = True

        # 3. Extremely low text → typical schematic/drawing pages
        elif text_len < 30:
            should_extract = True

        # 4. Complex diagrams/tables but only if low text
        elif num_blocks > 30 and text_len < 150:
            should_extract = True

        # 5. labels / artwork
        elif (
                any(k in text.lower() for k in ["label", "regulatory"])
                and vector_ops > 20
                and text_len < 600
            ):
            should_extract = True

        # 6. Filename hint PLUS low text
        elif any(k in base.lower() for k in ["schematic", "cad", "drawing", "layout", "wiring"]) and text_len < 80:
            should_extract = True

        #  7. Short PDFs (typical for less than 3 page doc)
        elif (
            len(pdf) <= 3
            and (vector_ops > 20 or text_len < 600)
        ):
            should_extract = True


        if not should_extract:
            continue

        # ---------------- Extract image ----------------
        try:
            pix = page.get_pixmap(dpi=dpi)
            img_path = f"{pdf_path}_page_{page_num}.png"
            pix.save(img_path)

            image_page_metadata.append({
                "pdf_file": base,
                "page": page_num,
                "local_image_path": img_path,
                "text_length": text_len,
                "raster_area": raster_area,
                "vector_ops": vector_ops,
                "blocks": num_blocks,
            })

        except Exception as e:
            print(f"[WARN] Image extraction failed for {base} page {page_num}: {e}")

    return image_page_metadata


# #### load_and_split_pdfs_text---updated: ## CAD and Schematic Support TOGGLE SWITCH
def load_and_split_pdfs_text(
        pdf_paths,
        CHUNK_SIZE,
        CHUNK_OVERLAP,
        extracted_texts=None,
        cad_schematics=True   # ⭐ NEW PARAM
    ):
    """
    pdf_paths: iterable of file paths
    extracted_texts: optional list

    Returns:
        chunks → text-only chunks
        image_page_metadata → selective CAD/schematic/table images 
                              (or empty list if cad_schematics=False)
    """

    docs = []
    image_page_metadata = []   # NEW
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " "],
        keep_separator=False,
    )

    # -------------------------------------------------------------
    # STEP 1 — Process PDFs (Text + optional CAD/Schematic images)
    # -------------------------------------------------------------
    for path in pdf_paths:
        if not str(path).lower().endswith(".pdf"):
            continue

        # ---- TEXT EXTRACTION ----
        loader = PyPDFLoader(str(path))
        raw_docs = loader.load()
        base = os.path.basename(str(path))

        for d in raw_docs:
            page = int(d.metadata.get("page", 1))
            d.metadata["source_file"] = base
            d.metadata["page"] = page
            d.metadata["citation"] = f"{base}#page={page}"

        docs.extend(raw_docs)

        # ---- OPTIONAL: CAD / SCHEMATIC IMAGE EXTRACTION ----
        if cad_schematics:   # ⭐ ONLY RUN IF ENABLED
            try:
                extracted_pages = extract_relevant_pdf_page_images(path)
                image_page_metadata.extend(extracted_pages)
            except Exception as e:
                print(f"[WARN] selective image extraction failed for {path}: {e}")
        else:
            # When disabled → return empty image list
            pass

    # -------------------------------------------------------------
    # STEP 2 — Process externally extracted text files (unchanged)
    # -------------------------------------------------------------
    if extracted_texts:
        for item in extracted_texts:
            if not isinstance(item, dict):
                continue

            if "filename" in item and "text" in item:
                filename = item["filename"]
                text = item["text"]
            elif len(item) == 1:
                filename, text = next(iter(item.items()))
            elif "text" in item:
                filename = item.get("filename") or "unknown"
                text = item["text"]
            else:
                filename, text = None, None
                for k, v in item.items():
                    if isinstance(v, str) and v.strip():
                        filename = k
                        text = v
                        break
                if text is None:
                    filename = item.get("filename") or "unknown"
                    text = " ".join(str(v) for v in item.values())

            metadata = {
                "source_file": os.path.basename(str(filename)),
                "page": 1,
                "citation": os.path.basename(str(filename))
            }

            docs.append(SimpleNamespace(page_content=text or "", metadata=metadata))

    # -------------------------------------------------------------
    # STEP 3 — Chunking
    # -------------------------------------------------------------
    chunks = splitter.split_documents(docs)

    # -------------------------------------------------------------
    # FINAL RETURN
    # -------------------------------------------------------------
    return chunks, image_page_metadata


## CAD and Schematic Support
from azure.storage.blob import BlobClient
import os
import re

def sanitize_blob_name(name: str) -> str:
    """Azure Blob safe-name converter (removes unsafe characters)."""
    name = name.replace(" ", "_")
    name = re.sub(r"[^A-Za-z0-9_\-./]", "_", name)
    return name


def upload_pdf_images_and_append_urls(
        image_page_metadata,
        image_urls,
        conn_str,
        container):
    """
    Uploads only relevant PDF page images to Azure Blob Storage and
    appends blob URLs to image_urls in the correct vision-RAG format.
    """

    for item in image_page_metadata:
        local_path = item["local_image_path"]
        pdf_file = item["pdf_file"]

        # Allow both "page" and "page_num" from upstream
        page = item.get("page") or item.get("page_num")

        # Sanitize PDF file name for folders
        safe_pdf_name = sanitize_blob_name(pdf_file)

        # Blob path: <pdf_file>/page_<n>.png
        image_filename = os.path.basename(local_path)
        safe_image_filename = sanitize_blob_name(image_filename)

        blob_name = f"{safe_pdf_name}/page_{page}.png"

        try:
            blob = BlobClient.from_connection_string(
                conn_str,
                container_name=container,
                blob_name=blob_name,
            )

            with open(local_path, "rb") as f:
                blob.upload_blob(f, overwrite=True)

            blob_url = blob.url

            # Append proper structure for RAG Vision pipeline
            image_urls.append({
                "url": blob_url,
                "image_file": safe_image_filename,
                "pdf_file": pdf_file,
                "page": page
            })

        except Exception as e:
            print(f"[ERROR] Upload failed for {local_path}: {e}")

    return image_urls


from concurrent.futures import ThreadPoolExecutor, as_completed
from azure.storage.blob import BlobClient
import os

def upload_pdf_images_and_append_urls_parallel(
        image_page_metadata,
        image_urls,
        conn_str,
        container,
        max_workers=8
):
    """
    Parallel upload of relevant PDF page images to Azure Blob Storage
    and append blob URLs to image_urls in correct vision-RAG format.
    """

    def upload_single(item):
        local_path = item["local_image_path"]
        pdf_file = item["pdf_file"]
        page = item.get("page") or item.get("page_num")

        safe_pdf_name = sanitize_blob_name(pdf_file)
        image_filename = os.path.basename(local_path)
        safe_image_filename = sanitize_blob_name(image_filename)

        blob_name = f"{safe_pdf_name}/page_{page}.png"

        try:
            blob = BlobClient.from_connection_string(
                conn_str,
                container_name=container,
                blob_name=blob_name,
            )

            with open(local_path, "rb") as f:
                blob.upload_blob(f, overwrite=True)

            return {
                "url": blob.url,
                "image_file": safe_image_filename,
                "pdf_file": pdf_file,
                "page": page
            }

        except Exception as e:
            print(f"[ERROR] Upload failed for {local_path}: {e}")
            return None

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = [executor.submit(upload_single, item) for item in image_page_metadata]

        for future in as_completed(futures):
            result = future.result()
            if result:
                image_urls.append(result)

    return image_urls



# for processing images (dont run)
def build_vectorstore2(embeddings, COSMOS_URL, COSMOS_KEY, DB_NAME_IMG, CONT_NAME_IMG):
    
    cosmos_client = CosmosClient(
        url=COSMOS_URL,
        credential=COSMOS_KEY,
        consistency_level=ConsistencyLevel.Eventual
    )

    # keep your existing policy helpers if your constructor requires them
    return AzureCosmosDBNoSqlVectorSearch(
        cosmos_client=cosmos_client,
        embedding=embeddings,
        database_name=DB_NAME_IMG,
        container_name=CONT_NAME_IMG,

        # if your version requires explicit policies, keep these as you already had:
        vector_embedding_policy={"vectorEmbeddings":[{"path":"/vector","dataType":"float32","dimensions":1536,"distanceFunction":"cosine"}]},
        indexing_policy={"includedPaths":[{"path":"/*"}],
                         "excludedPaths":[{"path":"/\"_etag\"/?"},{"path":"/vector/*"}],
                         "vectorIndexes":[{"path":"/vector","type":"quantizedFlat"}]},
        cosmos_container_properties={"partition_key":"/id"},
        cosmos_database_properties={}, # _db_props()

        # IMPORTANT: pass a dict, not a list
        vector_search_fields={
            "text_field": "text",
            "embedding_field": "vector",
            "metadata_field": "metadata"
        }
    )

# for image processing (dont run)
def extract_urls(mixed_list):
    urls = []

    for item in mixed_list:
        # If the item is a direct string → append
        if isinstance(item, str):
            urls.append(item)

        # If the item is a dict and has a 'url' key → append
        elif isinstance(item, dict) and "url" in item:
            urls.append(item["url"])

        # Ignore anything else
        else:
            continue

    return urls


from urllib.parse import urlparse
import os
import re

def extract_clean_image_name(blob_url: str):
    """
    Returns the FULL relative image name:
    <pdf_file>/page_4.png   OR   <pure_image.png>
    Removes SAS tokens.
    """

    parsed = urlparse(blob_url)
    path = parsed.path                     # /container/.../<whatever>
    parts = path.split("/")

    # Identify pdf folder if present
    pdf_file = None
    for part in parts:
        if part.lower().endswith(".pdf"):
            pdf_file = part
            break

    # Extract clean PNG filename without SAS
    image_filename = os.path.basename(path)   # page_4.png
    # Make FULL relative path
    if pdf_file:
        return f"{pdf_file}/{image_filename}"
    else:
        return image_filename



def process_single_image(url, index, MAX_RETRIES, INITIAL_BACKOFF, total, vision_deploy_name, client):
    """
    Single-thread processing with retry/backoff logic and progress prints.
    """
    print(f"[INFO] Processing image {index}/{total} → {url}")

    backoff = INITIAL_BACKOFF

    for attempt in range(1, MAX_RETRIES + 1):

        try:
            # Validate URL
            resp = requests.get(url, timeout=20)
            resp.raise_for_status()

            image_name = url.split("/")[-1]

            # Call Azure LLM Vision
            completion = client.chat.completions.create(
                model=vision_deploy_name,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "image_url", "image_url": url},
                            {
                                "type": "text",
                                "text": (
                                    "Perform OCR on this image and also provide a detailed "
                                    "description. Combine both into one response."
                                ),
                            },
                        ],
                    }
                ],
                max_tokens=2048,
            )

            extracted_text = completion.choices[0].message.content.strip()
            image_name = extract_clean_image_name(url)


            print(f"[SUCCESS] Completed image {index}/{total} → {url}")

            return Document(
                page_content=extracted_text,
                metadata={
                    "image_name": image_name,
                    "blob_url": url,
                    "source_type": "image",
                },
            )

        except Exception as e:
            print(f"[WARN] Attempt {attempt}/{MAX_RETRIES} failed for image {index} → {url}: {e}")

            if attempt == MAX_RETRIES:
                print(f"[ERROR] Giving up on image {index} → {url} after {MAX_RETRIES} attempts.")
                return None

            print(f"[INFO] Cooling down {backoff}s before retry (image {index})...")
            time.sleep(backoff)
            backoff *= 2  # exponential backoff




def load_and_process_images(image_urls,MAX_THREADS,MAX_RETRIES, INITIAL_BACKOFF, client, vision_deploy_name=CHAT_DEPLOY):
    docs = []
    total = len(image_urls)

    print(f"[START] Processing {total} images with up to {MAX_THREADS} threads.\n")

    with ThreadPoolExecutor(max_workers=MAX_THREADS) as executor:

        future_to_idx = {
            executor.submit(process_single_image, url, idx+1,MAX_RETRIES, INITIAL_BACKOFF, total, client, vision_deploy_name): idx
      
            for idx, url in enumerate(image_urls)
        }

        for future in as_completed(future_to_idx):
            idx = future_to_idx[future]
            url = image_urls[idx]

            try:
                result = future.result()
                if result is not None:
                    docs.append(result)
                else:
                    print(f"[ERROR] Image {idx+1}/{total} → returned None")
            except Exception as e:
                print(f"[FATAL] Unhandled error for image {idx+1}/{total} → {url}: {e}")

    print(f"\n[COMPLETE] Finished processing {total} images.\n")
    return docs


def build_vision_message_v5(inputs):
    """
    Simplified v5:
    - No task_type
    - No grey_mode
    - Single execution path
    - Evidence-driven extraction only
    """

    # ----------------------------------------------------
    # Extract fields
    # ----------------------------------------------------
    question = inputs.get("question", "")
    docs = inputs.get("context", [])
    image_urls = inputs.get("image", [])
    custom_prompt = inputs.get("custom_prompt")

    active_prompt = (
        custom_prompt.strip()
        if custom_prompt and str(custom_prompt).strip()
        else question
    )

    # ----------------------------------------------------
    # Build RAG text context
    # ----------------------------------------------------
    context_text = "\n\n".join(d.page_content for d in docs)

    # ----------------------------------------------------
    # Collect text filenames
    # ----------------------------------------------------
    text_files = list({d.metadata.get("source_file", "unknown") for d in docs})

    # ----------------------------------------------------
    # Normalize image filenames
    # ----------------------------------------------------
    def compute_image_file(img):
        """
        Manual.pdf/page_6.png → Manual.pdf_page_6.png
        """
        url = img["url"] if isinstance(img, dict) else img
        parts = url.split("/")
        pdf_file = parts[-2]
        png_name = parts[-1].split("?")[0]
        return f"{pdf_file}_{png_name}"

    image_files = list({compute_image_file(img) for img in image_urls})

    # ----------------------------------------------------
    # Evidence rules
    # ----------------------------------------------------
    evidence_block = f"""
EVIDENCE SOURCE RULES (follow EXACTLY):

1. Classify files ONLY by extension:
   - ".pdf" → text_files
   - ".jpg" / ".jpeg" / ".png" → image_files
   - PDF-derived images (xxx_page_7.png) → image_files

2. NEVER mix file types.
3. Include ONLY files that provide evidence.

Available files:
TEXT FILES: {text_files}
IMAGE FILES: {image_files}

JSON schema (MANDATORY):

"evidence_source": {{
    "type": "text" | "image" | "both",
    "files": {{
        "text_files": [],
        "image_files": []
    }}
}}
"""

    # ----------------------------------------------------
    # System prompt (single path)
    # ----------------------------------------------------
    system_prompt = f"""
You are an expert electrical safety compliance engineer specializing in IEC 61010-1.

Rules:
- Do NOT guess or infer missing information
- Use IEC 61010-1 only as interpretive guidance
- Base answers strictly on provided TRF text and images
- Every answer must be evidence-based

TASK:
{active_prompt}

Provide ONLY the following JSON keys:
1. "response": concise extracted answer (max 10 words)
2. "confidence": integer score from 1–100 based on evidence strength
3. "evidence_source": {{
      "type": "text" | "image" | "both",
      "files": {{
          "text_files": [],
          "image_files": []
      }}
}}

{evidence_block}

TRF TEXT CONTEXT:
{context_text}
""".strip()

    # ----------------------------------------------------
    # Build Vision message
    # ----------------------------------------------------
    content_blocks = [{"type": "text", "text": system_prompt}]

    for img in image_urls:
        url = img["url"] if isinstance(img, dict) else img
        content_blocks.append({
            "type": "image_url",
            "image_url": {"url": url}
        })

    return [{"role": "user", "content": content_blocks}]


######v5---final

def attach_supporting_refs_grey(
        docs,
        image_urls,
        llm_output,        # STRING with mixed text + JSON
        llm_classifier,   # backward compatibility, UNUSED
        vs,
        accuracy_level=None,
        grey=False,
        question=None,
        k=5):

    import json
    import re

    # ---------------------------------------------------
    # STEP 0 — Robust JSON extraction from raw string
    # ---------------------------------------------------
    def extract_json_from_text(text):
        # Case 1: valid {...} JSON block exists
        matches = re.findall(r"\{[\s\S]*\}", text)
        for block in reversed(matches):
            try:
                return json.loads(block)
            except Exception:
                continue

        # Case 2: loose JSON fields (missing outer braces)
        if '"response"' in text and '"confidence"' in text:
            try:
                recovered = "{" + text.split('"response"', 1)[1]
                recovered = recovered.replace("\n", " ").strip()
                return json.loads(recovered)
            except Exception:
                pass

        return None

    parsed = extract_json_from_text(llm_output)

    if not parsed:
        parsed = {
            "response": llm_output,
            "confidence": 0,
            "evidence_source": {
                "type": "text",
                "files": {}
            }
        }

    # ---------------------------------------------------
    # STEP 1 — Extract LLM-declared image files ONLY
    # ---------------------------------------------------
    evidence = parsed.get("evidence_source", {})
    source_type = evidence.get("type", "text").lower()
    evidence_files = evidence.get("files", {})

    image_files_needed = set()
    if isinstance(evidence_files, dict):
        image_files_needed = set(evidence_files.get("image_files", []))

    text_support = []
    image_support = []

    # ---------------------------------------------------
    # STEP 2 — TEXT SUPPORT (VECTOR SEARCH ONLY)
    # ---------------------------------------------------
    similarity_query = question.strip() if question else parsed.get("response", "")

    try:
        scored_results = vs.similarity_search_with_score(similarity_query, k=k)
    except Exception as e:
        print("Vector search failed:", str(e))
        scored_results = []

    for doc, score in scored_results:
        text_support.append({
            "filename": doc.metadata.get("source_file", "unknown"),
            "pdf_file": doc.metadata.get("source_file", "unknown"),
            "page": doc.metadata.get("page"),
            "similarity_score": float(score),
            "preview_text": doc.page_content[:500].strip()
        })

    # ---------------------------------------------------
    # STEP 3 — IMAGE SUPPORT
    # Filter image_urls using LLM image_files
    # ---------------------------------------------------
    if source_type in ("image", "both") and image_files_needed:

        # Build lookup from image_urls
        image_lookup = {}

        for img in image_urls or []:

            if isinstance(img, dict) and "url" in img:
                image_file = img.get("image_file")

                # Fallback: derive from URL
                if not image_file:
                    image_file = img["url"].split("?")[0].split("/")[-1]

                image_lookup[image_file] = {
                    "url": img.get("url"),
                    "image_file": image_file,
                    "pdf_file": img.get("pdf_file"),
                    "page": img.get("page")
                }

            elif isinstance(img, str):
                image_file = img.split("?")[0].split("/")[-1]
                image_lookup[image_file] = {
                    "url": img,
                    "image_file": image_file,
                    "pdf_file": None,
                    "page": None
                }

        # STRICT filtering by LLM-declared image_files
        for image_file in image_files_needed:
            if image_file in image_lookup:
                image_support.append(image_lookup[image_file])

    # ---------------------------------------------------
    # FINAL structured output
    # ---------------------------------------------------
    return {
        # "answer": parsed.get("response", llm_output),
        "answer":llm_output.strip(),
        "confidence": parsed.get("confidence", 0),

        "evidence_source": {
            "type": source_type,
            "files": {
                "image_files": list(image_files_needed)
            }
        },

        "text_support": text_support,
        "image_support": image_support,

        "grey_mode": grey,
        "accuracy_level": accuracy_level,
        "similarity_query": question,
        "vector_search_based_on": "question"
    }



def build_rag_image_pipeline_v5(
        retriever,
        llm,
        build_vision_message_v5,
        attach_supporting_refs,
        vs
    ):
    """
    FINAL v5:
    - custom_prompt preferred for retrieval & citations
    - fallback to question if custom_prompt is empty/missing
    """

    from langchain_core.runnables import (
        RunnableParallel, RunnableMap, RunnableLambda
    )
    from operator import itemgetter
    from langchain_core.output_parsers import StrOutputParser

    # --------------------------------------------------------
    # Normalize retriever output → list[Document]
    # --------------------------------------------------------
    def extract_docs(ctx):
        if isinstance(ctx, list):
            return ctx
        if isinstance(ctx, dict):
            for key in ("context", "documents", "docs"):
                if key in ctx and isinstance(ctx[key], list):
                    return ctx[key]
        return []

    # --------------------------------------------------------
    # Shared selector: custom_prompt → question (NEW)
    # --------------------------------------------------------
    def select_prompt(x):
        custom = (x.get("custom_prompt") or "").strip()
        if custom:
            return custom
        return x.get("question", "")

    rag_image = (

        # ----------------------------------------------------
        # STEP 1 — Retrieve context + passthrough inputs
        # ----------------------------------------------------
        RunnableParallel(
            # Retrieval uses selector
            context=RunnableLambda(select_prompt) | retriever,

            question=itemgetter("question"),
            custom_prompt=itemgetter("custom_prompt"),
            image=itemgetter("image"),
        )

        # ----------------------------------------------------
        # STEP 2 — Normalize context + build vision message
        # ----------------------------------------------------
        | RunnableMap({
            "inputs": RunnableLambda(lambda x: {
                **x,
                "context": extract_docs(x.get("context"))
            }),
            "messages": RunnableLambda(
                lambda x: build_vision_message_v5(x)
            )
        })

        # ----------------------------------------------------
        # STEP 3 — LLM call (raw string output)
        # ----------------------------------------------------
        | RunnableMap({
            "inputs": itemgetter("inputs"),
            "raw_llm_output": (
                itemgetter("messages")
                | llm
                | StrOutputParser()
            )
        })

        # ----------------------------------------------------
        # STEP 4 — Attach supporting references (UPDATED)
        # ----------------------------------------------------
        | RunnableLambda(
            lambda x: {
                **attach_supporting_refs(
                    docs=x["inputs"].get("context", []),
                    image_urls=x["inputs"].get("image", []),
                    llm_output=x.get("raw_llm_output", ""),
                    llm_classifier=None,
                    vs=vs,

                    # 🔽 SAME LOGIC as retrieval
                    question=select_prompt(x["inputs"]),
                    k=5
                ),

                # ----------------------------
                # Metadata (audit-safe)
                # ----------------------------
                "metadata": {
                    "question": x["inputs"].get("question"),
                    "custom_prompt": x["inputs"].get("custom_prompt"),
                },

                "raw_llm_output": x.get("raw_llm_output"),
            }
        )
    )

    return rag_image

@retry(
    retry=retry_if_exception_type(RateLimitError),
    stop=stop_never,   # Never stop, retry until available
    wait=wait_exponential(multiplier=2, min=5, max=300),
    before_sleep=print_retry_details
)
def run_single_task(task,rag_image):
    return rag_image.invoke(task)
    
@retry(
retry=retry_if_exception_type(RateLimitError),
stop=stop_never,   # Never stop, retry until available
wait=wait_exponential(multiplier=2, min=5, max=300),
before_sleep=print_retry_details
)
def run_single_task_stats(task, rag_image):
    with get_openai_callback() as cb:
        response = rag_image.invoke(task)
    
    response["_token_usage"] = {
        "prompt": cb.prompt_tokens,
        "completion": cb.completion_tokens,
        "total": cb.total_tokens
    }
    return response



def build_tasks_with_custom_prompt_letter(data):
    tasks = []
    item_refs = []

    # Process items
    for page in data.get("pages", []):
        for item in page.get("items", []):

            if not item.get("ai_fillable", False):
                continue

            if item.get("value", "").strip() != "":
                continue  # skip if value already filled

            question = item.get("question", "")
            custom = item.get("custom_prompt", "")

            tasks.append({
                "question": question,
                "custom_prompt": custom,
                "image": []  # remains empty, retriever will update later
            })

            item_refs.append(item)

    return tasks, item_refs



##dont run
import os
import re
from urllib.parse import urlparse
from concurrent.futures import ThreadPoolExecutor, as_completed



def extract_pdf_and_page_from_blob_url(blob_url: str):
    """
    Extract:
    - pdf_file: folder name ending with .pdf
    - image_file: pdf_file_page_x.png (UNDERSCORE FORMAT)
    - page: extracted page number
    """

    parsed = urlparse(blob_url)
    path = parsed.path   # /container/.../<pdf_folder>/<image>.png

    parts = path.split("/")

    pdf_file = None
    for part in parts:
        if part.lower().endswith(".pdf"):
            pdf_file = part
            break

    image_name = os.path.basename(path)

    m = re.search(r"page_(\d+)\.png$", image_name)
    page = int(m.group(1)) if m else None

    if pdf_file:
        # Build: pdf_file_page_x.png
        image_file = f"{pdf_file}/{image_name}"
        image_file = image_file.replace("/", "_")   # FIX ✔
    else:
        image_file = image_name

    return pdf_file, page, image_file



def build_image_entry(blob_url: str):
    pdf_file, page, image_file = extract_pdf_and_page_from_blob_url(blob_url)

    return {
        "pdf_file": pdf_file,
        "page": page,
        "image_file": image_file,
        "url": blob_url
    }


def process_single_task(task, retriever, task_index, total_tasks):
    """
    Processes ONE task:
    - runs retriever
    - builds 5 image entries
    """
    question = task.get("custom_prompt")

    print(f"[INFO] Processing task {task_index}/{total_tasks} → {question}")

    neighbors = retriever.invoke(question)

    images = []
    for d in neighbors[:5]:
        blob_url = d.metadata.get("blob_url")
        entry = build_image_entry(blob_url)
        images.append(entry)

    task_new = dict(task)
    task_new["image"] = images

    print(f"[SUCCESS] Completed task {task_index}/{total_tasks}")

    return task_new


def update_tasks_with_top5_images(tasks, retriever, max_threads=10):
    total = len(tasks)
    updated = [None] * total  # preserve order

    print(f"[START] Processing {total} tasks with {max_threads} parallel threads.\n")

    with ThreadPoolExecutor(max_workers=max_threads) as executor:

        future_map = {
            executor.submit(process_single_task, task, retriever, idx+1, total): idx
            for idx, task in enumerate(tasks)
        }

        for future in as_completed(future_map):
            idx = future_map[future]
            try:
                updated[idx] = future.result()
            except Exception as e:
                print(f"[ERROR] Task {idx+1} failed: {e}")

    print(f"\n[COMPLETE] Completed all {total} tasks.\n")

    return updated


# # for processing images (dont run)
# def build_vectorstore2(embeddings):
#     COSMOS_URL    = "https://rag-intertek.documents.azure.com:443/"
#     COSMOS_KEY    = "AbhkomWJLtf8TR7odpABPqx1OrjlmCcpTXlKr9Vvp3RulZmFGollxQflIp3LLUAFt4XcMh70RbRxACDbuxyZLg=="
#     # DB_NAME     = "ragdatabase_new"
#     # CONT_NAME   = "vectorstorecontainer_new"
#     DB_NAME     = "ragdatabase_new_itk2"
#     CONT_NAME   = "vectorstorecontainer_new_itk2"
    
#     # COSMOS_URL    = "https://csdb-intertek-esus-dev.documents.azure.com:443/"
#     # COSMOS_KEY    = "azcUeVxFxoYoFkChvWI8Wr8lMijOuWXDYQsvMf6O2LmT0Uv3Zs7lDPiXSxWYOjq00MFDbK88ApotACDbODLFXA=="
#     cosmos_client = CosmosClient(
#         url=COSMOS_URL,
#         credential=COSMOS_KEY,
#         consistency_level=ConsistencyLevel.Eventual
#     )

#     # keep your existing policy helpers if your constructor requires them
#     return AzureCosmosDBNoSqlVectorSearch(
#         cosmos_client=cosmos_client,
#         embedding=embeddings,
#         database_name=DB_NAME,
#         container_name=CONT_NAME,

#         # if your version requires explicit policies, keep these as you already had:
#         vector_embedding_policy={"vectorEmbeddings":[{"path":"/vector","dataType":"float32","dimensions":1536,"distanceFunction":"cosine"}]},
#         indexing_policy={"includedPaths":[{"path":"/*"}],
#                          "excludedPaths":[{"path":"/\"_etag\"/?"},{"path":"/vector/*"}],
#                          "vectorIndexes":[{"path":"/vector","type":"quantizedFlat"}]},
#         cosmos_container_properties={"partition_key":"/id"},
#         cosmos_database_properties={}, # _db_props()

#         # IMPORTANT: pass a dict, not a list
#         vector_search_fields={
#             "text_field": "text",
#             "embedding_field": "vector",
#             "metadata_field": "metadata"
#         }
#     )


import json
import re

# --------------------------------------------------
def extract_human_prefix(text):
    if not isinstance(text, str):
        return ""
    return text.split("{", 1)[0].strip()


# --------------------------------------------------
def extract_json_block(text):
    """
    Robust extraction: finds LAST valid JSON object
    """
    if not isinstance(text, str):
        return None

    matches = re.findall(r"\{[\s\S]*\}", text)
    for block in reversed(matches):
        try:
            return json.loads(block)
        except Exception:
            continue
    return None


# --------------------------------------------------
def extract_pdf_page_from_filename(fname):
    m = re.search(r"(.+\.pdf)_page_(\d+)\.png$", fname)
    if not m:
        return None, None
    return m.group(1), int(m.group(2))


def update_json_item_letter(item, result):
    import json

    answer_text = result.get("answer", "")
    raw_json = result.get("llm_raw_json")

    # -----------------------------
    # 1. Parse JSON safely
    # -----------------------------
    if raw_json:
        try:
            parsed = json.loads(raw_json)
        except Exception:
            parsed = {}
    else:
        parsed = extract_json_block(answer_text) or {}

    # Human prefix override
    human_prefix = extract_human_prefix(answer_text)
    if human_prefix and len(human_prefix.split()) > 3:
        parsed["response"] = human_prefix

    response = str(parsed.get("response", "")).strip()
    confidence = parsed.get("confidence", result.get("confidence", 0))

    # -----------------------------
    # 2. TBD short-circuit
    # -----------------------------
    if response.replace(" ", "").startswith("TBD"):
        item["value"] = response
        item["confidence"] = confidence
        return

    # -----------------------------
    # 3. Assign VALUE (CORE CHANGE)
    # -----------------------------
    item["value"] = response

    # -----------------------------
    # 4. Evidence handling (FIXED)
    # -----------------------------
    evidence = parsed.get("evidence_source", {}) or {}
    files = evidence.get("files") or {}

    # FIX: Prevent crash if "files" comes as a list
    if isinstance(files, list):
        # Attempt auto-recovery if list contains {"image_files": [...]}
        for entry in files:
            if isinstance(entry, dict) and "image_files" in entry:
                files = entry
                break
        else:
            files = {}

    json_image_files = files.get("image_files") or []

    # -----------------------------
    # 5. Text support
    # -----------------------------
    similarity_preview = (
        result.get("similarity_preview")
        or result.get("text_support")
        or []
    )

    item["text_support"] = [
        {
            "filename": hit.get("filename"),
            "page": hit.get("page"),
            "similarity_score": hit.get("similarity_score"),
            "preview_text": hit.get("preview_text"),
        }
        for hit in similarity_preview
    ]

    # -----------------------------
    # 6. Image support (safe)
    # -----------------------------
    pipeline_images = (
        result.get("image_support")
        or result.get("pipeline_output", {}).get("image_support")
        or []
    )

    pipeline_lookup = {d.get("image_file"): d for d in pipeline_images}
    task_images = item.get("image", []) or []
    task_lookup = {d.get("image_file"): d for d in task_images}

    image_support_final = []

    for fname in json_image_files:
        pdf_file, page = extract_pdf_page_from_filename(fname)

        enriched = {
            "filename": fname,
            "image_file": fname,
            "pdf_file": pdf_file,
            "page": page,
            "url": None
        }

        if fname in pipeline_lookup:
            enriched.update({k: v for k, v in pipeline_lookup[fname].items() if v is not None})

        if enriched["url"] is None and fname in task_lookup:
            enriched.update({k: v for k, v in task_lookup[fname].items() if v is not None})

        image_support_final.append(enriched)

    if not image_support_final:
        image_support_final = list(pipeline_images)

    item["image_support"] = image_support_final
    item["evidence_type"] = evidence.get("type") or "text"
    item["confidence"] = confidence

#### Updated v5.1
from concurrent.futures import ThreadPoolExecutor, as_completed
import time

def process_tasks_with_batches_parallel_grey(
        tasks,
        item_refs,
        rag_image,
        vs,
        batch_size=150,
        cooldown_sec=15,
        max_workers=6,
        use_llm_inGrey=False,
        stats=False
    ):

    total = len(tasks)
    processed = 0
    all_results = []

    # Stats
    total_prompt_tokens = 0
    total_completion_tokens = 0
    total_total_tokens = 0
    total_llm_calls = 0

    # -----------------------------------------------------
    # GREY fallback handler (VECTOR ONLY)
    # -----------------------------------------------------
    def handle_low_accuracy_item(item, task, vs, k=5):

        if use_llm_inGrey:
            return False

        if task.get("accuracy_level", True):
            return False

        question = task["question"]

        try:
            scored_results = vs.similarity_search_with_score(question, k=k)
        except Exception:
            scored_results = []

        text_sources = [
            {
                "filename": doc.metadata.get("source_file", "unknown"),
                "page": doc.metadata.get("page"),
                "similarity_score": float(score),
                "preview_text": doc.page_content[:500].strip(),
            }
            for doc, score in scored_results
        ]

        # Threshold logic
        has_text_info = any(score >= 0.76 for _, score in scored_results)

        question_l = question.lower()
        image_keywords = ["label", "marking", "diagram", "schematic", "figure", "block", "display"]
        has_image_hint = any(kw in question_l for kw in image_keywords)

        has_info = has_text_info or has_image_hint
        value = "TBD-Info available" if has_info else "TBD-Info not available"

        # Update item directly
        item.update({
            "value": value,
            "confidence": 0,
            "text_support": text_sources,
            "image_support": [],   # STRICT: no images in GREY
            "similarity_query": question,
            "vector_search_based_on": "question",
            "evidence_source": {
                "type": "text" if text_sources else "none",
                "files": {
                    "text_files": sorted({x["filename"] for x in text_sources}),
                    "image_files": []
                }
            }
        })

        all_results.append(item)
        return True

    # -----------------------------------------------------
    # MAIN LOOP
    # -----------------------------------------------------
    for start in range(0, total, batch_size):
        end = min(start + batch_size, total)
        batch = tasks[start:end]
        ref_batch = item_refs[start:end]

        print(f"\n🔵 Starting batch: {start+1} → {end}")

        with ThreadPoolExecutor(max_workers=max_workers) as exe:
            futures = {}

            for idx, task in enumerate(batch):
                item = ref_batch[idx]

                # 🔴 DECIDE GREY BEFORE LLM
                if handle_low_accuracy_item(item, task, vs):
                    processed += 1
                    print(f"Processed {processed}/{total} (vector-only GREY)")
                    continue

                # LLM path only if not GREY
                if stats:
                    fut = exe.submit(run_single_task_stats, task, rag_image)
                else:
                    fut = exe.submit(run_single_task, task, rag_image)

                futures[fut] = idx

            for future in as_completed(futures):
                idx = futures[future]
                task = batch[idx]
                item = ref_batch[idx]

                result = future.result()

                if stats and "_token_usage" in result:
                    usage = result["_token_usage"]
                    total_prompt_tokens += usage.get("prompt", 0)
                    total_completion_tokens += usage.get("completion", 0)
                    total_total_tokens += usage.get("total", 0)
                    total_llm_calls += 1

                result["question"] = task["question"]
                result["task_type"] = task.get("task_type")

                update_json_item_letter(item, result)

                all_results.append(result)
                processed += 1
                print(f"Processed {processed}/{total}")

        if end < total:
            print(f"⏳ Cooling down for {cooldown_sec} seconds...")
            time.sleep(cooldown_sec)

    print("\n✅ ALL TASKS COMPLETED.")

    if stats:
        return {
            "results": all_results,
            "stats": {
                "total_llm_calls": total_llm_calls,
                "total_prompt_tokens": total_prompt_tokens,
                "total_completion_tokens": total_completion_tokens,
                "total_tokens": total_total_tokens,
            }
        }

    return all_results

import subprocess
import os
import shutil
import sys

def convert_doc_to_docx(doc_path: str, output_dir: str = None) -> str:
    """
    Converts a .doc file to .docx using LibreOffice.
    Works on Linux, macOS, and Windows.
    """

    if not doc_path.lower().endswith(".doc"):
        raise ValueError("Input file must be a .doc file")

    if not os.path.exists(doc_path):
        raise FileNotFoundError(doc_path)

    soffice_path = shutil.which("soffice")

    if soffice_path is None and sys.platform.startswith("win"):
        soffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"

    if not soffice_path or not os.path.exists(soffice_path):
        raise RuntimeError("LibreOffice (soffice) not found")

    if output_dir is None:
        output_dir = os.getcwd()

    os.makedirs(output_dir, exist_ok=True)

    subprocess.run(
        [
            soffice_path,
            "--headless",
            "--convert-to",
            "docx",
            doc_path,
            "--outdir",
            output_dir,
        ],
        check=True
    )

    base = os.path.splitext(os.path.basename(doc_path))[0]
    return os.path.join(output_dir, base + ".docx")


from docx import Document

# Load the Word document
## idetify the file format if pf convert to docx and read table 51
# def extract_table1a(filename):
        
#     doc = Document(filename)
#     table = doc.tables[51]

#     import pandas as pd

#     data = []
#     for row in table.rows:
#         data.append([cell.text.strip() for cell in row.cells])

#     # Assuming first row is header
#     df = pd.DataFrame(data[:])
#     return df

def extract_table1a(filename):
    from docx import Document
    import pandas as pd

    doc = Document(filename)

    target_table = None

    for table in doc.tables:
        table_text = " ".join(
            cell.text.strip()
            for row in table.rows
            for cell in row.cells
        ).lower()

        if "table 1.a" in table_text and "list of components" in table_text:
            target_table = table
            break

    if target_table is None:
        # raise ValueError(
        #     "Table containing 'Table 1.A: List of components' not found"
        # )
        return None

    data = [
        [cell.text.strip() for cell in row.cells]
        for row in target_table.rows
    ]

    return pd.DataFrame(data)
# import os
# import pandas as pd
# from typing import Optional


# def extract_table1a(filename: Optional[str]) -> pd.DataFrame:
#     """
#     Extract Table 1.A from a DOCX file.
#     If DOCX is missing or Table 1.A is not found,
#     fall back to extracting from CDR Components Excel
#     located in the same directory.

#     Args:
#         filename: Path to CDR .docx file

#     Returns:
#         pd.DataFrame
#     """

#     # Directory for Excel fallback
#     src_directory = (
#         os.path.dirname(filename)
#         if filename and os.path.exists(filename)
#         else None
#     )

#     # -------------------------------
#     # 1️⃣ Attempt DOCX extraction
#     # -------------------------------
#     if filename and os.path.exists(filename):
#         try:
#             from docx import Document

#             doc = Document(filename)
#             target_table = None

#             for table in doc.tables:
#                 table_text = " ".join(
#                     cell.text.strip()
#                     for row in table.rows
#                     for cell in row.cells
#                 ).lower()

#                 if "table 1.a" in table_text and "list of components" in table_text:
#                     target_table = table
#                     break

#             if target_table:
#                 data = [
#                     [cell.text.strip() for cell in row.cells]
#                     for row in target_table.rows
#                 ]
#                 return pd.DataFrame(data)

#         except Exception as e:
#             print(f"DOCX extraction failed, falling back to Excel: {e}")

#     # --------------------------------
#     # 2️⃣ Fallback → Excel extraction
#     # --------------------------------
#     if src_directory:
#         df = load_cdr_components_df(src_directory)
#         if df is not None:
#             return df

#     # --------------------------------
#     # 3️⃣ Hard failure
#     # --------------------------------
#     raise RuntimeError(
#         "Failed to extract Table 1.A from both DOCX and Excel sources."
#     )


from typing import Optional
import os
import pandas as pd

import pandas as pd

def format_critical_components_df(raw_df: pd.DataFrame) -> pd.DataFrame:
    """
    Formats extracted IEC table into clean Critical Components table.
    - Drops first 3 and last 3 rows
    - Resets Item numbering starting from 1
    """

    # --------------------------------------------------
    # 1. Rename columns by position
    # --------------------------------------------------
    df = raw_df.rename(columns={
        0: "Name",
        2: "Manufacturer/ trademark",
        3: "Type / Model",
        5: "Technical data and securement means",
        9: "Mark(s) of conformity"
    })

    # --------------------------------------------------
    # 2. Keep only required columns
    # --------------------------------------------------
    df = df[
        [
            "Name",
            "Manufacturer/ trademark",
            "Type / Model",
            "Technical data and securement means",
            "Mark(s) of conformity"
        ]
    ]

    # --------------------------------------------------
    # 3. Remove header / note / empty rows
    # --------------------------------------------------
    df = df[
        df["Name"].notna()
        & ~df["Name"].str.contains(
            r"IEC|Clause|table|NOTE|List all|May include|licence",
            case=False,
            na=False
        )
    ]

    # --------------------------------------------------
    # 4. Drop first 3 and last 3 rows
    # --------------------------------------------------
    df = df.iloc[3:-3].reset_index(drop=True)

    # --------------------------------------------------
    # 5. Remove Item column if already present
    # --------------------------------------------------
    if "Item" in df.columns:
        df = df.drop(columns=["Item"])

    # --------------------------------------------------
    # 6. Add Item column as first column (reset counter)
    # --------------------------------------------------
    df.insert(0, "Item", range(1, len(df) + 1))

    # --------------------------------------------------
    # 7. Clean whitespace
    # --------------------------------------------------
    df = df.applymap(
        lambda x: " ".join(str(x).split()) if pd.notna(x) else x
    )

    return df


from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd


def insert_dataframe_below_anchor_dep(
    input_docx: str,
    output_docx: str,
    df: pd.DataFrame,
    anchor_text: str
):
    """
    Inserts a dataframe as a table immediately below the paragraph
    containing anchor_text. Preserves all existing formatting.
    Header row is formatted with Amber background.
    """

    doc = Document(input_docx)

    anchor_paragraph = None
    for para in doc.paragraphs:
        if anchor_text in para.text:
            anchor_paragraph = para
            break

    if anchor_paragraph is None:
        raise ValueError("Anchor text not found in document.")

    rows, cols = df.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.style = "Table Grid"

    # -----------------------------
    # Header row (Amber background)
    # -----------------------------
    for col_idx, col_name in enumerate(df.columns):
        cell = table.rows[0].cells[col_idx]
        cell.text = str(col_name)

        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "FFC000")  # Amber
        tc_pr.append(shd)

    # -----------------------------
    # Data rows
    # -----------------------------
    for row_idx in range(rows):
        for col_idx in range(cols):
            table.rows[row_idx + 1].cells[col_idx].text = str(df.iat[row_idx, col_idx])

    # -----------------------------
    # Move table below anchor
    # -----------------------------
    anchor_paragraph._p.addnext(table._element)

    doc.save(output_docx)

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd

# def insert_dataframe_below_anchor(
#     input_docx,
#     output_docx,
#     df,
#     anchor_text
# ):
#     doc = Document(input_docx)

#     anchor_paragraph = None
#     for para in doc.paragraphs:
#         if anchor_text in para.text:
#             anchor_paragraph = para
#             break

#     if anchor_paragraph is None:
#         raise ValueError("Anchor text not found in document.")

#     # -------------------------------------------------
#     # INSERT BLANK PARAGRAPH AFTER ANCHOR (XML SAFE)
#     # -------------------------------------------------
#     blank_p = OxmlElement("w:p")
#     anchor_paragraph._p.addnext(blank_p)

#     # -------------------------------------------------
#     # CREATE TABLE
#     # -------------------------------------------------
#     rows, cols = df.shape
#     table = doc.add_table(rows=rows + 1, cols=cols)
#     table.style = "Table Grid"

#     # -----------------------------
#     # Header row (Amber background)
#     # -----------------------------
#     for col_idx, col_name in enumerate(df.columns):
#         cell = table.rows[0].cells[col_idx]
#         cell.text = str(col_name)

#         tc_pr = cell._tc.get_or_add_tcPr()
#         shd = OxmlElement("w:shd")
#         shd.set(qn("w:fill"), "FFC000")  # Amber
#         tc_pr.append(shd)

#     # -----------------------------
#     # Data rows
#     # -----------------------------
#     for row_idx in range(rows):
#         for col_idx in range(cols):
#             table.rows[row_idx + 1].cells[col_idx].text = str(df.iat[row_idx, col_idx])

#     # -------------------------------------------------
#     # INSERT TABLE AFTER BLANK PARAGRAPH
#     # -------------------------------------------------
#     blank_p.addnext(table._element)

#     doc.save(output_docx)

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def insert_dataframe_below_anchor(
    input_docx,
    output_docx,
    df,
    anchor_text
):
    if df is None or df.empty:
        print("⚠️ DataFrame is empty. Skipping table insertion.")
        return

    doc = Document(input_docx)

    # 1️⃣ Find anchor paragraph
    anchor_paragraph = None
    for para in doc.paragraphs:
        if anchor_text.strip() in para.text.strip():
            anchor_paragraph = para
            break

    if anchor_paragraph is None:
        raise ValueError("Anchor text not found in document.")

    # 2️⃣ Create table FIRST (Word-safe)
    rows, cols = df.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.style = "Table Grid"

    # Header row
    for col_idx, col_name in enumerate(df.columns):
        cell = table.rows[0].cells[col_idx]
        cell.text = str(col_name)

        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "FFC000")
        tc_pr.append(shd)

    # Data rows
    for r in range(rows):
        for c in range(cols):
            table.rows[r + 1].cells[c].text = str(df.iat[r, c])

    # 3️⃣ Move table XML BELOW anchor paragraph
    anchor_paragraph._p.addnext(table._element)

    doc.save(output_docx)

    print("✅ Table inserted below anchor.")


from docx import Document

def replace_keys_with_values_no_format_change_dep(
    input_docx: str,
    output_docx: str,
    data: dict
):
    """
    Replaces placeholders with values WITHOUT changing formatting.
    Only replaces text inside the same Word run.
    """

    doc = Document(input_docx)

    # -------------------------------------------------
    # Build replacement map
    # -------------------------------------------------
    replacements = {}

    for page in data.get("pages", []):
        for item in page.get("items", []):
            key = str(item.get("key", "")).strip()
            value = str(item.get("value", "")).strip()

            if not key or not value:
                continue

            replacements[key] = value
            replacements[f"<{key}>"] = value

    # -------------------------------------------------
    # SAFE run-level replacement
    # -------------------------------------------------
    def replace_in_runs(paragraph):
        for run in paragraph.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)

    # Paragraphs
    for para in doc.paragraphs:
        replace_in_runs(para)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para)

    # -------------------------------------------------
    # Save output
    # -------------------------------------------------
    doc.save(output_docx)

from docx import Document

def replace_keys_with_values_no_format_change_v2(
    input_docx,
    output_docx,
    data
):
    """
    Replaces placeholders with values WITHOUT changing formatting.
    Works even if Word splits placeholder across multiple runs.
    Only replaces keys where ai_fillable == True.
    """

    doc = Document(input_docx)

    # -------------------------------------------------
    # Build replacement map (ai_fillable only)
    # -------------------------------------------------
    replacements = {}

    for page in data.get("pages", []):
        for item in page.get("items", []):

            if item.get("ai_fillable") is not True:
                continue

            key = str(item.get("key", "")).strip()
            value = str(item.get("value", "")).strip()

            if not key:
                continue

            replacements[key] = value
            replacements[f"<{key}>"] = value  # optional bracket form

    # -------------------------------------------------
    # Run-join fallback replacement
    # -------------------------------------------------
    def replace_in_para(para):
        if not para.runs:
            return

        full_text = "".join(run.text for run in para.runs)

        for old, new in replacements.items():
            if old in full_text:
                full_text = full_text.replace(old, new)

        # Clear runs but keep formatting
        for run in para.runs:
            run.text = ""

        para.runs[0].text = full_text

    # -------------------------------------------------
    # Replace in paragraphs
    # -------------------------------------------------
    for para in doc.paragraphs:
        replace_in_para(para)

    # -------------------------------------------------
    # Replace in tables
    # -------------------------------------------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

    doc.save(output_docx)

def replace_keys_with_values_no_format_change_all(
    input_docx: str,
    output_docx: str,
    data: dict
):
    """
    Replaces placeholders with values WITHOUT changing formatting.
    Works even if Word splits placeholder across multiple runs.
    """

    doc = Document(input_docx)

    # Build replacement map
    replacements = {}
    for page in data.get("pages", []):
        for item in page.get("items", []):
            key = str(item.get("key", "")).strip()
            value = str(item.get("value", "")).strip()
            if not key:
                continue
            replacements[key] = value
            replacements[f"<{key}>"] = value  # optional bracket form

    # --- RUN JOIN FALLBACK REPLACEMENT ---
    def replace_in_para(para):
        full_text = "".join(run.text for run in para.runs)
        for old, new in replacements.items():
            if old in full_text:
                full_text = full_text.replace(old, new)

        # Push replaced text back while keeping formatting
        for run in para.runs:
            run.text = ""  # clear
        if para.runs:
            para.runs[0].text = full_text  # write in first run

    # Replace in paragraphs
    for para in doc.paragraphs:
        replace_in_para(para)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

    doc.save(output_docx)


from docx import Document

def replace_header_keys_with_values(input_docx, output_docx, data):
    """
    Replaces header placeholders from data_header['pages'][*]['items'][*]
    'key' → 'value' without keeping guillemets.
    """

    doc = Document(input_docx)

    # Build replacement map (no fallback wrapping)
    replacements = {}
    for page in data.get("pages", []):
        for item in page.get("items", []):
            placeholder = str(item.get("key", "")).strip()
            value = str(item.get("value", "")).strip()
            if not placeholder or not value:
                continue
            # Direct exact replacement only
            replacements[placeholder] = value

    # Replace helper
    def replace_in_header(header):
        for para in header.paragraphs:
            for run in para.runs:
                for old, new in replacements.items():
                    if old in run.text:
                        run.text = run.text.replace(old, new)

        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for old, new in replacements.items():
                                if old in run.text:
                                    run.text = run.text.replace(old, new)

    # Replace in all section headers
    for section in doc.sections:
        replace_in_header(section.header)

    doc.save(output_docx)
    return output_docx


import os

# def ensure_docx(file_path: str) -> str:
#     """
#     Ensures the file is in DOCX format.
#     - If .doc  -> converts to .docx
#     - If .docx -> returns as-is
#     - Else     -> raises error
#     """

#     if not os.path.exists(file_path):
#         raise FileNotFoundError(file_path)

#     ext = os.path.splitext(file_path)[1].lower()

#     if ext == ".docx":
#         return file_path

#     if ext == ".doc":
#         return convert_doc_to_docx(file_path)

#     raise ValueError(f"Unsupported file type: {ext}. Only .doc or .docx allowed.")

import os
from typing import Optional

def ensure_docx(file_path: Optional[str]) -> Optional[str]:
    """
    Safely ensures the file is in DOCX format.

    - If file_path is None → returns None
    - If .docx → returns as-is
    - If .doc  → converts to .docx and returns path
    - If file does not exist / unsupported type → returns None
    """

    if not file_path:
        return None

    if not os.path.exists(file_path):
        return None

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".docx":
        return file_path

    if ext == ".doc":
        try:
            return convert_doc_to_docx(file_path)
        except Exception as e:
            print(f"⚠️ DOC → DOCX conversion failed: {e}")
            return None

    # Unsupported type → silently ignore
    return None

import pandas as pd

# def format_critical_components_df(raw_df: pd.DataFrame) -> pd.DataFrame:
#     """
#     Formats extracted IEC table into clean Critical Components table.
#     - Drops first 3 and last 3 rows
#     - Resets Item numbering starting from 1
#     """

#     # --------------------------------------------------
#     # 1. Rename columns by position
#     # --------------------------------------------------
#     df = raw_df.rename(columns={
#         0: "Name",
#         2: "Manufacturer/ trademark",
#         3: "Type / Model",
#         5: "Technical data and securement means",
#         9: "Mark(s) of conformity"
#     })

#     # --------------------------------------------------
#     # 2. Keep only required columns
#     # --------------------------------------------------
#     df = df[
#         [
#             "Name",
#             "Manufacturer/ trademark",
#             "Type / Model",
#             "Technical data and securement means",
#             "Mark(s) of conformity"
#         ]
#     ]

#     # --------------------------------------------------
#     # 3. Remove header / note / empty rows
#     # --------------------------------------------------
#     df = df[
#         df["Name"].notna()
#         & ~df["Name"].str.contains(
#             r"IEC|Clause|table|NOTE|List all|May include|licence",
#             case=False,
#             na=False
#         )
#     ]

#     # --------------------------------------------------
#     # 4. Drop first 3 and last 3 rows
#     # --------------------------------------------------
#     df = df.iloc[3:-3].reset_index(drop=True)

#     # --------------------------------------------------
#     # 5. Remove Item column if already present
#     # --------------------------------------------------
#     if "Item" in df.columns:
#         df = df.drop(columns=["Item"])

#     # --------------------------------------------------
#     # 6. Add Item column as first column (reset counter)
#     # --------------------------------------------------
#     df.insert(0, "Item", range(1, len(df) + 1))

#     # --------------------------------------------------
#     # 7. Clean whitespace
#     # --------------------------------------------------
#     df = df.applymap(
#         lambda x: " ".join(str(x).split()) if pd.notna(x) else x
#     )

#     return df

import pandas as pd

def format_critical_components_df(raw_df: pd.DataFrame) -> pd.DataFrame:
    """
    Formats extracted IEC table into clean Critical Components table.
    """

    # --------------------------------------------------
    # 0. Defensive guard
    # --------------------------------------------------
    if raw_df is None or raw_df.empty:
        return pd.DataFrame(
            columns=[
                "Item",
                "Name",
                "Manufacturer/ trademark",
                "Type / Model",
                "Technical data and securement means",
                "Mark(s) of conformity"
            ]
        )

    # --------------------------------------------------
    # 1. Rename columns by position
    # --------------------------------------------------
    df = raw_df.rename(columns={
        0: "Name",
        2: "Manufacturer/ trademark",
        3: "Type / Model",
        5: "Technical data and securement means",
        9: "Mark(s) of conformity"
    })

    # --------------------------------------------------
    # 2. Keep only required columns
    # --------------------------------------------------
    df = df[
        [
            "Name",
            "Manufacturer/ trademark",
            "Type / Model",
            "Technical data and securement means",
            "Mark(s) of conformity"
        ]
    ]

    # --------------------------------------------------
    # 3. Remove header / note / empty rows
    # --------------------------------------------------
    df = df[
        df["Name"].notna()
        & ~df["Name"].str.contains(
            r"IEC|Clause|table|NOTE|List all|May include|licence",
            case=False,
            na=False
        )
    ]

    # --------------------------------------------------
    # 4. Drop first 3 and last 3 rows (safe slice)
    # --------------------------------------------------
    if len(df) > 6:
        df = df.iloc[3:-3]
    df = df.reset_index(drop=True)

    # --------------------------------------------------
    # 5. Remove Item column if already present
    # --------------------------------------------------
    if "Item" in df.columns:
        df = df.drop(columns=["Item"])

    # --------------------------------------------------
    # 6. Add Item column
    # --------------------------------------------------
    df.insert(0, "Item", range(1, len(df) + 1))

    # --------------------------------------------------
    # 7. Clean whitespace
    # --------------------------------------------------
    df = df.applymap(
        lambda x: " ".join(str(x).split()) if pd.notna(x) else x
    )

    return df



from urllib.parse import unquote
import os

def attach_blob_urls_to_text_support(data, blob_urls):
    """
    Adds 'url' to each text_support item by matching filename
    WITHOUT considering file extension.
    """

    # Build base-name → url lookup
    blob_map = {}
    for url in blob_urls:
        fname = unquote(url.split("/")[-1])
        base = os.path.splitext(fname)[0]
        blob_map.setdefault(base, url)   # keep first match

    for table in data.get("Tables", []):
        for item in table.get("Items", []):
            for ts in item.get("text_support", []):
                ts["url"] = None
                fname = ts.get("filename")
                if not fname:
                    continue

                base = os.path.splitext(fname)[0]
                ts["url"] = blob_map.get(base)

    return data

from urllib.parse import unquote
import os

def attach_blob_urls_to_image_support(data, blob_urls):
    """
    Adds 'file_url' to each image_support item
    by matching pdf_file name WITHOUT extension.
    """

    # Build base-name → url lookup
    blob_map = {}
    for url in blob_urls:
        fname = unquote(url.split("/")[-1])
        base = os.path.splitext(fname)[0]
        blob_map.setdefault(base, url)

    for table in data.get("Tables", []):
        for item in table.get("Items", []):
            for img in item.get("image_support", []):
                img["file_url"] = None
                pdf_file = img.get("pdf_file")
                if not pdf_file:
                    continue

                base = os.path.splitext(pdf_file)[0]
                img["file_url"] = blob_map.get(base)

    return data


def extract_table9(filename):
        
    doc = Document(filename)
    table = doc.tables[9]

    import pandas as pd

    data = []
    for row in table.rows:
        data.append([cell.text.strip() for cell in row.cells])

    # Assuming first row is header
    df = pd.DataFrame(data[:])
    return df


import pandas as pd
import re

def fetch_actionable_requirements_dep(df):
    """
    Generic filter:
    - Keeps any meaningful verdict
    - Excludes pass / not applicable / blank / separators
    - Cleans spacing, tabs, casing
    - Removes header & standard rows
    """

    temp = df[[1, 4]].copy()
    temp.columns = ["Requirement", "Verdict"]

    # Normalize
    temp["Requirement"] = temp["Requirement"].astype(str).str.strip()

    temp["Verdict"] = (
        temp["Verdict"]
        .astype(str)
        .str.replace(r"\s+", "", regex=True)
        .str.upper()
    )

    # Generic non-actionable verdicts
    non_actionable_verdicts = {
        "P",          # Pass
        "PASS",
        "N/A",
        "NA",
        "",           # Blank
        "—",          # Separator
        "-"
    }

    # Header / standard noise
    non_data_requirements = {
        "IEC61010-1",
        "IEC 61010-1",
        "VERDICT",
        "REQUIREMENT+TEST",
        "REQUIREMENT + TEST"
    }

    filtered_df = temp[
        (~temp["Verdict"].isin(non_actionable_verdicts)) &
        (~temp["Requirement"].str.upper().isin(non_data_requirements))
    ]

    return filtered_df

import pandas as pd
import re

def fetch_actionable_requirements(df):
    """
    Generic filter:
    - Keeps any meaningful verdict
    - Excludes pass / not applicable / blank / separators
    - Cleans spacing, tabs, casing
    - Removes header & standard rows
    """

    # -----------------------------------
    # Fetch Clause (0), Requirement (1), Verdict (4)
    # -----------------------------------
    temp = df[[0, 1, 4]].copy()
    temp.columns = ["Clause", "Requirement", "Verdict"]

    # Normalize
    temp["Clause"] = temp["Clause"].astype(str).str.strip()
    temp["Requirement"] = temp["Requirement"].astype(str).str.strip()

    temp["Verdict"] = (
        temp["Verdict"]
        .astype(str)
        .str.replace(r"\s+", "", regex=True)
        .str.upper()
    )

    # Generic non-actionable verdicts
    non_actionable_verdicts = {
        "P",          # Pass
        "PASS",
        "N/A",
        "NA",
        "",           # Blank
        "—",          # Separator
        "-"
    }

    # Header / standard noise
    non_data_requirements = {
        "IEC61010-1",
        "IEC 61010-1",
        "VERDICT",
        "REQUIREMENT+TEST",
        "REQUIREMENT + TEST"
    }

    filtered_df = temp[
        (~temp["Verdict"].isin(non_actionable_verdicts)) &
        (~temp["Requirement"].str.upper().isin(non_data_requirements))
    ]

    return filtered_df


def get_first_doc_or_docx(folder_path):
    """
    Returns the full path of the first .doc or .docx file found in a folder.
    Returns None if no such file exists.
    """
    if not os.path.isdir(folder_path):
        raise FileNotFoundError(f"Folder not found: {folder_path}")

    for filename in sorted(os.listdir(folder_path)):
        if filename.lower().endswith((".doc", ".docx")):
            return os.path.join(folder_path, filename)

    return None

import os
import re
import pandas as pd
from typing import Optional, Tuple


# def sheet_name_is_component(sheet_name: str) -> bool:
#     """
#     Detects CDR Components sheets like:
#     - 4.0 Components
#     - Components
#     - Component List
#     """
#     normalized = sheet_name.lower()
#     normalized = re.sub(r"[^a-z]", "", normalized)

#     return "component" in normalized

# def sheet_has_required_columns(
#     excel_path: str,
#     sheet_name: str,
#     max_header_search_rows: int = 15
# ) -> bool:
#     """
#     Scans first N rows to find a valid header row
#     containing Photo + Conformity columns.
#     """

#     try:
#         preview = pd.read_excel(
#             excel_path,
#             sheet_name=sheet_name,
#             header=None,
#             nrows=max_header_search_rows
#         )

#         def normalize(val):
#             return (
#                 str(val)
#                 .lower()
#                 .replace(" ", "")
#                 .replace("#", "")
#                 .replace("_", "")
#             )

#         for row_idx in range(len(preview)):
#             row = preview.iloc[row_idx].tolist()
#             normalized = {normalize(c) for c in row if pd.notna(c)}

#             has_photo = any("photo" in c for c in normalized)
#             has_conform = any("conform" in c for c in normalized)

#             if has_photo and has_conform:
#                 print(f"        Header detected at row {row_idx}")
#                 return True

#         return False

#     except Exception as e:
#         print(f"        Column scan failed: {e}")
#         return False

# #CDR fallback

# # def find_cdr_components_excel(
# #     directory: str
# # ) -> Optional[Tuple[str, str]]:
# #     """
# #     Finds an Excel file containing a valid CDR Components sheet.

# #     Returns:
# #         (excel_path, sheet_name) if found
# #         None otherwise
# #     """

# #     if not os.path.isdir(directory):
# #         raise FileNotFoundError(f"Directory not found: {directory}")

# #     for filename in os.listdir(directory):

# #         # Skip Excel temp/lock files
# #         if filename.startswith("~$"):
# #             continue

# #         if not filename.lower().endswith(".xlsx"):
# #             continue

# #         excel_path = os.path.join(directory, filename)

# #         try:
# #             xls = pd.ExcelFile(excel_path)

# #             for sheet in xls.sheet_names:

# #                 if not sheet_name_is_component(sheet):
# #                     continue

# #                 if sheet_has_required_columns(excel_path, sheet):
# #                     print("[FOUND] CDR Components sheet detected")
# #                     print(f"        File : {filename}")
# #                     print(f"        Sheet: {sheet}")

# #                     return excel_path, sheet

# #         except Exception as e:
# #             print(f"[SKIP] {filename}: {e}")

# #     print("[WARN] No valid CDR Components sheet found")
# #     return None


# def find_cdr_components_excel(directory: str):
#     print("\n🔍 [CDR FINDER] Scanning directory:", directory)

#     if not os.path.isdir(directory):
#         print("❌ Directory does not exist")
#         return None

#     files = os.listdir(directory)
#     print("📂 Files found:", files)

#     for filename in files:

#         print(f"\n➡️ Checking file: {filename}")

#         if filename.startswith("~$"):
#             print("   ⏭️ Skipping temp file")
#             continue

#         if not filename.lower().endswith(".xlsx"):
#             print("   ⏭️ Not an .xlsx file")
#             continue

#         excel_path = os.path.join(directory, filename)
#         print("   📄 Excel candidate:", excel_path)

#         try:
#             xls = pd.ExcelFile(excel_path)
#             print("   📑 Sheets found:", xls.sheet_names)

#             for sheet in xls.sheet_names:
#                 print(f"      🔎 Checking sheet: '{sheet}'")

#                 if not sheet_name_is_component(sheet):
#                     print("         ❌ sheet_name_is_component = False")
#                     continue

#                 print("         ✅ sheet_name_is_component = True")

#                 if not sheet_has_required_columns(excel_path, sheet):
#                     print("         ❌ sheet_has_required_columns = False")
#                     continue

#                 print("         ✅ sheet_has_required_columns = True")
#                 print("[FOUND] CDR Components sheet detected")
#                 print(f"        File : {filename}")
#                 print(f"        Sheet: {sheet}")

#                 return excel_path, sheet

#         except Exception as e:
#             print(f"   ❌ Failed reading Excel: {e}")

#     print("\n⚠️ [WARN] No valid CDR Components sheet found")
#     return None


# import pandas as pd
# def extract_components_table(
#     excel_path: str,
#     sheet_name: str = "4.0 Components",
#     start_keyword: str = "photo"
# ) -> pd.DataFrame:
#     """
#     Extracts a table from an Excel sheet starting at the row
#     containing 'Photo#' (robust match) and stops at the first fully blank row.
#     """

#     # Read entire sheet (no headers)
#     df_raw = pd.read_excel(excel_path, sheet_name=sheet_name, header=None)

#     # -------- FIND HEADER ROW ROBUSTLY --------
#     start_row = None
#     for i, row in df_raw.iterrows():
#         row_str = row.astype(str).str.lower().str.strip()
#         if row_str.str.contains(start_keyword).any():
#             start_row = i
#             break

#     if start_row is None:
#         raise ValueError(
#             f"Could not find a row containing '{start_keyword}' in sheet '{sheet_name}'"
#         )

#     # -------- EXTRACT HEADERS --------
#     headers = (
#         df_raw.iloc[start_row]
#         .astype(str)
#         .str.strip()
#         .replace("nan", "")
#         .tolist()
#     )

#     # -------- EXTRACT DATA --------
#     data = df_raw.iloc[start_row + 1:].copy()
#     data.columns = headers

#     # -------- STOP AT FIRST COMPLETELY BLANK ROW --------
#     blank_rows = data.index[data.isna().all(axis=1)]
#     if not blank_rows.empty:
#         data = data.loc[: blank_rows[0] - 1]

#     # -------- CLEANUP --------
#     data = data.reset_index(drop=True)
#     data = data.dropna(axis=1, how="all")

#     return data 

# from typing import Optional
# import pandas as pd
# from typing import Optional
# import pandas as pd
# import os

# def load_cdr_components_df(
#     directory: str
# ) -> Optional[pd.DataFrame]:
#     """
#     Finds the CDR Components Excel file in the directory
#     and extracts the components table as a DataFrame.

#     Returns:
#         pd.DataFrame if successful
#         None if no valid CDR components Excel is found
#     """

#     print("\n🔍 [CDR EXCEL] Starting Excel-based component extraction")
#     print(f"📂 Directory: {directory}")

#     # -------------------------------
#     # Step 1: Find Excel + sheet
#     # -------------------------------
#     result = find_cdr_components_excel(directory)

#     if result is None:
#         print("❌ [CDR EXCEL] No valid CDR Components Excel found.")
#         print("📄 Files present:", os.listdir(directory))
#         return None

#     excel_path, sheet_name = result

#     print("✅ [CDR EXCEL] Excel file detected")
#     print(f"   📄 File : {excel_path}")
#     print(f"   📑 Sheet: {sheet_name}")

#     # -------------------------------
#     # Step 2: Extract table
#     # -------------------------------
#     try:
#         print("📊 [CDR EXCEL] Extracting components table...")

#         df = extract_components_table(
#             excel_path
#             # sheet_name=sheet_name  # enable if needed
#         )

#         if df is None:
#             print("❌ [CDR EXCEL] extract_components_table returned None")
#             return None

#         print("✅ [CDR EXCEL] Table extracted successfully")
#         print(f"   🔢 Shape: {df.shape}")
#         print("   🏷️ Columns:")
#         for col in df.columns:
#             print(f"      - {col}")

#         # Optional: show first few rows
#         print("   🔍 Preview:")
#         print(df.head(3))

#         return df

#     except Exception as e:
#         print("❌ [CDR EXCEL] Failed during table extraction")
#         print(f"   Reason: {e}")
#         return None
import pandas as pd
from typing import Optional
from pathlib import Path

def load_cdr_components_df(directory: str) -> Optional[pd.DataFrame]:
    """
    Loads CDR Components table from .xlsx / .xls / XML / HTML disguised files.
    Removes NOTES automatically.
    """

    print("\n🔍 [CDR EXCEL] Starting Excel-based component extraction")
    print(f"📂 Directory: {directory}")

    for path in Path(directory).iterdir():

        if path.name.startswith("~$"):
            continue
        if path.suffix.lower() not in (".xlsx", ".xls"):
            continue

        print(f"\n📄 Checking file: {path}")

        tables = None

        try:
            # -------------------------------
            # Explicit .xls handling
            # -------------------------------
            if path.suffix.lower() == ".xls":
                print("   ✅ Detected legacy .xls")
                tables = pd.read_excel(
                    path,
                    sheet_name=None,
                    header=None,
                    engine="xlrd"
                )

            else:
                # -------------------------------
                # Inspect header bytes
                # -------------------------------
                with open(path, "rb") as f:
                    head = f.read(200)

                is_zip = head.startswith(b"PK")
                is_xml = head.lstrip().startswith(b"<?xml")
                is_html = head.lstrip().lower().startswith(b"<html")

                if is_zip:
                    print("   ✅ Detected real XLSX")
                    tables = pd.read_excel(path, sheet_name=None, header=None)

                elif is_xml or is_html:
                    print("   ✅ Detected XML / HTML export")
                    dfs = pd.read_html(path)
                    tables = {f"sheet_{i}": df for i, df in enumerate(dfs)}

                else:
                    print("   ⚠️ Unknown format → trying CSV fallback")
                    df = pd.read_csv(path, header=None)
                    tables = {"csv": df}

        except Exception as e:
            print(f"   ❌ Failed to read file: {e}")
            continue

        # -------------------------------
        # Scan tables for Components header
        # -------------------------------
        for sheet, df_raw in tables.items():
            print(f"   🔎 Scanning: {sheet}")

            if df_raw is None or df_raw.empty:
                continue

            header_row = None
            for i, row in df_raw.iterrows():
                norm = (
                    row.astype(str)
                    .str.lower()
                    .str.replace(" ", "")
                    .str.strip()
                )

                if (
                    norm.str.contains("photo").any()
                    and norm.str.contains("name").any()
                    and norm.str.contains("manufacturer").any()
                    and norm.str.contains("conform").any()
                ):
                    header_row = i
                    break

            if header_row is None:
                continue

            print(f"   ✅ Header detected at row {header_row}")

            # -------------------------------
            # Extract table
            # -------------------------------
            headers = (
                df_raw.iloc[header_row]
                .astype(str)
                .str.strip()
                .replace("nan", "")
                .tolist()
            )

            data = df_raw.iloc[header_row + 1:].copy()
            data.columns = headers

            # Remove NOTES and everything below
            notes_idx = data[
                data.apply(
                    lambda r: r.astype(str)
                    .str.contains(r"^notes:", case=False, na=False)
                    .any(),
                    axis=1
                )
            ].index

            if not notes_idx.empty:
                data = data.loc[: notes_idx[0] - 1]

            data = data.dropna(axis=1, how="all")
            data = data.reset_index(drop=True)

            print("✅ [CDR EXCEL] Components table extracted")
            print(f"   🔢 Shape: {data.shape}")

            return data

    print("\n⚠️ [WARN] No valid CDR Components table found")
    return None


import os
from urllib.parse import unquote

def attach_blob_urls_to_text_support_letter(data, blob_urls):
    """
    Add 'url' to each text_support item by matching base filenames.
    Works with data.pages[].items[].text_support[] structure.
    """

    # Build filename → blob URL map
    blob_map = {}
    for url in blob_urls:
        fname = unquote(url.split("/")[-1])
        base = os.path.splitext(fname)[0]
        blob_map.setdefault(base, url)

    # Traverse updated JSON structure
    for page in data.get("pages", []):
        for item in page.get("items", []):
            for ts in item.get("text_support", []):
                ts["url"] = None  # default

                fname = ts.get("filename")
                if not fname:
                    continue

                base = os.path.splitext(fname)[0]
                ts["url"] = blob_map.get(base)

    return data
 

import os
from urllib.parse import unquote

def attach_blob_urls_to_image_support_letter(data, blob_urls):
    """
    Add 'file_url' to each image_support entry based on pdf_file basename.
    Works with data.pages[].items[].image_support[] structure.
    """

    # Build pdf base filename → blob URL map
    blob_map = {}
    for url in blob_urls:
        fname = unquote(url.split("/")[-1])
        base = os.path.splitext(fname)[0]
        blob_map.setdefault(base, url)

    # Traverse updated JSON structure
    for page in data.get("pages", []):
        for item in page.get("items", []):
            for img in item.get("image_support", []):
                img["file_url"] = None  # default

                pdf_file = img.get("pdf_file")
                if not pdf_file:
                    continue

                base = os.path.splitext(pdf_file)[0]
                img["file_url"] = blob_map.get(base)

    return data
 

from urllib.parse import unquote

def update_non_conforming_urls_from_blob(data, blob_urls):
    """
    Update 'nonConforming_urls' for the item where key == 'photograph'
    using a list of blob URLs.
    """

    nonConformin_urls = [
        {
            "id": idx + 1,
            "url": unquote(url).rstrip("'")
        }
        for idx, url in enumerate(blob_urls)
    ]

    for page in data.get("pages", []):
        for item in page.get("items", []):
            if item.get("key") == "photograph":
                item["nonConforming_urls"] = nonConformin_urls
                return data

    return data
 
