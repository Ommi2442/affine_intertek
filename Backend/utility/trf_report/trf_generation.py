
import re
import tempfile
import shutil
import requests
from urllib.parse import urlparse, unquote
from email import policy
from email.parser import BytesParser
import extract_msg
import uuid
import io
import openpyxl
import xlrd
# from utils import *
from azure.storage.blob import BlobClient
from azure.core.exceptions import ResourceNotFoundError, AzureError
# from templates import *
from utility.trf_report.trf_essential import *
from utility.trf_utils import *
import pandas as pd
import math
import copy
import time
from azure.cosmos import CosmosClient, PartitionKey, exceptions
import json, os
from azure.cosmos import CosmosClient, ConsistencyLevel
from typing import List, Dict, Any, Tuple
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
from langchain_azure_ai.vectorstores import AzureCosmosDBNoSqlVectorSearch
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from azure.cosmos import CosmosClient
from langchain_openai import AzureOpenAIEmbeddings
from operator import itemgetter
from langchain_core.runnables import (
    RunnableParallel, RunnableLambda, RunnableMap
)
from langchain_core.output_parsers import StrOutputParser
from langchain_openai import AzureChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from tenacity import retry, retry_if_exception_type, wait_exponential, stop_never, RetryCallState
from openai import RateLimitError  # Make sure this import exists
from types import SimpleNamespace
from concurrent.futures import ThreadPoolExecutor, as_completed
from azure.core.exceptions import HttpResponseError
import time
from langchain_community.callbacks import get_openai_callback
import os
import re
import json
import time
import shutil
import tempfile
from types import SimpleNamespace
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlparse, unquote
import os
import json
import re
import time
import pandas as pd
from urllib.parse import unquote, urlparse
from types import SimpleNamespace
from concurrent.futures import ThreadPoolExecutor, as_completed

# PDF / Image utilities
from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnableParallel, RunnableMap, RunnableLambda
from operator import itemgetter

# Azure OpenAI
from langchain_openai import AzureChatOpenAI

# Azure Cosmos Vector DB (same as ingest_pipeline)

from langchain_azure_ai.vectorstores import AzureCosmosDBNoSqlVectorSearch

# DOCX utilities
from docx import Document as WordDocument
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from dotenv import load_dotenv

from typing import Callable, Optional

from pathlib import Path

from utility.trf_report.prompts import (
    GUIDELINE_REFERENCE,
    EVIDENCE_BLOCK_TEMPLATE,
    get_grey_mode_prompt,
    get_base_system_prompt,
    get_remark_instruction,
    get_verdict_instruction,
    get_description_instruction,
    get_default_instruction
)


# Local helper module

load_dotenv()

# Chunking config
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150
EMBED_DIM = 1536
VECTOR_PATH = "/vector"

from projects.keyvault_load import *
load_keyvault_secrets()


# Load environment variables
AOAI_ENDPOINT      = os.getenv("aoai-endpoint")
AOAI_KEY           = os.getenv("aoai-key")
API_VERSION        = os.getenv("api-version")
EMBED_DEPLOY       = os.getenv("embed-deploy")
COSMOS_DB_TEXT     = os.getenv("cosmos-db-text")
COSMOS_CONT_TEXT   = os.getenv("cosmos-cont-text")
AZURE_CONN_STRING  = os.getenv("azure-conn-string")
BLOB_CONTAINER     = os.getenv("blob-container")
COSMOS_URL         = os.getenv("cosmos-url")
COSMOS_KEY         = os.getenv("cosmos-key")
CHAT_DEPLOY        = os.getenv("chat-deploy")
BLOB_CONT_NAME     = os.getenv("BLOB_CONT_NAME")
COSMOS_DB_IMAGE    = os.getenv("cosmos-db-image")
COSMOS_CONT_IMAGE  = os.getenv("cosmos-cont-image")
ENABLE_CAD_SCHEMATICS  = os.getenv("enable-cad-schematics")



def build_vectorstore_text(textDB_container_name):
    cosmos_client = CosmosClient(
        url=COSMOS_URL,
        credential=COSMOS_KEY
    )

    return AzureCosmosDBNoSqlVectorSearch(
        cosmos_client=cosmos_client,
        embedding=build_embeddings(AOAI_ENDPOINT,AOAI_KEY,API_VERSION,EMBED_DEPLOY),
        database_name=COSMOS_DB_TEXT,
        container_name=textDB_container_name,

        vector_embedding_policy={
            "vectorEmbeddings": [{
                "path": "/vector",
                "dataType": "float32",
                "dimensions": EMBED_DIM,
                "distanceFunction": "cosine"
            }]
        },
        indexing_policy={
            "includedPaths": [{"path": "/*"}],
            "excludedPaths": [{"path": "/\"_etag\"/?"}, {"path": "/vector/*"}],
            "vectorIndexes": [{"path": "/vector", "type": "quantizedFlat"}],
        },
        cosmos_container_properties={"partition_key": "/id"},
        cosmos_database_properties={},
        vector_search_fields={
            "text_field": "text",
            "embedding_field": "vector",
            "metadata_field": "metadata"
        }
    )

def build_vectorstore_image(imageDB_container_name):
    cosmos_client = CosmosClient(
        url=COSMOS_URL,
        credential=COSMOS_KEY
    )

    return AzureCosmosDBNoSqlVectorSearch(
        cosmos_client=cosmos_client,
        embedding=build_embeddings(AOAI_ENDPOINT,AOAI_KEY,API_VERSION,EMBED_DEPLOY),
        database_name=COSMOS_DB_IMAGE,
        container_name=imageDB_container_name,

        vector_embedding_policy={
            "vectorEmbeddings": [{
                "path": "/vector",
                "dataType": "float32",
                "dimensions": EMBED_DIM,
                "distanceFunction": "cosine"
            }]
        },
        indexing_policy={
            "includedPaths": [{"path": "/*"}],
            "excludedPaths": [{"path": "/\"_etag\"/?"}, {"path": "/vector/*"}],
            "vectorIndexes": [{"path": "/vector", "type": "quantizedFlat"}],
        },
        cosmos_container_properties={"partition_key": "/id"},
        cosmos_database_properties={},
        vector_search_fields={
            "text_field": "text",
            "embedding_field": "vector",
            "metadata_field": "metadata"
        }
    )

# -------------------------------------------------------------------------
# LLM INITIALIZATION (with JSON output enforced)
# -------------------------------------------------------------------------
llm = AzureChatOpenAI(
    azure_endpoint=AOAI_ENDPOINT,
    api_key=AOAI_KEY,
    openai_api_version=API_VERSION,
    azure_deployment=CHAT_DEPLOY,
    temperature=0.1,
).with_config({"response_format": "json_object"})


# =============================================================================
#   VISION PROMPT BUILDER (UPDATED V5 - GREY)
# =============================================================================

# def build_vision_message_grey(inputs, grey=False):
#     """
#     NEW VERSION — Includes correct normalized image filenames so the LLM
#     can output proper image_files in evidence_source.
#     """

#     # ----------------------------------------------------
#     # Extract fields
#     # ----------------------------------------------------
#     question = inputs.get("question", "")
#     docs = inputs.get("context", [])
#     image_urls = inputs.get("image", [])
#     custom_prompt = inputs.get("custom_prompt")
#     task_type = (inputs.get("task_type") or inputs.get("prompt_type") or "extract").lower()
#     accuracy_level = bool(inputs.get("accuracy_level", False))

#     # Build RAG text context
#     context_text = "\n\n".join([d.page_content for d in docs])

#     # ----------------------------------------------------
#     # Collect TEXT file names
#     # ----------------------------------------------------
#     text_files = list({d.metadata.get("source_file", "unknown") for d in docs})

#     # ----------------------------------------------------
#     # Normalize IMAGE filenames for LLM reference
#     # ----------------------------------------------------
    
#     def compute_image_file(img):

#         if isinstance(img, dict):
#             url = img.get("url")
#         else:
#             url = img

#         clean_url = url.split("?")[0]
#         filename = clean_url.split("/")[-1]
#         parts = clean_url.split("/")

#         # PDF extracted PNG
#         if len(parts) >= 2 and parts[-2].lower().endswith(".pdf"):
#             pdf_file = parts[-2]
#             m = re.search(r"page_(\d+)\.png$", filename)
#             page = int(m.group(1)) if m else None
#             return f"{pdf_file}_page_{page}.png"

#         # Standalone PNG
#         return filename



#     image_files = list({compute_image_file(img) for img in image_urls})

#     # ----------------------------------------------------
#     # IEC reference block
#     # ----------------------------------------------------
#     guideline_reference = """
#     You must obey these grounding rules at all times.
#     1. No guessing — If the requirement of a clause is not explicitly available in the TRF text and the provided images, you must NOT infer it.
#     2. However, you can use your judgement and expertise as an electrical safety compliance engineer to give responses.
#     3. IEC 61010-1 must guide interpretation only. Do NOT copy text from the standard.
#     4. Every remark must be evidence-based.
# """

#     # Determine ACTIVE PROMPT
#     active_prompt = custom_prompt.strip() if custom_prompt and str(custom_prompt).strip() else question

#     # ----------------------------------------------------
#     # Strict evidence source rule block
#     # ----------------------------------------------------
#     evidence_block = f"""
# EVIDENCE SOURCE RULES (follow EXACTLY):

# 1. Classify files ONLY by extension:
#    - ".pdf" → text_files
#    - ".jpg" ".jpeg" ".png" → image_files
#    - PDF-derived images like "xxx_page_7.png" ARE IMAGES → image_files

# 2. Do NOT put images in text_files
# 3. Do NOT put PDFs in image_files
# 4. Include ONLY files used as evidence
# 5. Choose ONLY from these:
#    TEXT FILES: {text_files}
#    IMAGE FILES: {image_files}

# Your JSON must be:

# "evidence_source": {{
#      "type": "text" | "image" | "both",
#      "files": {{
#           "text_files": [],
#           "image_files": []
#      }}
# }}
# """

#     # ----------------------------------------------------
#     # GREY MODE PROMPT
#     # ----------------------------------------------------
#     if grey and not accuracy_level:
#         system_prompt = f"""
# You are assisting with determining whether relevant information exists in the TRF content.

# Respond ONLY in this JSON:
# {{
#   "response": "TBD - Info available" | "TBD - No info available",
#   "confidence": <0-100>,
#   "evidence_source": {{
#       "type": "text" | "image" | "both",
#       "files": {{
#           "text_files": [],
#           "image_files": []
#       }}
#   }}
# }}

# {evidence_block}

# TRF TEXT CONTEXT:
# {context_text}

# QUESTION:
# {active_prompt}
# """.strip()

#     else:
#         # -----------------------------
#         # NORMAL MODE — Task-Type logic
#         # -----------------------------
#         if task_type == "remark":
#             instruction_block = f"""
# Provide ONLY:
# 1. "response": concise evidence-driven remark (Max 10 words)
# 2. "confidence": integer 1–100
# 3. "evidence_source": structure shown below
# {evidence_block}

# Remark Rules:
# - If input contains evidence, provide remark based on IEC clause
# - If component/feature not present → say so
# - If information missing → mention unavailable
# - If undecidable → keep blank
# """

#         elif task_type == "verdict":
#             instruction_block = f"""
# Provide ONLY:
# 1. "response": "P" | "N/A" | " "
# 2. "confidence": integer 1–100
# 3. "evidence_source": structure shown below
# {evidence_block}

# Verdict Rules:
# - "P" = Pass when evidence supports IEC clause
# - "N/A" = Not applicable
# - " " = Fail (contradicting evidence)
# """

#         elif task_type == "description":
#             instruction_block = f"""
# Provide ONLY:
# 1. "response": extracted description
# 2. "confidence": integer 1–100
# 3. "evidence_source": structure shown below
# {evidence_block}
# """

#         else:
#             instruction_block = f"""
# Provide ONLY:
# 1. "response": concise answer (max 10 words)
# 2. "confidence": integer 1–100
# 3. "evidence_source": structure shown below
# {evidence_block}
# """

#         # -----------------------------
#         # NORMAL MODE SYSTEM PROMPT
#         # -----------------------------
#         system_prompt = f"""
# You are an expert in IEC 61010-1 safety compliance. Use only TRF text + images. 
# No assumptions allowed.

# {guideline_reference}

# TASK:
# {active_prompt}

# OUTPUT FORMAT:
# {instruction_block}

# TRF TEXT CONTEXT:
# {context_text}
# """.strip()

#     # ----------------------------------------------------
#     # Build vision message (images + text prompt)
#     # ----------------------------------------------------
#     content_blocks = [{"type": "text", "text": system_prompt}]

#     for img in image_urls:
#         if isinstance(img, dict) and "url" in img:
#             content_blocks.append({"type": "image_url", "image_url": {"url": img["url"]}})
#         elif isinstance(img, str):
#             content_blocks.append({"type": "image_url", "image_url": {"url": img}})

#     return [{"role": "user", "content": content_blocks}]

### updated v5-- use this



def build_vision_message_grey(inputs, grey=False):
    """
    NEW VERSION — includes correct normalized image filenames so the LLM
    can output proper evidence_source.image_files.
    """

    # ----------------------------------------------------
    # Extract fields
    # ----------------------------------------------------
    question = inputs.get("question", "")
    docs = inputs.get("context", [])
    image_urls = inputs.get("image", [])
    custom_prompt = inputs.get("custom_prompt")
    task_type = (inputs.get("task_type") or inputs.get("prompt_type") or "extract").lower()
    accuracy_level = bool(inputs.get("accuracy_level", False))

    # Build RAG text context
    context_text = "\n\n".join([d.page_content for d in docs])

    # ----------------------------------------------------
    # Collect TEXT file names
    # ----------------------------------------------------
    text_files = list({d.metadata.get("source_file", "unknown") for d in docs})

    # ----------------------------------------------------
    # NORMALIZE IMAGE FILENAMES SO LLM CAN RETURN THEM
    # ----------------------------------------------------
    # def compute_image_file(img):
    #     """Convert a URL into a normalized image filename:
    #        e.g. Manual.pdf/page_6.png → Manual.pdf_page_6.png
    #     """
    #     if isinstance(img, dict) and "url" in img:
    #         url = img["url"]
    #     else:
    #         url = img  # plain string URL

    #     parts = url.split("/")
    #     pdf_file = parts[-2]                       # e.g. "Manual.pdf"
    #     png_name = parts[-1].split("?")[0]         # e.g. "page_6.png"

    #     # Combine → Manual.pdf_page_6.png
    #     return f"{pdf_file}_{png_name}"

    def compute_image_file(img):
        """
        Convert URL to canonical image filename.
        - PDF-derived images: Manual.pdf/page_6.png → Manual.pdf_page_6.png
        - Plain images: marking_label.JPG → marking_label.JPG
        """

        if isinstance(img, dict) and "url" in img:
            url = img["url"]
        else:
            url = img

        clean_url = url.split("?")[0]
        parts = clean_url.strip("/").split("/")

        filename = parts[-1]  # always real filename

        # ✅ PDF-derived image
        if len(parts) >= 2 and parts[-2].lower().endswith(".pdf"):
            pdf_file = parts[-2]
            return f"{pdf_file}_{filename}"

        # ✅ Plain image
        return filename

    image_files = list({compute_image_file(img) for img in image_urls})

    # Determine ACTIVE PROMPT
    active_prompt = custom_prompt.strip() if custom_prompt and str(custom_prompt).strip() else question

    # ----------------------------------------------------
    # Build evidence block with file lists
    # ----------------------------------------------------
    evidence_block = EVIDENCE_BLOCK_TEMPLATE.format(
        text_files=text_files,
        image_files=image_files
    )

    # ----------------------------------------------------
    # GREY MODE
    # ----------------------------------------------------
    if grey and not accuracy_level:
        system_prompt = get_grey_mode_prompt(
            evidence_block=evidence_block,
            context_text=context_text,
            active_prompt=active_prompt
        ).strip()

    else:
        # -----------------------------
        # TASK TYPE PROMPTING
        # -----------------------------
        if task_type == "remark":
            instruction_block = get_remark_instruction(evidence_block)

        elif task_type == "verdict":
            instruction_block = get_verdict_instruction(evidence_block)

        elif task_type == "description":
            instruction_block = get_description_instruction(evidence_block)

        else:
            instruction_block = get_default_instruction(evidence_block)

        # -----------------------------
        # NORMAL MODE SYSTEM PROMPT
        # -----------------------------
        system_prompt = get_base_system_prompt(
            guideline_reference=GUIDELINE_REFERENCE,
            active_prompt=active_prompt,
            instruction_block=instruction_block,
            context_text=context_text
        ).strip()

    # ----------------------------------------------------
    # BUILD VISION MESSAGE
    # ----------------------------------------------------
    content_blocks = [{"type": "text", "text": system_prompt}]

    for img in image_urls:
        # dict format
        if isinstance(img, dict) and "url" in img:
            content_blocks.append({"type": "image_url", "image_url": {"url": img["url"]}})

        # raw URL
        elif isinstance(img, str):
            content_blocks.append({"type": "image_url", "image_url": {"url": img}})

    return [{"role": "user", "content": content_blocks}]



# =============================================================================
#   ATTACH SUPPORTING REFERENCES (v5 FINAL)
# =============================================================================
# def normalize_image_file_fix(img):
#     if img.get("image_file"):
#         return img["image_file"]

#     url = img.get("url")
#     if not url:
#         return None

#     clean = url.split("?")[0]
#     parts = clean.strip("/").split("/")

#     filename = parts[-1]          # page_44.png
#     if len(parts) >= 2 and parts[-2].lower().endswith(".pdf"):
#         return f"{parts[-2]}_{filename}"

#     return filename


# def attach_supporting_refs_grey(
#         docs,
#         image_urls,
#         llm_output,
#         llm_classifier,   # unused but required for compatibility
#         vs,
#         accuracy_level=None,
#         grey=False,
#         question=None,
#         k=5):

#     import json
#     import re

#     # ---------------------------------------------------
#     # STEP 0 — Robust JSON extraction
#     # ---------------------------------------------------
#     def extract_json_from_text(text):
#         matches = re.findall(r"\{[\s\S]*\}", text)
#         for block in reversed(matches):
#             try:
#                 return json.loads(block)
#             except Exception:
#                 continue
#         return None

#     parsed = extract_json_from_text(llm_output) or {
#         "response": llm_output,
#         "confidence": 0,
#         "evidence_source": {"type": "text", "files": {}}
#     }

#     evidence = parsed.get("evidence_source", {})
#     source_type = evidence.get("type", "text").lower()
#     evidence_files = evidence.get("files", {})
#     image_files_needed = set(evidence_files.get("image_files", [])) if isinstance(evidence_files, dict) else set()

#     # ---------------------------------------------------
#     # STEP 2 — TEXT SUPPORT USING VECTOR STORE (CORRECT)
#     # ---------------------------------------------------
#     similarity_query = question.strip() if question else parsed.get("response", "")

#     try:
#         scored_results = vs.similarity_search_with_score(similarity_query, k=k)
#     except Exception as e:
#         print("Vector search failed:", str(e))
#         scored_results = []

#     # Full chunk, real similarity score
#     text_support = [
#         {
#             "filename": doc.metadata.get("source_file", "unknown"),
#             "pdf_file": doc.metadata.get("source_file", "unknown"),
#             "page": doc.metadata.get("page"),
#             "similarity_score": float(score),
#             "preview_text": doc.page_content  # FULL CHUNK — no truncation
#         }
#         for doc, score in scored_results
#     ]

#     # ---------------------------------------------------
#     # STEP 3 — IMAGE SUPPORT (unchanged)
#     # ---------------------------------------------------
#     image_support = []

#     if source_type in ("image", "both") and image_files_needed:

#         image_lookup = {}

#         for img in image_urls or []:
#             if isinstance(img, dict) and "url" in img:
#                 # image_file = img.get("image_file") or img["url"].split("?")[0].split("/")[-1]
#                 image_file = normalize_image_file_fix(img)
#                 image_lookup[image_file] = {
#                     "url": img.get("url"),
#                     "image_file": image_file,
#                     "pdf_file": img.get("pdf_file"),
#                     "page": img.get("page")
#                 }
#             elif isinstance(img, str):
#                 image_file = img.split("?")[0].split("/")[-1]
#                 image_lookup[image_file] = {
#                     "url": img,
#                     "image_file": image_file,
#                     "pdf_file": None,
#                     "page": None
#                 }

#         for image_file in image_files_needed:
#             if image_file in image_lookup:
#                 image_support.append(image_lookup[image_file])

#     # ---------------------------------------------------
#     # FINAL structured output
#     # ---------------------------------------------------
#     return {
#         "answer": llm_output.strip(),
#         "confidence": parsed.get("confidence", 0),

#         "evidence_source": {
#             "type": source_type,
#             "files": {
#                 "image_files": list(image_files_needed)
#             }
#         },

#         "text_support": text_support,
#         "image_support": image_support,

#         "grey_mode": grey,
#         "accuracy_level": accuracy_level,
#         "similarity_query": question,
#         "vector_search_based_on": "question"
#     }

def attach_supporting_refs_grey(
    docs,
    image_urls,
    llm_output,
    llm_classifier,   
    vs,
    accuracy_level=None,
    grey=False,
    question=None,
    k=5
):
    import json
    import re

    # ---------------------------------------------------
    # Helper: canonical filename (CRITICAL)
    # ---------------------------------------------------
    def canonical_image_key(name):
        if not name or not isinstance(name, str):
            return None
        name = name.strip()
        name = name.replace("\\", "/")
        name = name.split("/")[-1]
        name = name.replace("%20", "_")
        name = name.replace(" ", "_")
        return name.lower()

    # ---------------------------------------------------
    # STEP 0 — Robust JSON extraction
    # ---------------------------------------------------
    def extract_json_from_text(text):
        if not isinstance(text, str):
            return None
        matches = re.findall(r"\{[\s\S]*\}", text)
        for block in reversed(matches):
            try:
                return json.loads(block)
            except Exception:
                continue
        return None

    parsed = extract_json_from_text(llm_output) or {
        "response": llm_output,
        "confidence": 0,
        "evidence_source": {"type": "text", "files": {}}
    }

    evidence = parsed.get("evidence_source", {}) or {}
    source_type = str(evidence.get("type", "text")).lower()
    files = evidence.get("files") if isinstance(evidence.get("files"), dict) else {}
    image_files_needed = set(files.get("image_files", []))

    # ---------------------------------------------------
    # STEP 1 — TEXT SUPPORT (UNCHANGED)
    # ---------------------------------------------------
    similarity_query = question.strip() if question else parsed.get("response", "")

    try:
        scored_results = vs.similarity_search_with_score(similarity_query, k=k)
    except Exception:
        scored_results = []

    text_support = [
        {
            "filename": doc.metadata.get("source_file", "unknown"),
            "pdf_file": doc.metadata.get("source_file", "unknown"),
            "page": doc.metadata.get("page"),
            "similarity_score": float(score),
            "preview_text": doc.page_content
        }
        for doc, score in scored_results
    ]

    # ---------------------------------------------------
    # STEP 2 — IMAGE SUPPORT (STRICT, SAFE)
    # ---------------------------------------------------
    image_support = []

    if source_type in ("image", "both") and image_files_needed:

        # Build lookup ONLY from valid top-5 images
        image_lookup = {}

        for img in image_urls or []:
            if not isinstance(img, dict):
                continue

            image_file = img.get("image_file")
            url = img.get("url")

            
            if not image_file or not url:
                continue

            key = canonical_image_key(image_file)
            if not key:
                continue

            image_lookup[key] = {
                "filename": image_file,
                "image_file": image_file,
                "pdf_file": img.get("pdf_file"),
                "page": img.get("page"),
                "url": url
            }

        # Resolve ONLY LLM-cited images
        for image_file in image_files_needed:
            key = canonical_image_key(image_file)
            if not key:
                continue

            entry = image_lookup.get(key)
            if not entry:
                continue

            image_support.append(entry)

    # ---------------------------------------------------
    # FINAL OUTPUT (NO NULL URLs POSSIBLE)
    # ---------------------------------------------------
    valid_image_files = {img["image_file"] for img in image_support}

    return {
        "answer": llm_output.strip() if isinstance(llm_output, str) else "",
        "confidence": parsed.get("confidence", 0),

        "evidence_source": {
            "type": source_type,
            "files": {
                "image_files": sorted(valid_image_files)
            }
        },

        "text_support": text_support,
        "image_support": image_support,

        "grey_mode": grey,
        "accuracy_level": accuracy_level,
        "similarity_query": question,
        "vector_search_based_on": "question"
    }


def build_rag_image_pipeline_grey(
        retriever,
        llm,
        build_vision_message,
        attach_supporting_refs,
        vs
    ):

    # Normalize retriever output (list[Document])
    def extract_docs(ctx):
        if isinstance(ctx, list):
            return ctx
        if isinstance(ctx, dict):
            for key in ("context", "documents", "docs", "retrieved_docs"):
                if key in ctx and isinstance(ctx[key], list):
                    return ctx[key]
        return []

    rag_image = (

        # ----------------------------------------------------
        # STEP 1 — Retrieve context + passthrough fields
        # ----------------------------------------------------
        RunnableParallel(
            retrieved_docs=itemgetter("question") | retriever,   # PATCHED
            question=itemgetter("question"),
            image=itemgetter("image"),
            task_type=itemgetter("task_type"),
            custom_prompt=itemgetter("custom_prompt"),
            accuracy_level=itemgetter("accuracy_level"),
            grey=itemgetter("grey"),
        )

        # ----------------------------------------------------
        # STEP 2 — Normalize context + Build Vision Messages
        # ----------------------------------------------------
        | RunnableMap({
            "inputs": RunnableLambda(lambda x: {
                **x,
                "context": extract_docs(x.get("retrieved_docs")),  # PATCHED
                "retrieved_docs": extract_docs(x.get("retrieved_docs"))
            }),
            "messages": RunnableLambda(
                lambda x: build_vision_message(
                    {
                        **x,
                        "context": x.get("retrieved_docs", [])         # PATCHED
                    },
                    grey=x.get("grey", False)
                )
            )
        })

        # ----------------------------------------------------
        # STEP 3 — LLM CALL (RAW STRING)
        # ----------------------------------------------------
        | RunnableMap({
            "inputs": itemgetter("inputs"),
            "raw_llm_output": (
                itemgetter("messages")
                | llm
                | StrOutputParser()
            )
        })

        # ----------------------------------------------------
        # STEP 4 — Attach Supporting References
        # ----------------------------------------------------
        | RunnableLambda(
            lambda x: {

                **attach_supporting_refs(
                    docs=x["inputs"].get("retrieved_docs", []),  # PATCHED
                    image_urls=x["inputs"].get("image", []),
                    llm_output=x.get("raw_llm_output", ""),
                    llm_classifier=None,
                    vs=vs,
                    accuracy_level=x["inputs"].get("accuracy_level"),
                    grey=x["inputs"].get("grey", False),
                    question=x["inputs"].get("question"),
                    k=5
                ),

                "metadata": {
                    "question": x["inputs"].get("question"),
                    "task_type": x["inputs"].get("task_type"),
                    "custom_prompt": x["inputs"].get("custom_prompt"),
                    "accuracy_level": x["inputs"].get("accuracy_level"),
                    "grey": x["inputs"].get("grey"),
                },

                "raw_llm_output": x.get("raw_llm_output"),
            }
        )
    )

    return rag_image



# def build_embeddings(AOAI_ENDPOINT,AOAI_KEY,API_VERSION,EMBED_DEPLOY):
#     return AzureOpenAIEmbeddings(
#         azure_endpoint=AOAI_ENDPOINT,
#         api_key=AOAI_KEY,
#         openai_api_version=API_VERSION,
#         azure_deployment=EMBED_DEPLOY,
#     )

# vs=build_vectorstore_text()
# vs2=build_vectorstore_image()

# retriever = vs.as_retriever(search_kwargs={"k": 5})
# image_retriever_agent = vs2.as_retriever(search_kwargs={"k": 5})

# # Build the final RAG pipeline
# rag_image = build_rag_image_pipeline_grey(
#     retriever,
#     llm,
#     build_vision_message_grey,
#     attach_supporting_refs_grey,
#     vs
# )


# =============================================================================
#   JSON EXTRACTION + UPDATE JSON ITEM (v5.1)
# =============================================================================

import json
import re


def extract_human_prefix(text):
    if not isinstance(text, str):
        return ""
    return text.split("{", 1)[0].strip()


def extract_json_block(text):
    """
    Robust extraction: finds LAST valid JSON object in the string.
    """
    if not isinstance(text, str):
        return None

    matches = re.findall(r"\{[\s\S]*\}", text)
    for block in reversed(matches):
        try:
            return json.loads(block)
        except Exception:
            continue
    return None


def extract_pdf_page_from_filename(fname):
    m = re.search(r"(.+\.pdf)_page_(\d+)\.png$", fname)
    if not m:
        return None, None
    return m.group(1), int(m.group(2))


# def update_json_item_grey(item, result):

#     if not item.get("ai_fillable", False):
#         return

#     answer_text = result.get("answer", "")
#     raw_json = result.get("llm_raw_json")

#     # -----------------------------
#     # 1. Parse JSON safely
#     # -----------------------------
#     if raw_json:
#         try:
#             parsed = json.loads(raw_json)
#         except:
#             parsed = {}
#     else:
#         parsed = extract_json_block(answer_text) or {}

#     # Human prefix override
#     human_prefix = extract_human_prefix(answer_text)
#     if human_prefix and len(human_prefix.split()) > 3:
#         parsed["response"] = human_prefix

#     response = str(parsed.get("response", "")).strip()
#     confidence = parsed.get("confidence", result.get("confidence", 0))

#     # -----------------------------
#     # 2. Grey skip
#     # -----------------------------
#     if response.replace(" ", "").startswith("TBD"):
#         item["value"] = response
#         item["confidence"] = confidence
#         return

#     # -----------------------------
#     # 3. Assign value
#     # -----------------------------
#     if item.get("task_type") == "verdict":
#         item["value"] = normalize_verdict(response)
#     else:
#         item["value"] = response

#     # -----------------------------
#     # 4. SAFE extract image files from LLM JSON
#     # -----------------------------
#     evidence = parsed.get("evidence_source", {}) or {}
#     files = evidence.get("files", {})

#     if not isinstance(files, dict):
#         files = {"text_files": [], "image_files": []}

#     json_image_files = files.get("image_files") or []

#     # -----------------------------
#     # 5. Build text_support ALWAYS 5 hits
#     # -----------------------------
#     similarity_preview = (
#         result.get("similarity_preview")
#         or result.get("text_support")
#         or []
#     )

#     text_support_final = [
#         {
#             "filename": hit.get("filename"),
#             "page": hit.get("page"),
#             "similarity_score": hit.get("similarity_score"),
#             "preview_text": hit.get("preview_text")
#         }
#         for hit in similarity_preview
#     ]

#     # -----------------------------
#     # 6. Build IMAGE SUPPORT (FULL FIX)
#     # -----------------------------
#     image_support_final = []

#     pipeline_image_support = (
#         result.get("image_support")
#         or result.get("pipeline_output", {}).get("image_support")
#         or []
#     )

#     pipeline_lookup = {
#         d.get("image_file"): d for d in pipeline_image_support
#     }

#     task_images = item.get("image", []) or item.get("images", []) or []
#     task_lookup = {d.get("image_file"): d for d in task_images}

#     for fname in json_image_files:

#         pdf_file, page = extract_pdf_page_from_filename(fname)

#         enriched = {
#             "filename": fname,
#             "image_file": fname,
#             "pdf_file": pdf_file,
#             "page": page,
#             "url": None
#         }

#         if fname in pipeline_lookup:
#             enriched.update({
#                 k: v for k, v in pipeline_lookup[fname].items()
#                 if v is not None
#             })

#         if enriched["url"] is None and fname in task_lookup:
#             enriched.update({
#                 k: v for k, v in task_lookup[fname].items()
#                 if v is not None
#             })

#         image_support_final.append(enriched)

#     if not image_support_final:
#         image_support_final = list(pipeline_image_support)

#     # -----------------------------
#     # 7. Final assignment
#     # -----------------------------
#     item["text_support"] = text_support_final
#     item["image_support"] = image_support_final
#     item["evidence_type"] = evidence.get("type") or "text"
#     item["confidence"] = confidence

def canonical_image_key(name):
    if not name or not isinstance(name, str):
        return None
    name = name.strip()
    name = name.replace("\\", "/")
    name = name.split("/")[-1]
    name = name.replace("%20", "_")
    name = name.replace(" ", "_")
    return name.lower()


def update_json_item_grey(item, result):

    if not item.get("ai_fillable", False):
        return

    answer_text = result.get("answer", "")
    raw_json = result.get("llm_raw_json")

    # -----------------------------
    # 1. Parse JSON safely
    # -----------------------------
    if raw_json:
        try:
            parsed = json.loads(raw_json)
        except:
            parsed = {}
    else:
        parsed = extract_json_block(answer_text) or {}

    # Human prefix override
    human_prefix = extract_human_prefix(answer_text)
    if human_prefix and len(human_prefix.split()) > 3:
        parsed["response"] = human_prefix

    response = str(parsed.get("response", "")).strip()
    confidence = parsed.get("confidence", result.get("confidence", 0))

    # -----------------------------
    # 2. Grey skip
    # -----------------------------
    if response.replace(" ", "").startswith("TBD"):
        item["value"] = response
        item["confidence"] = confidence
        return

    # -----------------------------
    # 3. Assign value
    # -----------------------------
    if item.get("task_type") == "verdict":
        item["value"] = normalize_verdict(response)
    else:
        item["value"] = response

    # -----------------------------
    # 4. SAFE extract image files from LLM JSON
    # -----------------------------
    evidence = parsed.get("evidence_source", {}) or {}
    files = evidence.get("files", {})

    if not isinstance(files, dict):
        files = {"text_files": [], "image_files": []}

    json_image_files = files.get("image_files") or []

    # -----------------------------
    # 5. Build text_support ALWAYS 5 hits
    # -----------------------------
    similarity_preview = (
        result.get("similarity_preview")
        or result.get("text_support")
        or []
    )

    text_support_final = [
        {
            "filename": hit.get("filename"),
            "page": hit.get("page"),
            "similarity_score": hit.get("similarity_score"),
            "preview_text": hit.get("preview_text")
        }
        for hit in similarity_preview
    ]

    # -----------------------------
    # 6. Build IMAGE SUPPORT (FULL FIX)
    # -----------------------------
    image_support_final = []

    pipeline_image_support = (
        result.get("image_support")
        or result.get("pipeline_output", {}).get("image_support")
        or []
    )


    pipeline_lookup = {
        canonical_image_key(d.get("image_file")): d
        for d in pipeline_image_support
        if d.get("image_file") and d.get("url")
    }


    task_images = item.get("image", []) or item.get("images", []) or []
    # task_lookup = {d.get("image_file"): d for d in task_images}
    task_lookup = {
        canonical_image_key(d.get("image_file")): d
        for d in task_images
        if d.get("image_file") and d.get("url")
    }


    image_support_final = []

    for fname in json_image_files:

        key = canonical_image_key(fname)
        if not key:
            continue

        enriched = None

        # Priority 1: pipeline output (attach_supporting_refs_grey)
        if key in pipeline_lookup:
            enriched = pipeline_lookup[key]

        # Priority 2: task top-5 images (vs2)
        elif key in task_lookup:
            enriched = task_lookup[key]

        if not enriched or not enriched.get("url"):
            continue

        image_support_final.append({
            "filename": fname,
            "image_file": fname,
            "pdf_file": enriched.get("pdf_file"),
            "page": enriched.get("page"),
            "url": enriched.get("url"),
        })


    # -----------------------------
    # 7. Final assignment
    # -----------------------------
    item["text_support"] = text_support_final
    item["image_support"] = image_support_final
    item["evidence_type"] = evidence.get("type") or "text"
    item["confidence"] = confidence


# =============================================================================
#   NORMALIZE VERDICT
# =============================================================================

def normalize_verdict(value):
    if not isinstance(value, str):
        return "N/A"

    v = value.strip().lower()

    if v == "":
        return " "   # FAIL = blank

    if v.startswith("p"):
        return "P"
    if v.startswith("f"):
        return "F"

    return "N/A"


# =============================================================================
#   BUILD TASKS WITH CUSTOM PROMPTS (v5)
# =============================================================================

# def build_tasks_with_custom_prompt_grey(data, image_urls):
#     tasks = []
#     item_refs = []

#     # Normalize images (string + dict)
#     normalized_images = []
#     for img in image_urls:
#         if isinstance(img, dict):
#             image_file = (img.get("image_file") or "").split("?")[0]
#             normalized_images.append({
#                 "pdf_file": img.get("pdf_file"),
#                 "page": img.get("page"),
#                 "image_file": image_file,
#                 "url": img.get("url")
#             })
#             continue

#         raw = img.split("/")[-1]
#         clean = raw.split("?")[0]
#         normalized_images.append({
#             "pdf_file": None,
#             "page": None,
#             "image_file": clean,
#             "url": img
#         })

#     # Build tasks
#     for table in data.get("Tables", []):
#         for item in table.get("Items", []):

#             if not item.get("ai_fillable"):
#                 continue

#             custom = item.get("custom_prompt", "")
#             accuracy = bool(item.get("accuracy_level", False))

#             if "grey" in item:
#                 grey_flag = bool(item.get("grey"))
#             else:
#                 grey_flag = not accuracy

#             task_type = str(item.get("task_type", "extract")).strip().lower()
#             field = item.get("field") or ""

#             if isinstance(custom, str) and custom.strip():
#                 modified_question = custom.strip()

#             else:
#                 if task_type == "verdict":
#                     modified_question = (
#                         "Give the verdict as Pass, Fail, or N/A only.\n"
#                         f"{field}"
#                     )
#                 else:
#                     modified_question = field

#             tasks.append({
#                 "question": modified_question,
#                 "image": list(normalized_images),
#                 "task_type": task_type,
#                 "custom_prompt": custom,
#                 "accuracy_level": accuracy,
#                 "grey": grey_flag
#             })

#             item_refs.append(item)

#     return tasks, item_refs

def build_tasks_with_custom_prompt_grey(data, image_urls):
    tasks = []
    item_refs = []

    # --------------------------------------------------------
    # NORMALIZE image URLs
    # --------------------------------------------------------
    normalized_images = []
    for img in image_urls:
        if isinstance(img, dict):
            image_file = (img.get("image_file") or "").split("?")[0]
            normalized_images.append({
                "pdf_file": img.get("pdf_file"),
                "page": img.get("page"),
                "image_file": image_file,
                "url": img.get("url")
            })
        else:
            raw = img.split("/")[-1]
            clean = raw.split("?")[0]
            normalized_images.append({
                "pdf_file": None,
                "page": None,
                "image_file": clean,
                "url": img
            })

    # --------------------------------------------------------
    # Process all table items
    # --------------------------------------------------------
    for table in data.get("Tables", []):
        for item in table.get("Items", []):

            # ✅ FIX HERE
            # if not (item.get("ai_fillable") or item.get("task_type") == "remark"):
            if item.get("ai_fillable") is not True:
                continue

            custom = item.get("custom_prompt", "")
            accuracy = bool(item.get("accuracy_level", False))

            grey_flag = bool(item.get("grey")) if "grey" in item else not accuracy

            task_type = str(item.get("task_type", "extract")).strip().lower()
            field = item.get("field") or ""

            if isinstance(custom, str) and custom.strip():
                modified_question = custom.strip()
            else:
                if task_type == "verdict":
                    modified_question = (
                        "Give the verdict as Pass, Fail, or N/A only.\n"
                        f"{field}"
                    )
                else:
                    modified_question = field

            tasks.append({
                "question": modified_question,
                "image": list(normalized_images),
                "task_type": task_type,
                "custom_prompt": custom,
                "accuracy_level": accuracy,
                "grey": grey_flag
            })

            item_refs.append(item)

    return tasks, item_refs
 


# =============================================================================
#   AGENT TOOL — run_single_task_tool (Option 1)
# =============================================================================

# from langchain_core.tools import tool

# def generator_evaluator_agent(rag_image):

#     @tool("run_single_task")
#     def run_single_task_tool(task: dict):
#         """
#         Runs a single RAG image task with retry handling.
#         This calls your run_single_task from trf_Essentials.
#         """
#         return run_single_task(task, rag_image)

#     return run_single_task_tool


# # Create the tool instance for a given rag_image pipeline
# run_single_task_tool = generator_evaluator_agent(rag_image)

from langchain_core.tools import tool

def generator_evaluator_agent(rag_image):

    @tool("run_single_task")
    def run_single_task_tool(task: dict):
        """
        Runs a single RAG image task with retry handling.
        """
        return run_single_task(task, rag_image)

    return run_single_task_tool


# run_single_task_tool = generator_evaluator_agent(rag_image)

# def generator_evaluator_agent_tool(task):
#     return run_single_task_tool.invoke({"task": task})

def generator_evaluator_agent_tool(run_single_task_tool, task):
    return run_single_task_tool.invoke({"task": task})




# =============================================================================
#   PROCESS TASKS IN BATCHES WITH PARALLELISM (v5.1)
# =============================================================================

def process_tasks_with_batches_parallel_grey(
        tasks,
        item_refs,
        rag_image,
        vs,
        run_single_task_tool,
        batch_size=150,
        cooldown_sec=15,
        max_workers=6,
        use_llm_inGrey=False,
        stats=False
    ):

    total = len(tasks)
    processed = 0
    all_results = []

    total_prompt_tokens = 0
    total_completion_tokens = 0
    total_total_tokens = 0
    total_llm_calls = 0

    # -----------------------------------------------------
    # GREY fallback handler (vector-only)
    # -----------------------------------------------------
    def handle_low_accuracy_item(item, task, vs, k=5):

        if use_llm_inGrey:
            return False

        if task.get("accuracy_level", True):
            return False

        question = task["question"]

        try:
            scored_results = vs.similarity_search_with_score(question, k=k)
        except Exception:
            scored_results = []

        text_sources = [
            {
                "filename": doc.metadata.get("source_file", "unknown"),
                "page": doc.metadata.get("page"),
                "similarity_score": float(score),
                "preview_text": doc.page_content[:500].strip(),
            }
            for doc, score in scored_results
        ]

        has_text_info = any(score >= 0.79 for _, score in scored_results)

        question_l = question.lower()
        image_keywords = ["label", "marking", "diagram", "schematic", "figure", "block", "display"]
        has_image_hint = any(kw in question_l for kw in image_keywords)

        has_info = has_text_info or has_image_hint
        value = "TBD-Info available" if has_info else "TBD-Info not available"

        item.update({
            "value": value,
            "confidence": 0,
            "text_support": text_sources,
            "image_support": [],
            "similarity_query": question,
            "vector_search_based_on": "question",
            "evidence_source": {
                "type": "text" if text_sources else "none",
                "files": {
                    "text_files": sorted({x["filename"] for x in text_sources}),
                    "image_files": []
                }
            }
        })

        all_results.append(item)
        return True

    # -----------------------------------------------------
    # MAIN BATCH EXECUTION LOOP
    # -----------------------------------------------------
    for start in range(0, total, batch_size):
        end = min(start + batch_size, total)
        batch = tasks[start:end]
        ref_batch = item_refs[start:end]

        print(f"\n🔵 Starting batch: {start+1} → {end}")

        with ThreadPoolExecutor(max_workers=max_workers) as exe:
            futures = {}

            for idx, task in enumerate(batch):
                item = ref_batch[idx]

                if handle_low_accuracy_item(item, task, vs):
                    processed += 1
                    print(f"Processed {processed}/{total} (vector-only GREY)")
                    continue

                if stats:
                    fut = exe.submit(run_single_task_stats, task, rag_image)
                else:
                    # fut = exe.submit(generator_evaluator_agent_tool, task)
                    fut = exe.submit(generator_evaluator_agent_tool, run_single_task_tool, task)

                futures[fut] = idx

            for future in as_completed(futures):
                idx = futures[future]
                task = batch[idx]
                item = ref_batch[idx]

                try:
                    result = future.result()

                except Exception as e:

                    print(f"❌ LLM ERROR for task {idx+1}: {e}")
                    print("➡️  Using blank fallback and continuing...\n")

                    # Build a blank result so all later code works correctly
                    result = {
                        "answer": "",
                        "confidence": 0,
                        "evidence_source": {
                            "type": "none",
                            "files": {"text_files": [], "image_files": []},
                        },
                        "text_support": [],
                        "image_support": [],
                        "raw_llm_output": "",
                        "metadata": {
                            "question": task.get("question"),
                            "task_type": task.get("task_type"),
                            "custom_prompt": task.get("custom_prompt"),
                            "accuracy_level": task.get("accuracy_level"),
                            "grey": task.get("grey"),
                        },
                    }

                if stats and "_token_usage" in result:
                    usage = result["_token_usage"]
                    total_prompt_tokens += usage.get("prompt", 0)
                    total_completion_tokens += usage.get("completion", 0)
                    total_total_tokens += usage.get("total", 0)
                    total_llm_calls += 1

                result["question"] = task["question"]
                result["task_type"] = task.get("task_type")

                update_json_item_grey(item, result)

                all_results.append(result)
                processed += 1
                print(f"Processed {processed}/{total}")

        if end < total:
            print(f"⏳ Cooling down for {cooldown_sec} seconds...")
            time.sleep(cooldown_sec)

    print("\n✅ ALL TASKS COMPLETED.")

    if stats:
        return {
            "results": all_results,
            "stats": {
                "total_llm_calls": total_llm_calls,
                "total_prompt_tokens": total_prompt_tokens,
                "total_completion_tokens": total_completion_tokens,
                "total_tokens": total_total_tokens,
            }
        }

    return all_results


# =============================================================================
#   BUILD TOP-5 IMAGES PER TASK FROM IMAGE VECTORSTORE (vs2)
# =============================================================================

import os
import re
from urllib.parse import urlparse, unquote
from concurrent.futures import ThreadPoolExecutor, as_completed


def extract_pdf_and_page_from_blob_url(blob_url: str):
    """
    Extract:
    - pdf_file: folder name ending with .pdf
    - image_file: pdf_file_page_x.png (UNDERSCORE FORMAT)
    - page: extracted page number
    """

    parsed = urlparse(blob_url)
    path = parsed.path

    parts = path.split("/")

    pdf_file = None
    for part in parts:
        if part.lower().endswith(".pdf"):
            pdf_file = part
            break

    image_name = os.path.basename(path)
    m = re.search(r"page_(\d+)\.png$", image_name)
    page = int(m.group(1)) if m else None

    if pdf_file:
        image_file = f"{pdf_file}/{image_name}"
        image_file = image_file.replace("/", "_")
    else:
        image_file = image_name

    return pdf_file, page, image_file


def build_image_entry(blob_url: str):

    clean_url = blob_url.split("?")[0]
    filename = clean_url.split("/")[-1]
    parts = clean_url.split("/")

    # CASE 1: PDF extracted PNG
    if len(parts) >= 2 and parts[-2].lower().endswith(".pdf"):
        pdf_file = parts[-2]
        m = re.search(r"page_(\d+)\.png$", filename)
        page = int(m.group(1)) if m else None

        canonical_image_file = f"{pdf_file}_page_{page}.png"

        return {
            "pdf_file": pdf_file,
            "page": page,
            "image_file": canonical_image_file,
            "url": blob_url
        }

    # CASE 2: Standalone PNG
    canonical_image_file = filename  # unchanged

    return {
        "pdf_file": None,
        "page": None,
        "image_file": canonical_image_file,
        "url": blob_url
    }





def process_single_task(task, retriever, task_index, total_tasks):
    """
    Processes ONE task:
    - Run text-based retriever (vs2)
    - Extract 5 nearest image entries
    """
    question = task.get("custom_prompt")
    print(f"[INFO] Processing task {task_index}/{total_tasks} → {question}")

    neighbors = retriever.invoke(question)

    images = []
    for d in neighbors[:5]:
        blob_url = d.metadata.get("blob_url")
        entry = build_image_entry(blob_url)
        images.append(entry)

    task_new = dict(task)
    task_new["image"] = images

    print(f"[SUCCESS] Completed task {task_index}/{total_tasks}")

    return task_new


def update_tasks_with_top5_images(tasks, retriever, max_threads=10):
    total = len(tasks)
    updated = [None] * total

    print(f"[START] Processing {total} tasks with {max_threads} parallel threads.\n")

    with ThreadPoolExecutor(max_workers=max_threads) as executor:

        future_map = {
            executor.submit(process_single_task, task, retriever, idx + 1, total): idx
            for idx, task in enumerate(tasks)
        }

        for future in as_completed(future_map):
            idx = future_map[future]
            try:
                updated[idx] = future.result()
            except Exception as e:
                print(f"[ERROR] Task {idx+1} failed: {e}")

    print(f"\n[COMPLETE] Completed all {total} tasks.\n")

    return updated


# =============================================================================
#   VERDICT DEPENDENCY UPDATE
# =============================================================================

def update_verdict_dependencies(data):
    """
    Update all items with task_type = 'verdict_dependency'
    based on values of dependency_rows.
    """

    print("\n🔍 Updating verdict_dependency items...\n")

    for table in data["Tables"]:
        items = table["Items"]

        verdict_map = {}
        for item in items:
            if item["answer_column"] == 3:
                verdict_map[item["answer_row"]] = item.get("value")

        for item in items:
            if item.get("task_type") != "verdict_dependency":
                continue

            dep_rows = item.get("dependency_rows")
            if not dep_rows:
                continue

            dep_vals = [verdict_map.get(row) for row in dep_rows if row in verdict_map]
            dep_vals = [v if v is not None else "" for v in dep_vals]

            if any(v == "" for v in dep_vals):
                final_val = ""
            elif any(v == "F" for v in dep_vals):
                final_val = "F"
            elif all(v == "N/A" for v in dep_vals):
                final_val = "N/A"
            elif all(v in ("P", "N/A") for v in dep_vals):
                final_val = "P"
            else:
                final_val = ""

            old_val = item.get("value")
            item["value"] = final_val

            print(f"Updated: {item['field']} → {final_val} (was: {old_val})")

    print("\n✅ Verdict dependency update complete.\n")


# =============================================================================
#   PREFIXING FUNCTIONS
# =============================================================================

def apply_prefixes_to_items(data, targets):
    """
    Adds prefixes to data['Tables'][8]['Items'].
    """

    items = data["Tables"][8]["Items"]

    if len(items) < len(targets):
        raise ValueError("Items list has fewer elements than targets")

    for i, prefix in enumerate(targets):
        item = items[i]
        old_value = item.get("value", "")
        item["value"] = prefix + old_value

    return data


def prefix_summary_of_compliance(data):
    """
    Prefix to data['Tables'][5]['Items'][5]['value'].
    """
    prefix = "Summary of compliance with National Differences (List of countries addressed): "

    target_item = data["Tables"][5]["Items"][5]
    old_value = target_item.get("value", "")
    target_item["value"] = prefix + old_value

    return data


# =============================================================================
#   ADD BLOB URLS TO SUPPORT SECTIONS
# =============================================================================

def attach_blob_urls_to_text_support(data, blob_urls):
    """
    Add 'url' to each text_support item by matching base filenames.
    """

    blob_map = {}
    for url in blob_urls:
        fname = unquote(url.split("/")[-1])
        base = os.path.splitext(fname)[0]
        blob_map.setdefault(base, url)

    for table in data.get("Tables", []):
        for item in table.get("Items", []):
            for ts in item.get("text_support", []):
                ts["url"] = None
                fname = ts.get("filename")
                if not fname:
                    continue

                base = os.path.splitext(fname)[0]
                ts["url"] = blob_map.get(base)

    return data


def attach_blob_urls_to_image_support(data, blob_urls):
    """
    Add 'file_url' to each image_support entry based on pdf_file basename.
    """

    blob_map = {}
    for url in blob_urls:
        fname = unquote(url.split("/")[-1])
        base = os.path.splitext(fname)[0]
        blob_map.setdefault(base, url)

    for table in data.get("Tables", []):
        for item in table.get("Items", []):
            for img in item.get("image_support", []):
                img["file_url"] = None
                pdf_file = img.get("pdf_file")
                if not pdf_file:
                    continue

                base = os.path.splitext(pdf_file)[0]
                img["file_url"] = blob_map.get(base)

    return data


# =============================================================================
#   EXPORT: DATAFRAME → EXCEL
# =============================================================================

def export_results_to_excel(df, excel_output_path):
    """
    Saves the final DataFrame into an Excel file.
    """
    df.to_excel(excel_output_path, index=False)
    print(f"✅ Excel saved → {excel_output_path}")


# =============================================================================
#   EXPORT: JSON
# =============================================================================

def export_results_to_json(data, json_output_path):
    """
    Saves the updated JSON structure to disk.
    """
    with open(json_output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

    print(f"✅ JSON saved → {json_output_path}")



def insert_marking_images_from_json_into_docx_new(
    docx_path,
    output_path,
    image_paths,
    table_index,
    row,
    col,
    width_inches=5
):
    """
    Inserts the EXACT images downloaded from JSON
    into the DOCX, one after another.
    """

    doc = Document(docx_path)
    cell = doc.tables[table_index].cell(row, col)

    # Preserve existing text
    if cell.paragraphs and cell.paragraphs[-1].text.strip():
        cell.add_paragraph("")
        cell.add_paragraph("")

    for img_path in image_paths:
        p = cell.add_paragraph()
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))

    doc.save(output_path)



def extract_checkbox_indices_from_json(json_data, max_table=7):
    """
    Returns a sorted list of checkbox indices (int)
    that need to be ticked.
    """

    indices_to_tick = []

    for table in json_data.get("Tables", []):
        table_no = table.get("Table")

        # Stop after table 7
        if table_no is None or table_no > max_table:
            continue

        for item in table.get("Items", []):

            if not item.get("checkbox_value"):
                continue

            checkbox_index = item.get("checkbox_index")
            if checkbox_index is None:
                continue

            # Normalize checkbox_index → list[int]
            if isinstance(checkbox_index, list):
                index_list = checkbox_index

            elif isinstance(checkbox_index, str):
                # "[6,7]" → [6,7]
                index_list = [
                    int(x.strip())
                    for x in checkbox_index.strip("[]").split(",")
                    if x.strip().isdigit()
                ]

            else:
                # 3 or 3.0
                index_list = [int(checkbox_index)]

            # Evaluate each checkbox_value_<n>
            for idx in index_list:
                key = f"checkbox_value_{idx}"
                if item.get(key) is True:
                    indices_to_tick.append(idx)

    return sorted(set(indices_to_tick))



def apply_checkboxes_from_json(
    docx_path,
    output_path,
    json_data
):
    """
    Applies checkbox ticks based on JSON using
    set_checkbox_checked() exactly as-is.
    """

    checkbox_indices = extract_checkbox_indices_from_json(json_data)

    print(f"✔ Checkbox indices to tick: {checkbox_indices}")

    current_doc = docx_path

    for idx in checkbox_indices:
        set_checkbox_checked(
            docx_path=current_doc,
            out_path=output_path,
            checkbox_index=int(idx)
        )
        # Next checkbox should apply on updated file
        current_doc = output_path

    print("✅ All checkboxes applied.")




# =============================================================================
#   DOCX UPDATE (Arial 10 + Value Injection)
# =============================================================================

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def update_docx_tables_from_json_arial(docx_path, json_path, output_path):
    """
    Update DOCX tables using JSON values.
    - Updates only cells that have ai_fillable = True OR task_type = verdict_dependency.
    - Applies Arial 10 font to updated cells.
    - Saves to output_path (you will call the same name again to avoid temporary versions).
    """

    doc = Document(docx_path)

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    for table_obj in data["Tables"]:
        table_index = table_obj["Table"]

        # If table index missing → skip
        if table_index >= len(doc.tables):
            print(f"⚠️ Skipping missing table index: {table_index}")
            continue

        doc_table = doc.tables[table_index]

        for item in table_obj["Items"]:

            ai_fillable = item.get("ai_fillable", False)
            task_type = item.get("task_type")

            # Only update allowed fields
            # if not ai_fillable and task_type != "verdict_dependency":
            if not ai_fillable and task_type != "verdict_dependency" and not item.get("user_editable", False):
                continue

            row = item.get("answer_row")
            col = item.get("answer_column")
            value = item.get("value")

            if value is None:
                continue

            try:
                cell = doc_table.cell(row, col)
                cell.text = ""  # clear

                p = cell.paragraphs[0]
                run = p.add_run(str(value))
                font = run.font
                font.name = "Arial"
                font.size = Pt(10)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

            except Exception as e:
                print(f"⚠️ Error updating table {table_index} cell ({row},{col}): {e}")

    doc.save(output_path)
    print(f"✅ DOCX first update saved → {output_path}")


# =============================================================================
#   INSERT CHECKBOX (final update in SAME DOCX FILE)
# =============================================================================

def insert_legacy_checkbox_with_text(
    docx_path,
    output_path,
    sentence,
    table_index=9,
    row=4,
    col=0,
    new_lines=4
):
    """
    Inserts a legacy checkbox and text into the SAME DOCX.
    Used AFTER table updates.
    """

    doc = Document(docx_path)

    cell = doc.tables[table_index].cell(row, col)

    for _ in range(new_lines):
        cell.add_paragraph("")

    p = cell.add_paragraph()
    p_elm = p._p

    # fldChar BEGIN
    r1 = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r1.append(fld_begin)
    p_elm.append(r1)

    # instruction
    r2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "FORMCHECKBOX"
    r2.append(instr)
    p_elm.append(r2)

    # ffData checkbox
    r3 = OxmlElement("w:r")
    ffData = OxmlElement("w:ffData")
    checkBox = OxmlElement("w:checkBox")

    size = OxmlElement("w:size")
    size.set(qn("w:val"), "16")

    default = OxmlElement("w:default")
    default.set(qn("w:val"), "1")

    checkBox.append(size)
    checkBox.append(default)
    ffData.append(checkBox)
    r3.append(ffData)
    p_elm.append(r3)

    # fldChar END
    r4 = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r4.append(fld_end)
    p_elm.append(r4)

    p.add_run(" " + sentence)

    doc.save(output_path)
    print(f"✅ DOCX final update (checkbox added) → {output_path}")

# =============================================================================
#   UPDATE MARKING PLATE IN DOCX
# =============================================================================
import requests
from pathlib import Path

def fetch_marking_plate_via_vs2(
    retriever,
    output_filename="marking_plate.png"
):
    """
    Uses IMAGE VECTORSTORE (vs2) retriever to select the best marking plate.
    Returns the local downloaded image path.
    """

    QUERY = (
        "marking plate certification label rating label "
        "nameplate marking label safety marking"
    )

    print("🔍 Searching marking plate via vs2 vectorstore...")

    # 1 — Run vector search
    try:
        neighbors = retriever.invoke(QUERY)
    except Exception as e:
        print("❌ Vectorstore error:", e)
        return None

    if not neighbors:
        print("❌ No images returned from vs2 retriever.")
        return None

    # 2 — Take top-1 result
    doc = neighbors[0]
    blob_url = doc.metadata.get("blob_url")

    if not blob_url:
        print("❌ No blob_url found in vs2 result metadata.")
        return None

    print(f"✅ Selected marking plate URL:\n{blob_url}\n")

    # 3 — Download image
    try:
        img_data = requests.get(blob_url).content
    except Exception as e:
        print("❌ Error downloading image:", e)
        return None

    output_path = Path(output_filename)
    output_path.write_bytes(img_data)

    print(f"📥 Marking plate saved to: {output_path.resolve()}\n")

    return str(output_path.resolve())




import requests
from pathlib import Path

def download_marking_images_from_json_new(json_data: dict, project_data_dir):
    """
    Downloads ONLY the URLs present in JSON.marking_urls.
    Returns local image paths IN THE SAME ORDER.
    """

    image_paths = []

    for table in json_data.get("Tables", []):
        for item in table.get("Items", []):
            marking_urls = item.get("marking_urls")
            if not marking_urls:
                continue

            for entry in marking_urls:
                url = entry.get("url")
                img_id = entry.get("id")

                if not url:
                    continue

                filename = project_data_dir / f"marking_{img_id}.png"
                path = Path(filename)

                try:
                    resp = requests.get(url, timeout=30)
                    resp.raise_for_status()
                    path.write_bytes(resp.content)
                    image_paths.append(path)
                except Exception as e:
                    print(f"[ERROR] Failed to download {url}: {e}")

    return image_paths




from docx import Document
from docx.shared import Inches

# def fetch_marking_plate_urls_and_update_json_new(
#     retriever,
#     json_data: dict,
#     table_no: int,
#     question_row: int,
#     question_column: int,
#     top_k: int = 1
# ):
#     """
#     Fetches marking plate URLs from retriever and stores them
#     in JSON as the ONLY source of truth.
#     """

#     # Fetch from retriever
#     docs = retriever.invoke("marking plate")[:top_k]

#     marking_urls = []
#     for idx, d in enumerate(docs, start=1):
#         url = d.metadata.get("blob_url")
#         if url:
#             marking_urls.append({
#                 "id": idx,
#                 "url": url
#             })

#     if not marking_urls:
#         return json_data

#     # Update JSON at exact location
#     for table in json_data.get("Tables", []):
#         if table.get("Table") != table_no:
#             continue

#         for item in table.get("Items", []):
#             if (
#                 item.get("question_row") == question_row and
#                 item.get("question_column") == question_column
#             ):
#                 item["marking_urls"] = marking_urls
#                 return json_data

#     return json_data



def fetch_marking_plate_urls_and_update_json_new(
    retriever,
    json_data: dict,
    table_no: int,
    question_row: int,
    question_column: int,
    top_k: int = 1
):
    """
    Fetches marking plate URLs from retriever and stores them
    in JSON as the ONLY source of truth.
    Deterministic + robust marking plate detection.
    """

    # --------------------------------------------------
    # STEP 1 — High-recall query (do NOT reduce)
    # --------------------------------------------------
    QUERY = (
        "equipment marking label rating plate regulatory label "
        "electrical ratings certification serial number"
    )

    # Retrieve MORE candidates (important)
    docs = retriever.invoke(QUERY)[:20]

    if not docs:
        return json_data

    scored_candidates = []

    for d in docs:
        text = (d.page_content or "").lower()
        url = d.metadata.get("blob_url")

        if not url:
            continue

        score = 0

        # -------------------------------
        # STRONG POSITIVE SIGNALS
        # -------------------------------
        if re.search(r"\b(vdc|vac|hz|watt|amp|a)\b", text):
            score += 4

        if re.search(r"\b(sn|s/n|serial)\b", text):
            score += 3

        if re.search(r"\b(ce|ul|etl|fcc|intertek|csa|tuv)\b", text):
            score += 3

        # -------------------------------
        # WEAK POSITIVE SIGNALS
        # -------------------------------
        if "label" in text or "marking" in text:
            score += 2

        # -------------------------------
        # NEGATIVE SIGNALS (filter noise)
        # -------------------------------
        if any(k in text for k in ["figure", "table", "section", "page"]):
            score -= 2

        if len(text) > 2000:
            score -= 2  # large drawings / manuals

        # Threshold ensures TRUE marking plates only
        if score >= 5:
            scored_candidates.append((score, url))

    if not scored_candidates:
        return json_data

    # --------------------------------------------------
    # STEP 2 — Pick best candidates
    # --------------------------------------------------
    scored_candidates.sort(reverse=True, key=lambda x: x[0])

    marking_urls = [
        {"id": idx + 1, "url": url}
        for idx, (_, url) in enumerate(scored_candidates[:top_k])
    ]

    # --------------------------------------------------
    # STEP 3 — Update JSON (UNCHANGED)
    # --------------------------------------------------
    for table in json_data.get("Tables", []):
        if table.get("Table") != table_no:
            continue

        for item in table.get("Items", []):
            if (
                item.get("question_row") == question_row and
                item.get("question_column") == question_column
            ):
                item["marking_urls"] = marking_urls
                return json_data

    return json_data
 



# =============================================================================
#   MAIN PIPELINE: run_trf_generator
# =============================================================================

# def run_trf_generator(
#     blob_urls: list,
#     input_json_path: str,
#     input_docx_path: str,
#     output_json_path: str,
#     output_docx_path: str,
#     excel_output_path: str,
#     batch_size=150,
#     cooldown_sec=15,
#     max_workers=6
# ):
#     """
#     FULL TRF REPORT GENERATION PIPELINE
#     ------------------------------------
#     Uses EXACT notebook logic but wrapped in a single callable function.
#     Generates ONLY:
#         - 1 final JSON
#         - 1 final DOCX
#         - 1 Excel

#     No temporary Word files created.
#     """
    
#     vs=build_vectorstore_text()
#     vs2=build_vectorstore_image()

#     print("\n===============================================")
#     print("       STEP 1 — LOAD INPUT JSON")
#     print("===============================================")

#     with open(input_json_path, "r", encoding="utf-8") as f:
#         data = json.load(f)

#     # ---------------------------------------------------------
#     # Build tasks from JSON (grey logic included)
#     # ---------------------------------------------------------
#     print("\n===============================================")
#     print("       STEP 2 — BUILD TASKS")
#     print("===============================================")

#     tasks, item_refs = build_tasks_with_custom_prompt_grey(data, blob_urls)

#     # ---------------------------------------------------------
#     # Build retriever from vs2 for image top-5 suggestions
#     # ---------------------------------------------------------
#     print("\n===============================================")
#     print("       STEP 3 — COMPUTE TOP-5 IMAGE HINTS")
#     print("===============================================")

#     image_retriever_agent = vs2.as_retriever(search_kwargs={"k": 5})
#     tasks = update_tasks_with_top5_images(tasks, image_retriever_agent)

#     # tasks=tasks[:30]
#     # item_refs=item_refs[:30]

#     # ---------------------------------------------------------
#     # Batch processing with LLM / Grey mode
#     # ---------------------------------------------------------
#     print("\n===============================================")
#     print("       STEP 4 — RUN LLM PIPELINE")
#     print("===============================================")
    

#     results = process_tasks_with_batches_parallel_grey(
#         tasks,
#         item_refs,
#         rag_image,
#         vs,
#         batch_size=batch_size,
#         cooldown_sec=cooldown_sec,
#         max_workers=max_workers,
#         use_llm_inGrey=False,
#         stats=True
#     )

#     # Convert to DataFrame
#     df = results_to_dataframe(results["results"])

#     print("\n===============================================")
#     print("       STEP 5 — EXPORT EXCEL")
#     print("===============================================")

#     export_results_to_excel(df, excel_output_path)

#     # ---------------------------------------------------------
#     # Apply post-process rules
#     # ---------------------------------------------------------
#     print("\n===============================================")
#     print("       STEP 6 — APPLY POST-PROCESSING")
#     print("===============================================")

#     update_verdict_dependencies(data)

#     targets = [
#         "General product information and other remarks:\nDescription of unit:\n",
#         "Description of model differences:\n",
#         "Description of special features:\n(HV circuits, high pressure systems etc.)\n"
#     ]
#     data = apply_prefixes_to_items(data, targets)
#     data = prefix_summary_of_compliance(data)

#     # Attach blob URLs for evidence support
#     data = attach_blob_urls_to_text_support(data, blob_urls)
#     data = attach_blob_urls_to_image_support(data, blob_urls)

#     print("\n===============================================")
#     print("       STEP 7 — SAVE FINAL JSON")
#     print("===============================================")

#     export_results_to_json(data, output_json_path)

#     # ---------------------------------------------------------
#     # Update Word DOCX with JSON values
#     # ---------------------------------------------------------
#     print("\n===============================================")
#     print("       STEP 8 — UPDATE DOCX (VALUES + Arial 10)")
#     print("===============================================")

#     update_docx_tables_from_json_arial(
#         docx_path=input_docx_path,
#         json_path=output_json_path,
#         output_path=output_docx_path
#     )

#     # ---------------------------------------------------------
#     # Insert checkbox at designated position
#     # ---------------------------------------------------------
#     print("\n===============================================")
#     print("       STEP 9 — FINAL DOCX CHECKBOX UPDATE")
#     print("===============================================")

#     insert_legacy_checkbox_with_text(
#         docx_path=output_docx_path,     # update in-place
#         output_path=output_docx_path,   # overwrite same file
#         sentence='The product fulfils the requirements of IEC 61010-1:2010, IEC 61010-1:2010/AMD1:2016',
#         table_index=5,
#         row=12,
#         col=0,
#         new_lines=5
#     )

#     print("\n===============================================")
#     print("       STEP 9 — UPDATING MARKING PLATE")
#     print("===============================================")

#     fetch_marking_plate_urls_and_update_json_new(
#         retriever=image_retriever_agent,
#         docx_path=output_docx_path,
#         output_path=output_docx_path,
#         table_index=6,
#         row=0,
#         col=0,
#         width_inches=5,
#         filename="identified_marking_plate.png"
#     )

   


#     return {
#         "json": output_json_path,
#         "docx": output_docx_path,
#         "excel": excel_output_path
#     }


import json
from copy import deepcopy

def update_pta_from_iec_correct(
    pta_path: str,
    iec_path: str,
    output_path: str
):
    """
    CORRECT + SAFE UPDATE

    Unique key = (field, task_type)

    Rules:
    - Update existing PTA items
    - Append ONLY if key does not exist ANYWHERE in PTA
    - Append ONCE
    - No cross-table duplication
    - Idempotent
    """

    with open(pta_path, "r", encoding="utf-8") as f:
        pta = json.load(f)

    with open(iec_path, "r", encoding="utf-8") as f:
        iec = json.load(f)

    # -------------------------------------------------
    # 1. Build GLOBAL PTA key index
    # -------------------------------------------------
    pta_key_index = {}   # (field, task_type) -> item

    for table in pta.get("Tables", []):
        for item in table.get("Items", []):
            field = str(item.get("field", "")).strip()
            task_type = item.get("task_type")
            if field:
                pta_key_index[(field, task_type)] = item

    # -------------------------------------------------
    # 2. Build IEC key index with table reference
    # -------------------------------------------------
    iec_items = []  # (table_ref, key, item)

    for table in iec.get("Tables", []):
        for item in table.get("Items", []):
            field = str(item.get("field", "")).strip()
            task_type = item.get("task_type")
            if field:
                iec_items.append((table, (field, task_type), item))

    # -------------------------------------------------
    # 3. Update existing PTA items
    # -------------------------------------------------
    for _, key, iec_item in iec_items:
        if key in pta_key_index:
            pta_item = pta_key_index[key]
            for k, v in iec_item.items():
                if v is not None:
                    pta_item[k] = deepcopy(v)

    # -------------------------------------------------
    # 4. Append ONLY truly missing keys (once)
    # -------------------------------------------------
    for iec_table, key, iec_item in iec_items:
        if key not in pta_key_index:
            # find matching PTA table by table number
            target_table_no = iec_table.get("Table")

            for pta_table in pta.get("Tables", []):
                if pta_table.get("Table") == target_table_no:
                    pta_table["Items"].append(deepcopy(iec_item))
                    pta_key_index[key] = iec_item  # prevent future dupes
                    break

    # -------------------------------------------------
    # 5. Write output
    # -------------------------------------------------
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(pta, f, indent=2, ensure_ascii=False)

    return output_path


import json
import os

def update_pta_with_multiple_iec(
    base_pta_path: str,
    iec_paths: list[str],
    final_output_path: str,
    project_data_dir: Path

):
    print(f"[START] Base PTA: {base_pta_path}")
    temp_pta = base_pta_path

    for i, iec_path in enumerate(iec_paths, start=1):
        is_last = i == len(iec_paths)
        next_pta = Path(final_output_path) if is_last else project_data_dir / f"_tmp_pta_{i}.json"

        print(f"[STEP {i}/{len(iec_paths)}] Applying IEC: {iec_path}")
        print(f"[INFO] Input PTA : {temp_pta}")
        print(f"[INFO] Output PTA: {next_pta}")

        update_pta_from_iec_correct(
            pta_path=temp_pta,
            iec_path=iec_path,
            output_path=next_pta,
        )

        # ✅ HARD GUARANTEE: file must exist
        if not Path(next_pta).exists():
            raise RuntimeError(f"[ERROR] PTA output not created: {next_pta}")

        print(f"[DONE] PTA dumped successfully → {next_pta}\n")
        temp_pta = next_pta

    print(f"[FINAL] Loading merged PTA: {final_output_path}")

    with open(final_output_path, "r", encoding="utf-8") as f:
        final_data = json.load(f)

    print("[SUCCESS] PTA merge completed")
    return final_data



def delete_cosmos_container(
    endpoint: str,
    key: str,
    database_name: str,
    container_name: str
):
    """
    Deletes a Cosmos DB container.
    """

    client = CosmosClient(endpoint, credential=key)

    try:
        database = client.get_database_client(database_name)
        database.delete_container(container_name)
        
        print("\n===============================================")
        print(f"       STEP 12 — DELETING THE VECTOR COSMOS DB Container '{container_name}'")
        print("===============================================")

        print(f"Container '{container_name}' deleted successfully.")


    except ResourceNotFoundError:
        print(f"Container '{container_name}' or database '{database_name}' not found.")

    except Exception as e:
        raise RuntimeError(f"Failed to delete container: {e}")


# =============================================================================
#   MAIN PIPELINE: run_trf_generator
# =============================================================================

def trf_gen_partwise(
    blob_urls: list,
    vs,
    vs2,
    rag_image,
    run_single_task_tool,
    input_json_path: str,
    output_json_path: str,
    excel_output_path: str,
    batch_size=150,
    cooldown_sec=15,
    max_workers=6
):
    """
    FULL TRF REPORT GENERATION PIPELINE
    ------------------------------------
    Uses EXACT notebook logic but wrapped in a single callable function.
    Generates ONLY:
        - 1 final JSON
        - 1 final DOCX
        - 1 Excel

    No temporary Word files created.
    """

    print("\n===============================================")
    print("       STEP 1 — LOAD INPUT JSON")
    print("===============================================")

    BASE_DIR = Path(__file__).resolve().parent

    input_json_path = BASE_DIR / input_json_path

    print("Resolved path:", input_json_path)

    with open(input_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # ---------------------------------------------------------
    # Build tasks from JSON (grey logic included)
    # ---------------------------------------------------------
    print("\n===============================================")
    print("       STEP 2 — BUILD TASKS")
    print("===============================================")

    tasks, item_refs = build_tasks_with_custom_prompt_grey(data, blob_urls)

    # ---------------------------------------------------------
    # Build retriever from vs2 for image top-5 suggestions
    # ---------------------------------------------------------
    print("\n===============================================")
    print("       STEP 3 — COMPUTE TOP-5 IMAGE HINTS")
    print("===============================================")

    image_retriever_agent = vs2.as_retriever(search_kwargs={"k": 5})
    tasks = update_tasks_with_top5_images(tasks, image_retriever_agent)

    # tasks=tasks[:30]
    # item_refs=item_refs[:30]

    # ---------------------------------------------------------
    # Batch processing with LLM / Grey mode
    # ---------------------------------------------------------
    print("\n===============================================")
    print("       STEP 4 — RUN LLM PIPELINE")
    print("===============================================")

    results = process_tasks_with_batches_parallel_grey(
        tasks,
        item_refs,
        rag_image,
        vs,
        run_single_task_tool,
        batch_size=batch_size,
        cooldown_sec=cooldown_sec,
        max_workers=max_workers,
        use_llm_inGrey=False,
        stats=True
    )

    # Convert to DataFrame
    df = results_to_dataframe(results["results"])

    print("\n===============================================")
    print("       STEP 5 — EXPORT EXCEL")
    print("===============================================")

    # export_results_to_excel(df, excel_output_path)

    print("\n===============================================")
    print("       STEP 6 — SAVE JSON")
    print("===============================================")

    export_results_to_json(data, output_json_path)
    return {
        "json": output_json_path,
        "excel": excel_output_path
    }



def run_trf_generation(
    blob_urls: list,
    textDB_container_name,
    imageDB_container_name,
    input_docx_path: str,
    output_docx_path: str,
    base_pta_path: str,
    input_json_paths: list,
    project_data_dir: Path,
    batch_size=150,
    final_output_path: str | None = None,
    cooldown_sec=15,
    max_workers=6,
    on_first_json_generated: Optional[Callable[[], None]] = None,  # 🔔 EVENT HOOK
):
    iec_paths = []
    first_event_fired = False

    def build_embeddings(AOAI_ENDPOINT,AOAI_KEY,API_VERSION,EMBED_DEPLOY):
        return AzureOpenAIEmbeddings(
            azure_endpoint=AOAI_ENDPOINT,
            api_key=AOAI_KEY,
            openai_api_version=API_VERSION,
            azure_deployment=EMBED_DEPLOY,
        )

    vs=build_vectorstore_text(textDB_container_name)
    vs2=build_vectorstore_image(imageDB_container_name)

    retriever = vs.as_retriever(search_kwargs={"k": 5})
    image_retriever_agent = vs2.as_retriever(search_kwargs={"k": 5})

    # Build the final RAG pipeline
    rag_image = build_rag_image_pipeline_grey(
        retriever,
        llm,
        build_vision_message_grey,
        attach_supporting_refs_grey,
        vs
    )

    run_single_task_tool = generator_evaluator_agent(rag_image)

    # ---------------------------------------------------------
    # STEP 1–6 — PARTWISE JSON GENERATION
    # ---------------------------------------------------------
    for idx, input_json_path in enumerate(input_json_paths, start=1):
        print(f"[INFO] ({idx}/{len(input_json_paths)}) Processing: {input_json_path}")

        input_json_path = Path(input_json_path)
        base_name = input_json_path.stem   # e.g. pta_final_6_2_part1

        output_json_path = project_data_dir / f"{base_name}_output.json"
        excel_output_path = project_data_dir / f"{base_name}.xlsx"

        output_json_path.parent.mkdir(parents=True, exist_ok=True)

        result = trf_gen_partwise(
            blob_urls=blob_urls,
            vs=vs,
            vs2=vs2,
            rag_image=rag_image,
            run_single_task_tool=run_single_task_tool,
            input_json_path=str(input_json_path),
            output_json_path=str(output_json_path),
            excel_output_path=str(excel_output_path),
            batch_size=batch_size,
            cooldown_sec=cooldown_sec,
            max_workers=max_workers,
        )

        print(f"[DONE] JSON generated: {output_json_path}")
        print(f"[DONE] Excel generated: {excel_output_path}")


        iec_paths.append(result["json"])

        # 🔔 FIRE EVENT AFTER FIRST JSON ONLY
        if not first_event_fired:
            first_event_fired = True
            if on_first_json_generated:
                try:
                    on_first_json_generated()
                except Exception as e:
                    # Never break TRF pipeline due to notification failure
                    print(f"[WARN] First JSON event callback failed: {e}")

    # ---------------------------------------------------------
    # STEP 7 — CREATE FINAL JSON
    # ---------------------------------------------------------
    print("\n===============================================")
    print("       STEP 7 — CREATE FINAL JSON")
    print("===============================================")

    data = update_pta_with_multiple_iec(
        base_pta_path,
        iec_paths=iec_paths,
        final_output_path=final_output_path,
        project_data_dir= project_data_dir
    )

    print("\n===============================================")
    print("       STEP 8 — APPLY POST-PROCESSING")
    print("===============================================")

    update_verdict_dependencies(data)

    targets = [
        "General product information and other remarks:\nDescription of unit:\n",
        "Description of model differences:\n",
        "Description of special features:\n(HV circuits, high pressure systems etc.)\n"
    ]
    data = apply_prefixes_to_items(data, targets)
    data = prefix_summary_of_compliance(data)

    # Attach blob URLs for evidence support
    data = attach_blob_urls_to_text_support(data, blob_urls)
    data = attach_blob_urls_to_image_support(data, blob_urls)

    print("\n===============================================")
    print("       STEP 9 — SAVE FINAL JSON")
    print("===============================================")

    export_results_to_json(data, final_output_path)

    # # ---------------------------------------------------------
    # # Update Word DOCX with JSON values
    # # ---------------------------------------------------------
    print("\n===============================================")
    print("       STEP 10 — UPDATE DOCX (VALUES + Arial 10)")
    print("===============================================")

    update_docx_tables_from_json_arial(
        docx_path=input_docx_path,
        json_path=final_output_path,
        output_path=output_docx_path
    )

    # ---------------------------------------------------------
    # Insert checkbox at designated position
    # ---------------------------------------------------------
    print("\n===============================================")
    print("       STEP 11 — FINAL DOCX CHECKBOX UPDATE")
    print("===============================================")

    insert_legacy_checkbox_with_text(
        docx_path=output_docx_path,     # update in-place
        output_path=output_docx_path,   # overwrite same file
        sentence='The product fulfils the requirements of IEC 61010-1:2010, IEC 61010-1:2010/AMD1:2016',
        table_index=5,
        row=12,
        col=0,
        new_lines=5
    )

    print("\n===============================================")
    print("       STEP 12 — UPDATING MARKING PLATE And INTERTEK LOGO")
    print("===============================================")

    # insert_marking_plate_via_vs2(
    #     retriever=image_retriever_agent,
    #     docx_path=output_docx_path,
    #     output_path=output_docx_path,
    #     table_index=6,
    #     row=0,
    #     col=0,
    #     width_inches=5,
    #     filename="identified_marking_plate.png"
    # )

    data = fetch_marking_plate_urls_and_update_json_new(
        retriever=image_retriever_agent,
        json_data=data,
        table_no=6,
        question_row=0,
        question_column=0,
        top_k=1
    )

    image_paths = download_marking_images_from_json_new(data, project_data_dir)

    export_results_to_json(data, final_output_path)

    insert_marking_images_from_json_into_docx_new(
        docx_path=output_docx_path,
        output_path=output_docx_path,
        image_paths=image_paths,
        table_index=6,
        row=0,
        col=0,
        width_inches=2
    )

    apply_checkboxes_from_json(
        docx_path=output_docx_path,
        output_path=output_docx_path,
        json_data=data
    )

    delete_cosmos_container(
    endpoint=COSMOS_URL,
    key=COSMOS_KEY,
    database_name=COSMOS_DB_TEXT,
    container_name=textDB_container_name
    )

    delete_cosmos_container(
    endpoint=COSMOS_URL,
    key=COSMOS_KEY,
    database_name=COSMOS_DB_IMAGE,
    container_name=imageDB_container_name
    )

    print("\n✅ TRF GENERATION COMPLETED SUCCESSFULLY")





