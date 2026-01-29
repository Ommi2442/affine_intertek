import os
import time
from urllib.parse import urlparse, unquote
from azure.cosmos import CosmosClient, ConsistencyLevel
import json

# Import helpers from your existing utils
from utility.trf_utils import *
import re
import tempfile
import shutil
import requests
from urllib.parse import urlparse, unquote
from email import policy
from email.parser import BytesParser
import extract_msg
import fitz
import uuid
from langchain_core.documents import Document
import io
import openpyxl
import xlrd
# from utils import *
from azure.storage.blob import BlobClient
from azure.core.exceptions import ResourceNotFoundError, AzureError
# from templates import *
import pandas as pd
import math
import copy
import time
from azure.cosmos import CosmosClient, PartitionKey, exceptions
import json, os
from azure.cosmos import CosmosClient, ConsistencyLevel
from typing import List, Dict, Any, Tuple
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
from langchain_azure_ai.vectorstores import AzureCosmosDBNoSqlVectorSearch
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from azure.cosmos import CosmosClient
from operator import itemgetter
from langchain_core.runnables import (
    RunnableParallel, RunnableLambda, RunnableMap
)
from langchain_core.output_parsers import StrOutputParser
from langchain_openai import AzureChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from tenacity import retry, retry_if_exception_type, wait_exponential, stop_never, RetryCallState
from openai import RateLimitError  # Make sure this import exists
from types import SimpleNamespace
from concurrent.futures import ThreadPoolExecutor, as_completed
from azure.core.exceptions import HttpResponseError
import time
from langchain_community.callbacks import get_openai_callback
pd.set_option('display.max_colwidth', None)  # Don't truncate cell text
pd.set_option('display.max_rows', None)      # Show all rows (optional)
pd.set_option('display.max_columns', None)
from dotenv import load_dotenv
from pathlib import Path
from openai import AzureOpenAI
from pypdf import PdfReader
from azure.cosmos.exceptions import CosmosResourceNotFoundError
import shutil


from projects.keyvault_load import *
load_keyvault_secrets()

load_dotenv()

# Chunking config
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150
EMBED_DIM = 1536
VECTOR_PATH = "/vector"
TOP_K = 5


# Load environment variables
AOAI_ENDPOINT      = os.getenv("aoai-endpoint")
AOAI_KEY           = os.getenv("aoai-key")
API_VERSION        = os.getenv("api-version")
EMBED_DEPLOY       = os.getenv("embed-deploy")
COSMOS_DB_TEXT     = os.getenv("cosmos-db-text")
COSMOS_CONT_TEXT   = os.getenv("cosmos-cont-text")
AZURE_CONN_STRING  = os.getenv("azure-conn-string")
BLOB_CONTAINER     = os.getenv("blob-container")
COSMOS_URL         = os.getenv("cosmos-url")
COSMOS_KEY         = os.getenv("cosmos-key")
CHAT_DEPLOY        = os.getenv("chat-deploy")
BLOB_CONTAINER     = os.getenv("blob-container")
COSMOS_DB_IMAGE    = os.getenv("cosmos-db-image")
COSMOS_CONT_IMAGE  = os.getenv("cosmos-cont-image")
ENABLE_CAD_SCHEMATICS  = os.getenv("enable-cad-schematics")



print('AOAI_ENDPOINT', AOAI_ENDPOINT)

# ----------------------------------------------------------------------------------------
# Azure OpenAI Client (shared for whole pipeline)
# ----------------------------------------------------------------------------------------
aoai_client = AzureOpenAI(
    api_key=AOAI_KEY,
    api_version=API_VERSION,
    azure_endpoint=AOAI_ENDPOINT,
)


# ----------------------------------------------------------------------------------------
# Embedding Builder — same as notebook
# ----------------------------------------------------------------------------------------

def build_embeddings():
    return AzureOpenAIEmbeddings(
        azure_endpoint=AOAI_ENDPOINT,
        api_key=AOAI_KEY,
        openai_api_version=API_VERSION,
        azure_deployment=EMBED_DEPLOY,
    )


# ----------------------------------------------------------------------------------------
# Vector Store Builder (for TEXT)
# EXACT logic from notebook (not altered)
# ----------------------------------------------------------------------------------------


def build_vectorstore_text(textDB_container_name):
    cosmos_client = CosmosClient(
        url=COSMOS_URL,
        credential=COSMOS_KEY
    )

    return AzureCosmosDBNoSqlVectorSearch(
        cosmos_client=cosmos_client,
        embedding=build_embeddings(),
        database_name=COSMOS_DB_TEXT,
        container_name=textDB_container_name,

        vector_embedding_policy={
            "vectorEmbeddings": [{
                "path": "/vector",
                "dataType": "float32",
                "dimensions": EMBED_DIM,
                "distanceFunction": "cosine"
            }]
        },
        indexing_policy={
            "includedPaths": [{"path": "/*"}],
            "excludedPaths": [{"path": "/\"_etag\"/?"}, {"path": "/vector/*"}],
            "vectorIndexes": [{"path": "/vector", "type": "quantizedFlat"}],
        },
        cosmos_container_properties={"partition_key": "/id"},
        cosmos_database_properties={},
        vector_search_fields={
            "text_field": "text",
            "embedding_field": "vector",
            "metadata_field": "metadata"
        }
    )


# ----------------------------------------------------------------------------------------
# Vector Store Builder (for IMAGES)
# EXACT logic from notebook (not altered)
# ----------------------------------------------------------------------------------------
def build_vectorstore_image(imageDB_container_name):
    cosmos_client = CosmosClient(
        url=COSMOS_URL,
        credential=COSMOS_KEY
    )

    return AzureCosmosDBNoSqlVectorSearch(
        cosmos_client=cosmos_client,
        embedding=build_embeddings(),
        database_name=COSMOS_DB_IMAGE,
        container_name=imageDB_container_name,

        vector_embedding_policy={
            "vectorEmbeddings": [{
                "path": "/vector",
                "dataType": "float32",
                "dimensions": EMBED_DIM,
                "distanceFunction": "cosine"
            }]
        },
        indexing_policy={
            "includedPaths": [{"path": "/*"}],
            "excludedPaths": [{"path": "/\"_etag\"/?"}, {"path": "/vector/*"}],
            "vectorIndexes": [{"path": "/vector", "type": "quantizedFlat"}],
        },
        cosmos_container_properties={"partition_key": "/id"},
        cosmos_database_properties={},
        vector_search_fields={
            "text_field": "text",
            "embedding_field": "vector",
            "metadata_field": "metadata"
        }
    )


# ingestion_tool.py (continued)
# --------------------------------------------------------
# SECTION 2.2 — PDF Text Loading, CAD/Schematic Image Extraction,
#                Chunking Logic, Image Upload to Azure Blob Storage
# --------------------------------------------------------


# -------------------------------------------------------------------------
# Utility: Sanitize blob names (same as in notebook)
# -------------------------------------------------------------------------
def sanitize_blob_name(name: str) -> str:
    name = name.replace(" ", "_")
    name = re.sub(r"[^A-Za-z0-9_\-./]", "_", name)
    return name



# def is_editable_pdf(pdf_path) -> bool:
#     """
#     True if PDF contains AcroForm (fillable fields).
#     """
#     reader = PdfReader(pdf_path)
#     return bool(reader.get_fields())


# -------------------------------------------------------------------------
# STRICT CAD/Schematic Page Detector (EXACT logic from notebook)
# -------------------------------------------------------------------------
# def extract_relevant_pdf_page_images(pdf_path, dpi=200):
#     """
#     STRICT selection of pages containing diagrams / CAD / schematics.
#     EXTENDED:
#     - If PDF is editable → convert ALL pages to images
#     - Else → keep EXACT notebook logic
#     """

#     import fitz
#     import os

#     pdf = fitz.open(pdf_path)
#     base = os.path.basename(pdf_path)

#     image_page_metadata = []

#     # -------------------------------------------------
#     # NEW: Editable PDF shortcut
#     # -------------------------------------------------
#     if is_editable_pdf(pdf_path):
#         for i, page in enumerate(pdf):
#             page_num = i + 1
#             try:
#                 pix = page.get_pixmap(dpi=dpi)
#                 img_path = f"{pdf_path}_page_{page_num}.png"
#                 pix.save(img_path)

#                 image_page_metadata.append({
#                     "pdf_file": base,
#                     "page": page_num,
#                     "local_image_path": img_path,
#                     "reason": "editable_pdf_full_page"
#                 })
#             except Exception as e:
#                 print(f"[WARN] Image extraction failed for {base} page {page_num}: {e}")

#         return image_page_metadata

#     # -------------------------------------------------
#     # ORIGINAL NOTEBOOK LOGIC (UNTOUCHED)
#     # -------------------------------------------------
#     for i, page in enumerate(pdf):
#         page_num = i + 1

#         text = page.get_text().strip()
#         images = page.get_images(full=True)
#         drawings = page.get_drawings()
#         blocks = page.get_text("blocks")

#         text_len = len(text)
#         num_blocks = len(blocks)
#         vector_ops = len(drawings)

#         raster_area = 0
#         for img in images:
#             try:
#                 w = img[2]
#                 h = img[3]
#                 raster_area += (w * h)
#             except:
#                 continue

#         should_extract = False

#         if raster_area > 500000:
#             should_extract = True
#         elif vector_ops > 150:
#             should_extract = True
#         elif text_len < 30:
#             should_extract = True
#         elif num_blocks > 30 and text_len < 150:
#             should_extract = True
#         elif (
#             any(k in text.lower() for k in ["label", "regulatory"])
#             and vector_ops > 20
#             and text_len < 600
#         ):
#             should_extract = True
#         elif any(k in base.lower() for k in ["schematic", "cad", "drawing", "layout", "wiring"]) and text_len < 80:
#             should_extract = True
#         elif len(pdf) <= 3 and (vector_ops > 20 or text_len < 600):
#             should_extract = True

#         if not should_extract:
#             continue

#         try:
#             pix = page.get_pixmap(dpi=dpi)
#             img_path = f"{pdf_path}_page_{page_num}.png"
#             pix.save(img_path)

#             image_page_metadata.append({
#                 "pdf_file": base,
#                 "page": page_num,
#                 "local_image_path": img_path,
#                 "text_length": text_len,
#                 "raster_area": raster_area,
#                 "vector_ops": vector_ops,
#                 "blocks": num_blocks,
#             })

#         except Exception as e:
#             print(f"[WARN] Image extraction failed for {base} page {page_num}: {e}")

#     return image_page_metadata



def extract_relevant_pdf_page_images(pdf_path, dpi=200):
    """
    STRICT selection of pages containing diagrams / CAD / schematics.
    EXACT code copied from your notebook. No modifications.
    """

    import fitz
    pdf = fitz.open(pdf_path)
    base = os.path.basename(pdf_path)

    image_page_metadata = []

    for i, page in enumerate(pdf):
        page_num = i + 1

        text = page.get_text().strip()
        images = page.get_images(full=True)
        drawings = page.get_drawings()
        blocks = page.get_text("blocks")

        text_len = len(text)
        num_blocks = len(blocks)
        vector_ops = len(drawings)

        raster_area = 0
        for img in images:
            try:
                w = img[2]
                h = img[3]
                raster_area += (w * h)
            except:
                continue

        # --- EXACT strict rules from notebook ---
        should_extract = False

        if raster_area > 750000:
            should_extract = True
        elif vector_ops > 150 and text_len < 800:
            should_extract = True
        elif text_len < 30:
            should_extract = True
        elif num_blocks > 30 and text_len < 150:
            should_extract = True
        elif (
            any(k in text.lower() for k in ["label", "regulatory"])
            and vector_ops > 20
            and text_len < 600
        ):
            should_extract = True
        elif any(k in base.lower() for k in ["schematic", "cad", "drawing", "layout", "wiring"]) and text_len < 80:
            should_extract = True
        elif len(pdf) <= 3 and (vector_ops > 20 or text_len < 600):
            should_extract = True

        if not should_extract:
            continue

        # ---- Extract the page image EXACTLY like notebook ----
        try:
            pix = page.get_pixmap(dpi=dpi)
            img_path = f"{pdf_path}_page_{page_num}.png"
            pix.save(img_path)

            image_page_metadata.append({
                "pdf_file": base,
                "page": page_num,
                "local_image_path": img_path,
                "text_length": text_len,
                "raster_area": raster_area,
                "vector_ops": vector_ops,
                "blocks": num_blocks,
            })

        except Exception as e:
            print(f"[WARN] Image extraction failed for {base} page {page_num}: {e}")

    return image_page_metadata



# -------------------------------------------------------------------------
# PDF Loader + Text Chunking (EXACT logic from notebook)
# -------------------------------------------------------------------------
def load_and_split_pdfs_text(
    pdf_paths,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    extracted_texts=None,
    cad_schematics=True
):
    """
    EXACT implementation from your notebook.
    Returns:
        chunks → list of text Document objects
        image_page_metadata → list of schematic image metadata
    """

    docs = []
    image_page_metadata = []

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " "],
        keep_separator=False,
    )

    # ----- STEP 1: PDF TEXT EXTRACTION -----
    for path in pdf_paths:
        if not str(path).lower().endswith(".pdf"):
            continue

        loader = PyPDFLoader(str(path))
        raw_docs = loader.load()
        base = os.path.basename(str(path))

        for d in raw_docs:
            page = int(d.metadata.get("page", 1))
            d.metadata["source_file"] = base
            d.metadata["page"] = page
            d.metadata["citation"] = f"{base}#page={page}"

        docs.extend(raw_docs)

        # ----- STEP 2: CAD/Schematic Image Extraction -----
        if cad_schematics:
            try:
                extracted = extract_relevant_pdf_page_images(path)
                image_page_metadata.extend(extracted)
            except Exception as e:
                print(f"[WARN] selective image extraction failed for {path}: {e}")

    # ----- STEP 3: External extracted text -----
    # -------------------------------------------------------------
    # STEP 2 — Process externally extracted text files (unchanged)
    # -------------------------------------------------------------
    if extracted_texts:
        for item in extracted_texts:
            if not isinstance(item, dict):
                continue

            if "filename" in item and "text" in item:
                filename = item["filename"]
                text = item["text"]
            elif len(item) == 1:
                filename, text = next(iter(item.items()))
            elif "text" in item:
                filename = item.get("filename") or "unknown"
                text = item["text"]
            else:
                filename = item.get("filename") or "unknown"
                text = None
                for k, v in item.items():
                    if isinstance(v, str) and v.strip():
                        filename = k
                        text = v
                        break
                if text is None:
                    text = " ".join(str(v) for v in item.values())

            metadata = {
                "source_file": os.path.basename(str(filename)),
                "page": 1,
                "citation": os.path.basename(str(filename))
            }

            # ORIGINAL WORKING VERSION — KEEP SimpleNamespace
            docs.append(
                SimpleNamespace(
                    page_content=text or "",
                    metadata=metadata
                )
            )



    # ----- STEP 4: Chunking -----
    chunks = splitter.split_documents(docs)

    return chunks, image_page_metadata


# -------------------------------------------------------------------------
# Upload extracted PDF page images → Azure Blob Storage
# (EXACT logic from notebook)
# -------------------------------------------------------------------------
# def upload_pdf_images_and_append_urls(
#     image_page_metadata,
#     image_urls,
#     conn_str,
#     container
# ):
#     """
#     Takes CAD/schematic page images extracted above,
#     uploads each PNG to blob storage,
#     appends URLs in same format as notebook.
#     """

#     for item in image_page_metadata:
#         local_path = item["local_image_path"]
#         pdf_file = item["pdf_file"]
#         page = item.get("page") or item.get("page_num")

#         safe_pdf_name = sanitize_blob_name(pdf_file)
#         safe_image_filename = sanitize_blob_name(os.path.basename(local_path))

#         blob_name = f"{safe_pdf_name}/page_{page}.png"

#         try:
#             blob = BlobClient.from_connection_string(
#                 conn_str,
#                 container_name=container,
#                 blob_name=blob_name,
#             )

#             with open(local_path, "rb") as f:
#                 blob.upload_blob(f, overwrite=True)

#             blob_url = blob.url

#             # EXACT return structure from notebook
#             image_urls.append({
#                 "url": blob_url,
#                 "image_file": safe_image_filename,
#                 "pdf_file": pdf_file,
#                 "page": page
#             })

#         except Exception as e:
#             print(f"[ERROR] Upload failed for {local_path}: {e}")

#     return image_urls



# def upload_pdf_images_and_append_urls(
#     image_page_metadata,
#     image_urls,
#     conn_str,
#     container,
#     max_workers=8
# ):
#     """
#     Takes CAD/schematic page images extracted above,
#     uploads each PNG to blob storage in parallel,
#     appends URLs in same format as notebook.
#     """

#     def upload_single(item):
#         local_path = item["local_image_path"]
#         pdf_file = item["pdf_file"]
#         page = item.get("page") or item.get("page_num")

#         safe_pdf_name = sanitize_blob_name(pdf_file)
#         safe_image_filename = sanitize_blob_name(os.path.basename(local_path))

#         blob_name = f"{safe_pdf_name}/page_{page}.png"

#         try:
#             blob = BlobClient.from_connection_string(
#                 conn_str,
#                 container_name=container,
#                 blob_name=blob_name,
#             )

#             with open(local_path, "rb") as f:
#                 blob.upload_blob(f, overwrite=True)

#             blob_url = blob.url

#             return {
#                 "url": blob_url,
#                 "image_file": safe_image_filename,
#                 "pdf_file": pdf_file,
#                 "page": page
#             }

#         except Exception as e:
#             print(f"[ERROR] Upload failed for {local_path}: {e}")
#             return None

#     # -------------------------------------------------
#     # PARALLEL EXECUTION
#     # -------------------------------------------------
#     with ThreadPoolExecutor(max_workers=max_workers) as executor:
#         futures = [
#             executor.submit(upload_single, item)
#             for item in image_page_metadata
#         ]

#         for future in as_completed(futures):
#             result = future.result()
#             if result:
#                 image_urls.append(result)

#     return image_urls





from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from azure.storage.blob import BlobClient
import os

def upload_pdf_images_and_append_urls(
    image_page_metadata,
    image_urls,
    conn_str,
    container,
    max_workers,
    project_id      
):
    """
    Takes CAD/schematic page images extracted above,
    uploads each PNG to blob storage in parallel,
    appends URLs in same format as notebook.

    Blob path:
    Documents/<project_id>/pdf_images/<pdf_file>/page_<n>.png
    """

    # -------------------------------------------------
    # BASE BLOB PATH (matches reference logic)
    # -------------------------------------------------
    base_blob_path = (
        Path("Documents") /
        str(project_id) /
        "pdf_images"
    ).as_posix()   # Azure requires forward slashes

    def upload_single(item):
        local_path = item["local_image_path"]
        pdf_file = item["pdf_file"]
        page = item.get("page") or item.get("page_num")

        if not pdf_file or page is None:
            return None

        safe_pdf_name = sanitize_blob_name(pdf_file)

        # ✅ NEW BLOB PATH
        blob_name = f"{base_blob_path}/{safe_pdf_name}/page_{page}.png"

        try:
            blob = BlobClient.from_connection_string(
                conn_str,
                container_name=container,
                blob_name=blob_name,
            )

            with open(local_path, "rb") as f:
                blob.upload_blob(f, overwrite=True)

            return {
                "url": blob.url,
                "image_file": os.path.basename(local_path),
                "pdf_file": pdf_file,
                "page": page
            }

        except Exception as e:
            print(f"[ERROR] Upload failed for {local_path}: {e}")
            return None

    # -------------------------------------------------
    # PARALLEL EXECUTION (UNCHANGED)
    # -------------------------------------------------
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = [
            executor.submit(upload_single, item)
            for item in image_page_metadata
        ]

        for future in as_completed(futures):
            result = future.result()
            if result:
                image_urls.append(result)

    return image_urls
 




# ingestion_tool.py (continued)
# --------------------------------------------------------
# SECTION 3.1 — Image URL Extraction + Agent-Based OCR Pipeline
# --------------------------------------------------------


# --------------------------------------------------------
# Extract image name from blob URL — EXACT as notebook
# --------------------------------------------------------
def extract_clean_image_name(blob_url: str):
    parsed = urlparse(blob_url)
    path = parsed.path
    parts = path.split("/")

    pdf_file = None
    for part in parts:
        if part.lower().endswith(".pdf"):
            pdf_file = part
            break

    image_filename = os.path.basename(path)

    if pdf_file:
        return f"{pdf_file}/{image_filename}"
    else:
        return image_filename


# --------------------------------------------------------
# Wrapper: turn mixed list into plain list of URLs
# --------------------------------------------------------
def extract_urls(mixed_list):
    urls = []

    for item in mixed_list:
        if isinstance(item, str):
            urls.append(item)
        elif isinstance(item, dict) and "url" in item:
            urls.append(item["url"])

    return urls


# --------------------------------------------------------
# The ORIGINAL process_single_image replaced by an AGENT CALL
# --------------------------------------------------------
MAX_THREADS = 10
MAX_RETRIES = 5
INITIAL_BACKOFF = 3


def _direct_llm_fallback(url, index, total, vision_deploy_name):
    """
    Only used if agent not triggered correctly.
    Matches original notebook LLM behavior exactly.
    """

    resp = requests.get(url, timeout=20)
    resp.raise_for_status()

    completion = aoai_client.chat.completions.create(
        model=vision_deploy_name,
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": url},
                    {
                        "type": "text",
                        "text": (
                            "Perform OCR on this image and also provide a detailed "
                            "description. Combine both into one response."
                        ),
                    },
                ],
            }
        ],
        max_tokens=2048,
    )

    extracted_text = completion.choices[0].message.content.strip()
    image_name = extract_clean_image_name(url)

    return Document(
        page_content=extracted_text,
        metadata={
            "image_name": image_name,
            "blob_url": url,
            "source_type": "image",
        },
    )


# --------------------------------------------------------
# Agent Function Schema (EXACT as notebook)
# --------------------------------------------------------
image_desc_schema = {
    "name": "image_description_agent",
    "description": "OCR + description on an Azure Blob image URL.",
    "parameters": {
        "type": "object",
        "properties": {
            "url": {"type": "string"},
            "vision_deploy_name": {"type": "string"},
            "index": {"type": "integer"},
            "total": {"type": "integer"},
        },
        "required": ["url", "vision_deploy_name"],
    }
}


# --------------------------------------------------------
# run_function_agent_llm — EXACT as notebook
# --------------------------------------------------------
def run_function_agent_llm(client, user_prompt: str, function_schema: dict, python_callback: callable):
    response = client.chat.completions.create(
        model=CHAT_DEPLOY,
        messages=[{"role": "user", "content": user_prompt}],
        functions=[function_schema],
    )

    msg = response.choices[0].message

    if msg.function_call:
        fn_args = json.loads(msg.function_call.arguments)
        return python_callback(**fn_args)

    return msg.content


# --------------------------------------------------------
# Callback executed when LLM triggers function_call
# --------------------------------------------------------
def image_desc_callback(url, vision_deploy_name, index=1, total=1):
    """
    Internally calls _direct_llm_fallback so that return format matches EXACT original logic.
    """
    try:
        return _direct_llm_fallback(url, index, total, vision_deploy_name)
    except Exception as e:
        print(f"[ERROR] Fallback LLM failed for {url}: {e}")
        return None


# --------------------------------------------------------
# Main Agent Wrapper — same structure as notebook
# --------------------------------------------------------
def image_desc_agent(blob_url, vision_deploy_name="gpt-4.1", index=1, total=1):

    payload = {
        "url": blob_url,
        "vision_deploy_name": vision_deploy_name,
        "index": index,
        "total": total,
    }

    return run_function_agent_llm(
        aoai_client,
        user_prompt="Use the image_description_agent tool on: " + json.dumps(payload),
        function_schema=image_desc_schema,
        python_callback=image_desc_callback,
    )


# --------------------------------------------------------
# NEW process_single_image → replaced with AGENT call
# --------------------------------------------------------
def process_single_image(url, index, total, vision_deploy_name):
    """
    Calls the agent and ALWAYS returns a Document object.
    This guarantees downstream ingestion will not break.
    """

    print(f"[INFO] Processing image {index}/{total} → {url}")
    backoff = INITIAL_BACKOFF

    for attempt in range(1, MAX_RETRIES + 1):

        try:
            # Validate URL
            resp = requests.get(url, timeout=20)
            resp.raise_for_status()

            # Call the agent
            result = image_desc_agent(
                blob_url=url,
                vision_deploy_name=vision_deploy_name,
                index=index,
                total=total
            )

            # -------------------------------------------------------------------
            # NORMALIZE OUTPUT → ALWAYS A Document (required for add_ids_to_chunks)
            # -------------------------------------------------------------------
            cleaned_name = extract_clean_image_name(url)

            # CASE 1: Agent correctly returned Document
            if isinstance(result, Document):
                return result

            # CASE 2: Agent returned dict (tool responses sometimes do this)
            if isinstance(result, dict):
                text = result.get("text") or result.get("content") or str(result)
                return Document(
                    page_content=text,
                    metadata={"image_name": cleaned_name, "blob_url": url, "source_type": "image"}
                )

            # CASE 3: Agent returned string
            if isinstance(result, str):
                return Document(
                    page_content=result,
                    metadata={"image_name": cleaned_name, "blob_url": url, "source_type": "image"}
                )

            # CASE 4: Anything else (fallback)
            return Document(
                page_content=str(result),
                metadata={"image_name": cleaned_name, "blob_url": url, "source_type": "image"}
            )

        except Exception as e:
            print(f"[WARN] Attempt {attempt}/{MAX_RETRIES} failed for image {index} → {url}: {e}")

            if attempt == MAX_RETRIES:
                print(f"[ERROR] Giving up after {MAX_RETRIES} attempts → {url}")
                return Document(
                    page_content="OCR extraction failed.",
                    metadata={"image_name": extract_clean_image_name(url), "blob_url": url, "source_type": "image", "error": str(e)}
                )

            print(f"[INFO] Cooling down {backoff}s before retry…")
            time.sleep(backoff)
            backoff *= 2



# --------------------------------------------------------
# Multi-threaded loader (unchanged except for agent call inside)
# --------------------------------------------------------
def load_and_process_images(image_urls, vision_deploy_name=CHAT_DEPLOY):
    docs = []
    total = len(image_urls)

    print(f"[START] Processing {total} images with up to {MAX_THREADS} threads.\n")

    with ThreadPoolExecutor(max_workers=MAX_THREADS) as executor:

        futures = {
            executor.submit(process_single_image, url, idx + 1, total, vision_deploy_name): idx
            for idx, url in enumerate(image_urls)
        }

        for future in as_completed(futures):
            idx = futures[future]
            url = image_urls[idx]

            try:
                result = future.result()
                if result is not None:
                    docs.append(result)
                else:
                    print(f"[ERROR] Image {idx+1}/{total} → returned None")

            except Exception as e:
                print(f"[FATAL] Unexpected error for {url}: {e}")

    print(f"\n[COMPLETE] Finished processing {total} images.\n")
    return docs


# ingestion_tool.py (continued)
# --------------------------------------------------------
# SECTION 4 — Full Orchestration: run_full_ingestion()
# --------------------------------------------------------

def clear_cosmos_container(database_name, container_name):
    """
    Clears all items from a Cosmos DB container.
    EXACT logic from notebook.
    """
    client = CosmosClient(COSMOS_URL, credential=COSMOS_KEY)

    container = client.get_database_client(database_name).get_container_client(container_name)

    print(f"[INFO] Deleting all items in {database_name}.{container_name} ...")

    try:
        items = container.read_all_items()
        for item in items:
            try:
                container.delete_item(item=item, partition_key=item["id"])
            except Exception as e:
                print(f"[WARN] Failed deleting item {item.get('id')}: {e}")

        print("[SUCCESS] All documents deleted successfully!\n")

    except Exception as e:
        print(f"[ERROR] Could not list/delete items: {e}")




def get_or_create_vector_container_serverless(
    client: CosmosClient,
    DB_NAME: str,
    CONT_NAME: str,
    VECTOR_PATH: str,
    EMBED_DIM: int,
):
    """
    Serverless-safe Cosmos vector container creator.
    - Database must already exist
    - Creates container only if missing
    - Does NOT delete existing container
    - Does NOT set throughput (serverless limitation)
    """

    print("→ Using existing database:", DB_NAME)
    db = client.get_database_client(DB_NAME)

    try:
        db.read()
    except exceptions.CosmosResourceNotFoundError:
        raise RuntimeError(f"[FATAL] Database does not exist: {DB_NAME}")

    # ---- Vector embedding policy ----
    vector_embedding_policy = {
        "vectorEmbeddings": [
            {
                "path": VECTOR_PATH,
                "dataType": "float32",
                "dimensions": EMBED_DIM,
                "distanceFunction": "cosine",
            }
        ]
    }

    # ---- Indexing policy ----
    indexing_policy = {
        "includedPaths": [{"path": "/*"}],
        "excludedPaths": [
            {"path": "/\"_etag\"/?"},
            {"path": f"{VECTOR_PATH}/*"}
        ],
        "vectorIndexes": [
            {
                "path": VECTOR_PATH,
                "type": "quantizedFlat"
            }
        ],
    }

    # ---- Check container ----
    try:
        container = db.get_container_client(CONT_NAME)
        container.read()
        print("✔ Container exists:", CONT_NAME)
        return container

    except exceptions.CosmosResourceNotFoundError:
        print("→ Creating vector container (serverless-safe):", CONT_NAME)

        container = db.create_container(
            id=CONT_NAME,
            partition_key=PartitionKey(path="/id"),
            indexing_policy=indexing_policy,
            vector_embedding_policy=vector_embedding_policy
        )

        print("✔ Vector container created:", CONT_NAME)
        return container


from pathlib import Path
from typing import List, Dict, Any
from PyPDF2 import PdfReader
from PyPDF2.generic import IndirectObject
import fitz  # PyMuPDF
import base64
from dotenv import load_dotenv
import os
from pathlib import Path
import json
from openai import OpenAI
import os
import importlib
# importlib.reload(configs)
from openai import AzureOpenAI
 

def _resolve_indirect(obj):
    if isinstance(obj, IndirectObject):
        return obj.get_object()
    return obj


def _has_freetext_annotations(reader: PdfReader, max_pages: int = 3) -> bool:
    """Return True if any of the first `max_pages` pages contain /FreeText annotations.

    Many "editable"-looking PDFs (Fill & Sign / typewriter tools) store user-entered
    text as /FreeText annotations under each page's /Annots array. These are NOT
    AcroForm/XFA fields, so they won't appear under /Root -> /AcroForm.
    """

    try:
        pages = reader.pages
        n = min(len(pages), max_pages)

        for i in range(n):
            page = pages[i]
            annots = page.get("/Annots")
            annots = _resolve_indirect(annots) or []

            for a in annots:
                aobj = _resolve_indirect(a)
                subtype = aobj.get("/Subtype")
                # subtype is usually a NameObject like '/FreeText'
                if subtype and str(subtype) == "/FreeText":
                    return True
        return False
    except Exception:
        return False
        
def is_editable_form_pdf(path: Path) -> bool:
    """
    Return True if the PDF appears to contain *editable* content:
      - AcroForm fields (/Root -> /AcroForm -> /Fields)
      - XFA forms (/Root -> /AcroForm -> /XFA)
      - FreeText annotations (/Annots -> /Subtype /FreeText)

    Note: FreeText annotations are NOT form fields; they are an annotation layer.
    Ensures the file handle is closed (important on Windows) by using a 'with' block.
    """
    try:
        # Open the file explicitly so it gets closed after we're done
        with open(path, "rb") as f:
            reader = PdfReader(f)

            root = _resolve_indirect(reader.trailer.get("/Root"))
            if not root:
                return False

            acroform = root.get("/AcroForm")
            # If there's no AcroForm, it still may be "editable" via FreeText annotations.
            if not acroform:
                return _has_freetext_annotations(reader)

            acroform = _resolve_indirect(acroform)

            fields = acroform.get("/Fields")
            if isinstance(fields, IndirectObject):
                fields = fields.get_object()

            # If we have any form fields, it's editable
            if fields and len(fields) > 0:
                return True

            # Check for XFA forms as well
            xfa = acroform.get("/XFA")
            if isinstance(xfa, IndirectObject):
                xfa = xfa.get_object()

            if xfa:
                return True

            # No fields/XFA found. Fall back to FreeText annotations.
            return _has_freetext_annotations(reader)
    except Exception:
        return False

def flatten_and_get_images(input_path: str, output_path: str, dpi: int = 200):
    """
    Flattens the PDF AND returns a list of page images (pixmaps).
    Uses context managers so files are always closed, even on errors.
    """
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    pixmaps = []

    # Ensure both documents are closed reliably
    with fitz.open(input_path) as src_doc, fitz.open() as out_doc:
        for page in src_doc:
            pix = page.get_pixmap(matrix=mat)
            pixmaps.append(pix)

            new_page = out_doc.new_page(width=pix.width, height=pix.height)
            new_page.insert_image(new_page.rect, pixmap=pix)

        out_doc.save(output_path)

    return pixmaps

# ============================
# 3) Convert pixmaps to PNG paths
# ============================

def save_pixmaps_to_images(pixmaps: List[fitz.Pixmap], out_dir: Path, stem: str):
    out_dir.mkdir(parents=True, exist_ok=True)
    image_paths = []

    for i, pix in enumerate(pixmaps, start=1):
        img_path = out_dir / f"{stem}_page{i}.png"
        pix.save(str(img_path))
        image_paths.append(img_path)

    return image_paths


prompt_lm = """
You are reading a scanned "Client Information Sheet".

Extract ALL fields into a STRICT JSON object with this exact shape:

{
  "Applicant": {
    "Legal Entity Name": string or null,
    "DBA": string or null,
    "Street Address": string or null,
    "City, State, Postal Code, Country": string or null,
    "Phone Number": string or null,
    "Email": string or null,
    "Contacts": [
      {
        "Name": string or null,
        "Role": string or null,
        "Phone Number": string or null,
        "Email": string or null
      }
    ]
  },
  "Bill-To": {
    "Legal Entity Name": string or null,
    "Street Address": string or null,
    "City, State, Postal Code, Country": string or null,
    "Accounts Payable Contact": string or null,
    "Phone Number": string or null,
    "Email": string or null
  },
  "Manufacturer": {
    "Legal Entity Name": string or null,
    "Street Address": string or null,
    "City, State, Postal Code, Country": string or null,
    "Contacts": [
      {
        "Name": string or null,
        "Role": string or null,
        "Phone Number": string or null,
        "Email": string or null
      }
    ],
    "Estimated Production Date": string or null,
    "Labeling Method": string or null
  },
  "Completed By": string or null,
  "Dates": {
    "Form Completion": string or null
  },
  "Signatures": string or null
}

Rules:
- DO NOT put fields like "Legal Entity Name" at the root.
- All applicant fields must be inside "Applicant".
- All bill-to fields must be inside "Bill-To".
- All manufacturer/production facility fields must be inside "Manufacturer".
- Use null if something is missing or unreadable.
- Return ONLY JSON, no extra text.
"""


def image_to_data_uri(path: Path) -> str:
    b = path.read_bytes()
    b64 = base64.b64encode(b).decode("utf-8")
    return f"data:image/png;base64,{b64}"


def extract_page_with_llm(img_path: Path) -> str:
    """
    Sends a single PNG page to GPT-4.1 (vision-enabled)
    and extracts structured JSON from the form.
    """
    data_uri = image_to_data_uri(img_path)

    prompt = prompt_lm

    response = aoai_client.chat.completions.create(
        model="gpt-4.1",
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": data_uri}},
                ],
            }
        ]
    )

    return response.choices[0].message.content




def process_pdfs(
    src_root: str = "src_files",
    images_root: str = "page_images",
    dpi: int = 200,
) -> Dict[str, Any]:

    src_root = Path(src_root)
    images_root = Path(images_root)

    images_root.mkdir(exist_ok=True)

    results: Dict[str, Any] = {}

    print(f"Scanning {src_root.resolve()} ...")

    for pdf_path in src_root.glob("*.pdf"):
        if not pdf_path.is_file():
            continue

        print(f"\nChecking {pdf_path.name}...")

        if not is_editable_form_pdf(pdf_path):
            print(" → Not editable. Skipping.")
            continue

        print(" → Editable PDF detected.")

        # Extract images directly from original PDF
        pixmaps = []

        with fitz.open(pdf_path) as doc:
            zoom = dpi / 72.0
            mat = fitz.Matrix(zoom, zoom)

            for page in doc:
                pix = page.get_pixmap(matrix=mat)
                pixmaps.append(pix)

        # Save images
        img_dir = images_root / pdf_path.stem
        img_paths = save_pixmaps_to_images(pixmaps, img_dir, pdf_path.stem)

        # OCR / LLM extraction
        llm_outputs = []
        for img_path in img_paths:
            print(f"   → Sending {img_path.name} to GPT-4.1...")
            output = extract_page_with_llm(img_path)
            llm_outputs.append(output)

        # results[pdf_path.name] = {
        #     "original": pdf_path,
        #     "images": img_paths,
        #     "extracted": llm_outputs,
        # }
        results[pdf_path.name] = {
            "original": pdf_path,
            "images": img_paths,
            "extracted": llm_outputs,
            "is_editable": True
        }


    return results


def extract_cis(src_root, images_root):
    results = process_pdfs(
        src_root=src_root,
        images_root=images_root,
        dpi=200
    )

    all_cis = []
    editable_pdfs = []

    for pdfs in results.keys():
        dic_cis = {}
        dic_cis['filename'] = pdfs
        dic_cis['text'] = results[pdfs]['extracted'][0]
        all_cis.append(dic_cis)

        editable_pdfs.append(pdfs)

    return all_cis, editable_pdfs

# with open("src_files\\all_cis_info.txt", "w", encoding="utf-8") as f:
#     json.dump(all_cis, f, indent=4, default=str)

def copy_extracted_images_to_src(page_images_root: str, src_root: str):
    """
    Copies all extracted page images into src_files directory
    so they get ingested like normal uploaded images.
    """

    page_images_root = Path(page_images_root)
    src_root = Path(src_root)

    src_root.mkdir(exist_ok=True)

    copied = 0

    for subdir in page_images_root.iterdir():
        if not subdir.is_dir():
            continue

        for img in subdir.glob("*.png"):
            dest = src_root / img.name

            # avoid overwrite
            if dest.exists():
                continue

            shutil.copy2(img, dest)
            copied += 1

    print(f"✅ Copied {copied} extracted page images into {src_root}")


def append_cis_images_to_image_metadata(images_root: str, image_page_metadata: list):
    """
    Adds CIS extracted page images into CAD schematic image pipeline
    so they get uploaded + OCR + embedded exactly the same way.
    """

    images_root = Path(images_root)

    for subdir in images_root.iterdir():
        if not subdir.is_dir():
            continue

        pdf_name = subdir.name + ".pdf"   # fake source name for metadata consistency

        for img in subdir.glob("*.png"):
            image_page_metadata.append({
                "pdf_file": pdf_name,
                "page": None,
                "local_image_path": str(img),
                "reason": "editable_pdf_page"
            })



def clean_text(text: str) -> str:
    if not text or not isinstance(text, str):
        return ""

    t = text

    # -------------------------------------------------
    # 1. Remove URLs (inline and bracketed)
    # -------------------------------------------------
    t = re.sub(r"<https?://[^>]+>", " ", t)
    t = re.sub(r"https?://\S+", " ", t)

    # -------------------------------------------------
    # 2. Remove common email headers / reply markers
    # -------------------------------------------------
    t = re.sub(
        r"(^|\n)(from|sent|to|cc|subject):.*",
        " ",
        t,
        flags=re.IGNORECASE
    )

    # -------------------------------------------------
    # 3. Remove legal / disclaimer style blocks
    #    (generic wording, multiline)
    # -------------------------------------------------
    disclaimer_patterns = [
        r"confidentiality notice.*",
        r"this email.*confidential.*",
        r"this message.*confidential.*",
        r"export (control|notification).*",
        r"unauthorized.*prohibited.*",
        r"intended recipient.*",
        r"do not (print|forward|distribute).*",
    ]

    for p in disclaimer_patterns:
        t = re.sub(p, " ", t, flags=re.IGNORECASE | re.DOTALL)

    # -------------------------------------------------
    # 4. Remove spam / scanner / tracker messages
    # -------------------------------------------------
    scanner_patterns = [
        r"scanned for (spam|viruses).*",
        r"click here.*",
        r"report this email.*",
        r"external sender.*",
    ]

    for p in scanner_patterns:
        t = re.sub(p, " ", t, flags=re.IGNORECASE | re.DOTALL)

    # -------------------------------------------------
    # 5. Remove excessive punctuation / separators
    # -------------------------------------------------
    t = re.sub(r"_+", " ", t)
    t = re.sub(r"-{3,}", " ", t)
    t = re.sub(r"={3,}", " ", t)

    # -------------------------------------------------
    # 6. Normalize whitespace
    # -------------------------------------------------
    t = re.sub(r"\n{3,}", "\n\n", t)
    t = re.sub(r"[ \t]{2,}", " ", t)

    # -------------------------------------------------
    # 7. Drop very short junk lines
    # -------------------------------------------------
    lines = []
    for line in t.splitlines():
        stripped = line.strip()
        if len(stripped) < 10:
            continue
        if stripped.lower() in {"click here", "here", "thanks", "thank you"}:
            continue
        lines.append(stripped)

    t = "\n".join(lines)

    return t.strip()


def normalize_whitespace_only(text: str) -> str:
    if not text:
        return ""

    t = text
    t = re.sub(r"\r\n", "\n", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    t = re.sub(r"[ \t]{2,}", " ", t)

    return t.strip()


def clean_extracted_texts(extracted_texts):
    """
    Applies clean_text safely based on file type.
    Preserves tables and structured content.
    """
    cleaned_items = []

    for item in extracted_texts:
        if not isinstance(item, dict):
            continue

        filename = item.get("filename", "unknown")
        text = item.get("text", "")

        if not text:
            continue

        fname = filename.lower()

        # --------------------------------------------
        # Email-like files → aggressive cleanup
        # --------------------------------------------
        if fname.endswith((".eml", ".msg")):
            cleaned_text = clean_text(text)

        # --------------------------------------------
        # Spreadsheet / table text → VERY LIGHT cleanup
        # --------------------------------------------
        elif fname.endswith((".xlsx", ".xls", ".csv")):
            cleaned_text = normalize_whitespace_only(text)

        # --------------------------------------------
        # Everything else → moderate cleanup
        # --------------------------------------------
        else:
            cleaned_text = clean_text(text)

        cleaned_items.append({
            "filename": filename,
            "text": cleaned_text
        })

    return cleaned_items



def ingest_files_from_blob_urls_create_embeddings(download_dir,blob_urls: list, project_id: str,
textDB_container_name, imageDB_container_name, keep_files: bool = True, verbose: bool = True):
    """
    Single-call function that performs the full notebook ingestion pipeline.
    - blob_urls: list of full blob URLs 
    - download_dir: local directory to store downloads and converted pdfs
    - keep_files: whether to keep local downloaded files
    - verbose: print progress

    NOTE: This function will DELETE ALL DOCUMENTS in the target Cosmos container (as per notebook).
    """

    print('######### download_dir ########', download_dir)

    print('######## blob_urls #########', blob_urls)

    print('######## project_id #########', project_id)

    print('######## textDB_container_name #########', textDB_container_name)

    print('######## imageDB_container_name #########', imageDB_container_name)





    # 1) Cosmos client and container
    # client = CosmosClient(COSMOS_URL, credential=COSMOS_KEY, consistency_level=ConsistencyLevel.Eventual)
    # db_client = client.get_database_client(COSMOS_DB_TEXT)
    # container_client = db_client.get_container_client(textDB_container_name)

    client = CosmosClient(COSMOS_URL, credential=COSMOS_KEY)

    container_client = get_or_create_vector_container_serverless(
        client=client,
        DB_NAME=COSMOS_DB_TEXT,
        CONT_NAME=textDB_container_name,
        VECTOR_PATH=VECTOR_PATH,
        EMBED_DIM=EMBED_DIM
    )

    # 2) Delete all items (not reversible)
    try:
        items = container_client.read_all_items()
        for item in items:
            try:
                container_client.delete_item(item=item, partition_key=item["id"])
            except Exception as e:
                print(f"[WARN] Failed to delete item {item.get('id')}: {e}")
        print("All documents deleted successfully!")
    except Exception as e:
        print(f"[WARN] Could not enumerate/delete items: {e}")

    BASE_DIR = Path(__file__).resolve().parent.parent
    DOWNLOAD_DIR = BASE_DIR / "data" / project_id / "src_files"
    IMAGES_ROOT = BASE_DIR / "data" / project_id / "image_files"

    # 3) Process blob URLs (download/convert/extract)
    extracted_texts, image_urls_raw, downloaded_pdf_paths, converted_pdf_paths = process_blob_urls_2(
        blob_urls, AZURE_CONN_STRING, BLOB_CONTAINER,
        download_dir=DOWNLOAD_DIR, keep_files=keep_files, verbose=verbose
    )

    pdf_paths = downloaded_pdf_paths + converted_pdf_paths

    print(f"[INFO] Extracted text files: {len(extracted_texts)}")
    print(f"[INFO] Initial image URLs: {len(image_urls_raw)}")
    print(f"[INFO] Total PDFs: {len(pdf_paths)}\n")

    cis_info, editable_pdfs = extract_cis(src_root=DOWNLOAD_DIR,images_root=IMAGES_ROOT)

    # ✅ Remove editable PDFs from PDF ingestion list (they are handled via OCR images)
    pdf_paths = [
        p for p in pdf_paths
        if os.path.basename(p) not in editable_pdfs
    ]
    
    copy_extracted_images_to_src(
    page_images_root=IMAGES_ROOT,
    src_root=DOWNLOAD_DIR)

    extracted_texts=clean_extracted_texts(extracted_texts)
    
    extracted_texts += cis_info

    # print("###################### CIS INFO ##################################")
    # print(cis_info)
    # print("###################################################################")


    print("\n======================================")
    print("   STEP 2 — LOAD + SPLIT PDF TEXT      ")
    print("======================================\n")

    chunks, image_page_metadata = load_and_split_pdfs_text(
        pdf_paths,
        CHUNK_SIZE,
        CHUNK_OVERLAP,
        extracted_texts=extracted_texts,
        cad_schematics=ENABLE_CAD_SCHEMATICS,
    )

    print(f"[INFO] Total text chunks produced: {len(chunks)}")
    print(f"[INFO] CAD/Schematic pages extracted: {len(image_page_metadata)}\n")

    # ✅ Add CIS extracted images into CAD schematic pipeline
    append_cis_images_to_image_metadata(
        images_root=IMAGES_ROOT,
        image_page_metadata=image_page_metadata
    )


    print("\n======================================")
    print("   STEP 3 — CREATE TEXT VECTOR STORE   ")
    print("======================================\n")

    # Clear DB (EXACT notebook logic)
    clear_cosmos_container(COSMOS_DB_TEXT, textDB_container_name)

    vectorstore_text = build_vectorstore_text(textDB_container_name)

    # Assign UUIDs to each chunk
    chunks_uuid = add_ids_to_chunks(chunks)

    # Ingest in parallel
    ingest_to_cosmos_parallel(vectorstore_text, chunks_uuid, batch_size=10, max_workers=10)

    print("\n[SUCCESS] Text ingestion completed.\n")


    print("\n=====================================================")
    print("   STEP 4 — UPLOAD CAD/SCHEMATIC PAGE IMAGES         ")
    print("=====================================================\n")

    # Deduplicate local image paths
    seen = set()
    unique_metadata = []

    for item in image_page_metadata:
        key = item["local_image_path"]
        if key not in seen:
            seen.add(key)
            unique_metadata.append(item)

    image_urls = upload_pdf_images_and_append_urls(
        image_page_metadata=unique_metadata,
        image_urls=image_urls_raw,
        conn_str=AZURE_CONN_STRING,
        container=BLOB_CONTAINER,
        max_workers=8,
        project_id=project_id   
    )


    # Turn into flat list
    img_links = extract_urls(image_urls)

    print(f"[INFO] Total images for OCR (blob URLs): {len(img_links)}\n")


    print("\n==============================================")
    print("   STEP 5 — IMAGE OCR USING AGENT PIPELINE     ")
    print("==============================================\n")

    docs_image = load_and_process_images(img_links, vision_deploy_name=CHAT_DEPLOY)

    print(f"[SUCCESS] Finished OCR for {len(docs_image)} images.\n")


    print("\n==============================================")
    print("   STEP 6 — CREATE IMAGE VECTOR STORE          ")
    print("==============================================\n")

    image_client = CosmosClient(COSMOS_URL, credential=COSMOS_KEY)

    container_client_image = get_or_create_vector_container_serverless(
        client=image_client,
        DB_NAME=COSMOS_DB_IMAGE,
        CONT_NAME=imageDB_container_name,
        VECTOR_PATH=VECTOR_PATH,
        EMBED_DIM=EMBED_DIM
    )

    # Clear DB (EXACT notebook logic)
    clear_cosmos_container(COSMOS_DB_IMAGE, imageDB_container_name)

    vectorstore_image = build_vectorstore_image(imageDB_container_name)

    # Assign UUIDs to image docs
    docs_image_uuid = add_ids_to_chunks(docs_image)

    # Ingest
    ingest_to_cosmos_parallel(vectorstore_image, docs_image_uuid, batch_size=10, max_workers=10)

    print("\n[SUCCESS] Image ingestion completed.\n")


    print("\n==============================================")
    print("         FULL INGESTION PIPELINE DONE         ")
    print("==============================================\n")

    
    return {
            "project_id": project_id,
            "image_urls": img_links,
            "pdf_paths": pdf_paths,
            "downloaded_pdfs": downloaded_pdf_paths,
            "converted_pdfs": converted_pdf_paths,
            "image_page_metadata": image_page_metadata,
            "chunks_count": len(chunks)
        }



